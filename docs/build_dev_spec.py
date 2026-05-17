#!/usr/bin/env python3
"""Generate Health_Interop_Agentic_Framework_Dev_Spec.docx — MVP-driven.

Structure: MVP use cases upfront, requirements derived from use cases,
general/reference material in appendices.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)

# ── helpers ──────────────────────────────────────────────────────────

HEADER_BG = RGBColor(0x1a, 0x1a, 0x2e)
HEADER_FG = RGBColor(0xff, 0xff, 0xff)

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph(style='Normal')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = HEADER_FG
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto',
            qn('w:fill'): '1a1a2e',
        })
        shading.append(bg)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if ri % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                bg = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear', qn('w:color'): 'auto',
                    qn('w:fill'): 'f5f5fa',
                })
                shading.append(bg)
    doc.add_paragraph()
    return table

def add_note(text, label="Note"):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x52, 0x9B)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x00, 0x52, 0x9B)
    return p


# ═════════════════════════════════════════════════════════════════════
# TITLE
# ═════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Health Interoperability\nAgentic Framework')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Product Requirements Specification')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'This document defines WHAT the framework must enable and WHY,\n'
    'organized around three MVP use cases that drive every requirement.\n'
    'Implementation decisions (the HOW) are left to the engineering team.\n\n'
    'May 2026'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════
# HOW TO READ
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('How to Read This Document', level=1)

add_para(
    'This document uses three labels to classify every capability, '
    'property, and concept:'
)

add_table(
    ['Label', 'Meaning'],
    [
        ['Framework (exists)',
         'Already exists in the %AI Framework as shipped in IRIS for '
         'Health 2026.2. Verified and working. No new development needed.'],
        ['Framework (proposed)',
         'Does not exist in the current %AI Framework. This document '
         'proposes it as a product requirement. The engineering team '
         'decides how to implement it.'],
        ['Application-level',
         'Something application developers or end-user builders create '
         'on top of the framework. The framework must enable it but does '
         'not ship it. The Health Interop Copilot (Appendix C) provides '
         'concrete examples.'],
    ]
)

add_para(
    'The document is structured in four tiers:'
)
add_bullet(
    'Section 1-2: Product overview, personas, and design principles.',
    bold_prefix='Context. '
)
add_bullet(
    'Section 3: The three use cases that define the MVP. '
    'Every requirement in this document traces back to at least one of these.',
    bold_prefix='MVP Use Cases. '
)
add_bullet(
    'Section 4-7: Framework capabilities required to deliver the MVP use cases. '
    'Each requirement states what is needed, why, which use case drives it, and '
    'whether the framework already provides it.',
    bold_prefix='MVP Requirements. '
)
add_bullet(
    'Appendices A-D: Framework class catalog, general requirements beyond MVP, '
    'reference implementation details, and glossary.',
    bold_prefix='Reference. '
)


# ═════════════════════════════════════════════════════════════════════
# 1. PRODUCT OVERVIEW
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('1. Product Overview', level=1)

add_para(
    'The Health Interoperability Agentic Framework is an AI orchestration '
    'layer embedded in InterSystems IRIS for Health and Health Connect. It '
    'enables integration engineers to build, review, validate, optimize, '
    'and operate interoperability solutions through natural-language '
    'interaction, while keeping all state-changing operations under '
    'explicit user control.'
)
add_para(
    'The framework composes five runtime concepts: Agents, MCP Servers, '
    'ToolSets, Tools, and Skills, over a grounded Catalog. Agents reason '
    'and plan. Tools execute. MCP Servers group tools into functional '
    'domains and enforce boundaries. Catalogs and Skills supply the '
    'grounded context that prevents the agent from fabricating class names, '
    'paths, or references that do not exist in the target IRIS instance.'
)

doc.add_heading('1.1 Two Personas', level=2)

add_table(
    ['Persona', 'Where They Work', 'What They Do'],
    [
        ['Developer (DX)',
         'VS Code, any IDE, command line',
         'Author Tool classes in ObjectScript or Python. Write Skill '
         'documents. Build catalog embeddings. Package as IPM modules. '
         'Developers define what the copilot CAN do.'],
        ['Builder (End User)',
         'IRIS Management Portal, browser UI',
         'Create Agents with system prompts. Assemble MCP Servers from '
         'ToolSets. Attach Skills. Configure LLM providers. Chat with '
         'the copilot. Builders decide HOW the copilot BEHAVES for '
         'their use case.'],
    ]
)

add_para(
    'A builder should never need to open an IDE. A developer should '
    'never need to configure an agent through the Management Portal '
    'to test their tools.'
)

doc.add_heading('1.2 Design Principles', level=2)

add_table(
    ['Principle', 'What It Means'],
    [
        ['Catalog-grounded generation',
         'Agents never invent class names, business hosts, or DTL '
         '(Data Transformation Language) paths. All references come '
         'from the Catalog via vector search or class introspection.'],
        ['Explicit execution',
         'Nothing runs without an explicit Tool call. The Agent plans; '
         'the Tool executes. No hidden side-effect path.'],
        ['Human gating',
         'Any Tool that changes IRIS state pauses for explicit user '
         'approval before executing.'],
        ['Namespace isolation',
         'Every Tool runs in the namespace of the authenticated '
         'session. Cross-namespace operations require explicit grants.'],
        ['Provider independence',
         'Swapping LLM providers does not require changes to tools, '
         'skills, or MCP definitions.'],
        ['Extensible by design',
         'New tools, skills, MCPs, and catalogs can be added by '
         'developers without modifying framework code.'],
        ['Enterprise observability',
         'Every tool invocation is captured in an immutable audit '
         'record. Streaming events show real-time progress.'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# 2. CONTAINMENT HIERARCHY
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('2. Architecture: Containment Hierarchy', level=1)

add_para(
    'The runtime forms a strict containment hierarchy. An Agent owns '
    'its conversation context and decides what to do next. To act, '
    'the Agent issues a Tool call against a ToolSet exposed by an '
    'MCP Server. Skills inject domain knowledge; the Catalog supplies '
    'grounded references to real IRIS artifacts.'
)

add_para(
    'Agent\n'
    '  +-- MCP Servers\n'
    '  |     +-- ToolSets\n'
    '  |           +-- Tools\n'
    '  +-- Skills\n'
    '  +-- Catalog Access'
)

add_para(
    'Agents never modify IRIS state directly. All mutations are '
    'performed by Tools, which apply input validation, namespace '
    'checks, approval gates, and audit logging before touching IRIS.'
)

add_para(
    'All components are loosely coupled: Agents do not know how Tools '
    'execute; MCP Servers do not know how Agents reason; Tools do not '
    'know which Agent invoked them.'
)


# ═════════════════════════════════════════════════════════════════════
# 3. MVP USE CASES
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('3. MVP Use Cases', level=1)

add_para(
    'The MVP is defined by three use cases. Every requirement in this '
    'document traces back to at least one of them. If a capability is '
    'not needed by any of these three use cases, it is deferred to '
    'Appendix B (General Requirements Beyond MVP).'
)

# ── UC1 ──────────────────────────────────────────────────────────────

doc.add_heading('3.1 Use Case 1: Guided Production Build', level=2)

add_para(
    'The user asks: "I need to receive ADT admission messages from '
    'our HIS and send observation reports to the downstream LIS."'
)

doc.add_heading('Step 1: Prompt Refinement', level=3)
add_para(
    'The agent identifies missing information and asks targeted '
    'questions: What HL7 version? What transport (MLLP port, file '
    'drop, TCP)? What fields need to map from ADT to ORU? What '
    'happens to messages that fail transformation? What throughput '
    'is expected? The agent guides the user to a complete '
    'specification without requiring them to know IRIS terminology.'
)

doc.add_heading('Step 2: Catalog Search and Plan', level=3)
add_para(
    'The agent searches the Catalog (a vector database of IRIS '
    'classes) to find the right business host classes, adapter classes, '
    'and existing transformation templates. It presents a structured '
    'plan: every component name, the class it will use (citing which '
    'catalog result and why), key settings, data flow, field mappings, '
    'and the complete test message. The plan ends with "Ready to build. '
    'Shall I proceed?"'
)

doc.add_heading('Step 3: Build', level=3)
add_para(
    'On approval, the agent executes in a strict sequence, calling '
    'tools to build the production end to end. For complex productions '
    'this includes:'
)
add_bullet(
    'DTL (Data Transformation Language) with complete field mappings, '
    'foreach blocks for repeating segments, conditional logic, and '
    'subtransform references'
)
add_bullet(
    'BPL (Business Process Language) for orchestration workflows '
    'with compensation handlers, async callbacks, and code activities'
)
add_bullet(
    'Routing Rules with constraint-based message routing, transform '
    'chaining, and dead-letter handling'
)
add_bullet(
    'Lookup Tables for code mappings (facility codes, provider '
    'identifiers, insurance plan maps)'
)
add_bullet(
    'HL7 Schema validation ensuring source and target structure '
    'paths are correct before any mapping'
)
add_bullet(
    'FHIR R4 resource construction for bundles, references, and '
    'search parameter configuration'
)
add_bullet(
    'SDA (Summary Document Architecture) as the intermediary hub '
    'for cross-format transformations (HL7 to SDA to FHIR)'
)

add_para(
    'The tools must understand these artifacts as a connected system, '
    'not isolated components. A routing rule references a DTL by class '
    'name, a DTL references lookup tables by name, a BPL calls '
    'sub-DTLs and routes to specific business hosts. The build '
    'sequence enforces these dependencies.'
)

doc.add_heading('Step 4: Test and Validate', level=3)
add_para(
    'The agent creates sample messages and sends them to the '
    'production to check: production exists and is running, all hosts '
    'are enabled, no Event Log errors, messages flowed through the '
    'pipeline. Failed checks are corrected and validation re-run. '
    'The completion report shows the test message sent, transformation '
    'output, processing time, and any errors encountered.'
)

doc.add_heading('Step 5: Catalog Update', level=3)
add_para(
    'After a successful build, any new DTL, BPL, or routing rule '
    'classes created during the session are automatically indexed '
    'into the Catalog. This ensures the next user who asks for a '
    'similar transformation can discover and reuse what was just '
    'built. The catalog grows with every successful build.'
)

doc.add_heading('What This Use Case Requires from the Framework', level=3)

add_table(
    ['Capability', 'Why UC1 Needs It', 'Status'],
    [
        ['Catalog vector search',
         'The agent must find the right EnsLib.* business host and '
         'adapter classes for the user\'s transport and message type. '
         'Training data does not know which classes exist on the '
         'user\'s specific IRIS instance.',
         'Framework (exists via %AI.RAG.KnowledgeBase)'],
        ['Plan-and-approve flow',
         'The agent must present a structured plan and wait for user '
         'approval before executing any build operations.',
         'Framework (proposed) - no built-in approval gate on tools'],
        ['Multi-tool sequencing',
         'A production build requires 5-10 tool calls in a specific '
         'order (create production, add hosts, create DTL, compile, '
         'create routing rule, start, validate).',
         'Framework (exists) - agent turn loop handles sequencing'],
        ['Approval gates for mutations',
         'Creating a production, adding a business host, and starting '
         'a production are state-changing operations that must not '
         'execute without user consent.',
         'Framework (proposed) - RequiresConfirmation flag needed'],
        ['Post-build validation',
         'The agent must verify the build succeeded before claiming '
         'success. Production running, hosts enabled, no errors, '
         'messages flowed.',
         'Application-level tool (framework must support the pattern)'],
        ['Catalog auto-growth',
         'New DTLs and routing rules created during the build must be '
         'indexed so future sessions can discover them.',
         'Framework (proposed) - no auto-reindex after artifact creation'],
        ['Streaming with tool cards',
         'A 90-second multi-step build must show real-time progress. '
         'Each tool call renders as a visible card in the chat UI.',
         'Framework (proposed) - no tool lifecycle event streaming'],
        ['Iteration limits and token budget',
         'A confused agent must not loop indefinitely. The build must '
         'complete within a defined iteration and token ceiling.',
         'Framework (proposed) - no agent-level caps'],
        ['Domain skills',
         'The agent needs deep knowledge of Productions, DTL, BPL, '
         'Routing Rules, HL7 schemas, and SDA to build correctly.',
         'Framework (exists via %AI.Agent.Skill, pending %OnNew fix)'],
    ]
)


# ── UC2 ──────────────────────────────────────────────────────────────

doc.add_heading('3.2 Use Case 2: Production Review and Optimization', level=2)

add_para(
    'The user asks: "Review the LAB.Production and tell me what it '
    'does and how to improve it."'
)

doc.add_heading('What the Agent Does', level=3)

add_para(
    'The agent calls tools to inspect every business host in the '
    'production: its class, settings, connections, and role in the data '
    'flow. It queries the Event Log for recent error patterns and '
    'throughput data. It then explains in plain language:'
)
add_bullet('What each host does and the data flow from source to destination')
add_bullet('Which settings are at defaults versus customized')
add_bullet(
    'Which features are unused (pool sizing, retry intervals, alert '
    'triggers, archive flags)'
)
add_bullet(
    'Prioritized recommendations, each citing the specific setting, '
    'its current value, the recommended value, and the expected impact'
)

add_para(
    'Example recommendation: "Your TCP service uses the default '
    'StayConnected=0 which reopens the socket on every message. For '
    'a high-volume HL7 feed, setting StayConnected=-1 keeps the '
    'connection open and reduces latency."'
)

doc.add_heading('What This Use Case Requires from the Framework', level=3)

add_table(
    ['Capability', 'Why UC2 Needs It', 'Status'],
    [
        ['Production introspection tools',
         'The agent must read a production\'s full definition: every '
         'business host, class, settings (distinguishing host-level '
         'from adapter-level), and connections.',
         'Application-level tools (framework must support %AI.Tool)'],
        ['Event Log and monitoring tools',
         'The agent must query Ens.Util.Log for error patterns and '
         'Ens.MessageHeader for throughput data.',
         'Application-level tools'],
        ['Catalog-grounded recommendations',
         'Recommendations must reference real IRIS features, not '
         'generic advice. The catalog provides class descriptions '
         'and key settings documentation.',
         'Framework (exists via %AI.RAG.KnowledgeBase)'],
        ['Read-only tool safety',
         'Review is entirely read-only. The framework must distinguish '
         'read-only tools (no approval needed) from mutating tools.',
         'Framework (proposed) - MutatingOperation flag needed'],
        ['Skills for best practices',
         'The agent needs knowledge of production best practices, '
         'common anti-patterns, and setting-level recommendations.',
         'Framework (exists via %AI.Agent.Skill, pending %OnNew fix)'],
    ]
)


# ── UC3 ──────────────────────────────────────────────────────────────

doc.add_heading('3.3 Use Case 3: Complex HL7-to-HL7 Transformations', level=2)

add_para(
    'The user needs to transform HL7 v2.5 ADT^A01 messages into '
    'ORU^R01 observation reports with field mappings that span '
    'multiple segments.'
)

doc.add_heading('What the Agent Does', level=3)

add_para(
    'The agent retrieves HL7 schema maps for both source (ADT_A01) '
    'and target (ORU_R01) structures BEFORE writing any DTL. This is '
    'critical because ADT_A01 has flat PID paths while ORU_R01 has '
    'nested paths like PIDgrpgrp(1).PIDgrp.PID. Guessed paths produce '
    'silent empty output.'
)
add_para(
    'The agent builds the complete DTL XML in a single tool call with '
    'all field mappings. It then runs a test with a sample message to '
    'verify every mapped field produces output. If any field is empty, '
    'the agent checks the schema paths and corrects them.'
)
add_para(
    'For complex scenarios (conditional logic, repeating segments, '
    'lookup tables), the agent uses DTL skill knowledge to generate '
    'foreach blocks, subtransform references, and lookup table entries.'
)

doc.add_heading('What This Use Case Requires from the Framework', level=3)

add_table(
    ['Capability', 'Why UC3 Needs It', 'Status'],
    [
        ['HL7 schema introspection tools',
         'The agent must look up exact segment paths for any HL7 '
         'message type and version. ADT_A01 and ORU_R01 have '
         'fundamentally different path structures. Guessing paths '
         'guarantees empty output.',
         'Application-level tools'],
        ['One-shot DTL creation',
         'The complete DTL XML must be built in a single tool call. '
         'Iterative updates burn rate-limited API tokens and produce '
         'incomplete transforms.',
         'Application-level tools'],
        ['DTL dry-run and testing',
         'After creation, the agent must run a test message through '
         'the DTL and verify field-level output. This requires a '
         'tool that executes a DTL without a running production.',
         'Application-level tools'],
        ['Lookup table management',
         'Real-world transformations depend on lookup tables for '
         'value translation (facility codes, insurance plans). The '
         'agent must create and populate these.',
         'Application-level tools'],
        ['DTL skill knowledge',
         'The agent needs deep knowledge of DTL syntax, virtual '
         'document paths, foreach patterns, subtransforms, and '
         'common pitfalls.',
         'Framework (exists via %AI.Agent.Skill, pending %OnNew fix)'],
        ['Anti-fabrication enforcement',
         'The agent must not claim a transformation works unless the '
         'test tool confirmed field-level output.',
         'Application-level (agent prompt + audit trail)'],
    ]
)


# ── UC TRACEABILITY ──────────────────────────────────────────────────

doc.add_heading('3.4 Use Case Traceability Matrix', level=2)

add_para(
    'This matrix maps every MVP capability to the use cases that '
    'require it. Capabilities needed by multiple use cases are '
    'highest priority.'
)

add_table(
    ['Capability', 'UC1\nBuild', 'UC2\nReview', 'UC3\nTransform', 'Status'],
    [
        ['Catalog vector search', 'Yes', 'Yes', 'No', 'Framework (exists)'],
        ['Approval gates (RequiresConfirmation)', 'Yes', 'No', 'Yes', 'Framework (proposed)'],
        ['MutatingOperation flag', 'Yes', 'Yes', 'Yes', 'Framework (proposed)'],
        ['Iteration limits / Token budget', 'Yes', 'Yes', 'Yes', 'Framework (proposed)'],
        ['Streaming tool lifecycle events', 'Yes', 'Yes', 'Yes', 'Framework (proposed)'],
        ['Skill sub-agents (%OnNew fix)', 'Yes', 'Yes', 'Yes', 'Framework (bug fix)'],
        ['%AI.INC visibility', 'Yes', 'Yes', 'Yes', 'Framework (bug fix)'],
        ['Post-build validation pattern', 'Yes', 'No', 'Yes', 'Application-level'],
        ['Catalog auto-growth', 'Yes', 'No', 'Yes', 'Framework (proposed)'],
        ['Production CRUD tools', 'Yes', 'Yes', 'No', 'Application-level'],
        ['DTL/BPL/Routing Rule tools', 'Yes', 'No', 'Yes', 'Application-level'],
        ['HL7 schema introspection tools', 'Yes', 'No', 'Yes', 'Application-level'],
        ['Monitoring / Event Log tools', 'No', 'Yes', 'No', 'Application-level'],
        ['Test message tools', 'Yes', 'No', 'Yes', 'Application-level'],
        ['Production skill', 'Yes', 'Yes', 'No', 'Application-level'],
        ['DTL skill', 'Yes', 'No', 'Yes', 'Application-level'],
        ['HL7v2 skill', 'Yes', 'Yes', 'Yes', 'Application-level'],
        ['Conversation persistence', 'Yes', 'Yes', 'Yes', 'Framework (exists)'],
        ['Namespace isolation', 'Yes', 'Yes', 'Yes', 'Framework (exists)'],
        ['Audit trail', 'Yes', 'Yes', 'Yes', 'Application-level (via Policy)'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# 4. MVP FRAMEWORK REQUIREMENTS
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('4. MVP Framework Requirements', level=1)

add_para(
    'This section lists the framework capabilities required to deliver '
    'the three MVP use cases. Each requirement states what is needed, '
    'why, and which use case drives it. Requirements already satisfied '
    'by the existing %AI Framework (see Appendix A) are noted but not '
    'repeated in detail.'
)


doc.add_heading('4.1 Agent Orchestration', level=2)

add_para(
    'The %AI.Agent class (Framework, exists) provides the core execution '
    'engine. All three MVP use cases require additional orchestration '
    'controls.'
)

add_table(
    ['Requirement', 'Status', 'Driven By', 'What Is Needed', 'Why'],
    [
        ['Declarative configuration',
         'Framework (exists)',
         'All UCs',
         '%AI.Agent supports Parameter and XData blocks. Subclasses '
         'declare their system prompt, tools, and provider binding.',
         'Builders need to configure agents without writing ObjectScript.'],
        ['Iteration limit',
         'Framework (proposed)',
         'All UCs',
         'An agent-level property that caps LLM turns per user request. '
         'When reached, the agent stops and returns what it has.',
         'UC1 involves 5-10 tool calls. Without a ceiling a confused LLM '
         'can loop indefinitely, consuming tokens and blocking the user.'],
        ['Token budget',
         'Framework (proposed)',
         'All UCs',
         'An agent-level property that caps total tokens (input + output '
         'across all turns). Reaching the budget triggers a graceful stop.',
         'A single UC1 build can consume 30,000+ tokens. Without a budget '
         'costs are unpredictable.'],
        ['Parallel tool batching',
         'Framework (proposed)',
         'UC1',
         'When the agent plans multiple independent tool calls in one '
         'turn, execute them concurrently.',
         'UC1 Step 2 often needs both a catalog search and a schema '
         'lookup simultaneously. Sequential execution doubles wait time.'],
        ['Approval policy',
         'Framework (proposed)',
         'UC1, UC3',
         'An agent-level property governing mutation approval: manual '
         '(all mutations need approval), auto (none), or hybrid (per '
         'tool). Default: manual.',
         'UC1 creates productions; UC3 creates DTLs. Both are '
         'state-changing and must pause for user consent. UC2 is '
         'read-only and never triggers approval.'],
    ]
)


doc.add_heading('4.2 Tool Metadata', level=2)

add_para(
    'The %AI.Tool class (Framework, exists) auto-discovers public '
    'ClassMethods and exposes them as tools. The following metadata '
    'properties are needed on every tool for orchestration, UI, and '
    'security. Properties marked Framework (exists) are already '
    'available; Framework (proposed) are new requirements.'
)

add_table(
    ['Property', 'Status', 'Driven By', 'What It Is', 'Why'],
    [
        ['Name',
         'Framework (exists)',
         'All UCs',
         'Method name on a %AI.Tool subclass. The LLM sees this name.',
         'Already auto-discovered.'],
        ['Description',
         'Framework (exists)',
         'All UCs',
         'Method Description comment. The LLM reads this to decide when '
         'to call the tool and what to expect.',
         'The contract between tool author and LLM. A vague description '
         'causes the LLM to misuse the tool.'],
        ['FormalSpec',
         'Framework (exists)',
         'All UCs',
         'ObjectScript parameter signature. The framework auto-generates '
         'a JSON Schema from this at compile time.',
         'Tool authors do not hand-write JSON Schemas.'],
        ['InputSchema',
         'Framework (exists)',
         'All UCs',
         'JSON Schema for input parameters. Auto-generated from FormalSpec.',
         'The orchestrator validates inputs before invocation.'],
        ['RequiresConfirmation',
         'Framework (proposed)',
         'UC1, UC3',
         'Boolean flag. When true, the orchestrator pauses and shows '
         'an approval gate in the UI. The user sees the tool name, what '
         'it will do, and the key arguments. Nothing executes until '
         'the user clicks Approve.',
         'CreateProduction, AddBusinessHost, CreateDTL, '
         'StartProduction all change IRIS state and must have user consent.'],
        ['MutatingOperation',
         'Framework (proposed)',
         'All UCs',
         'Boolean flag. True if the tool changes IRIS state (creates, '
         'updates, or deletes). Drives audit classification, UI '
         'treatment, and approval gating.',
         'UC2 is entirely read-only. The framework must distinguish '
         'GetProduction (read) from CreateProduction (write) for audit '
         'and safety.'],
        ['TimeoutSeconds',
         'Framework (proposed)',
         'UC1, UC3',
         'Maximum wall-clock seconds before the orchestrator cancels '
         'the tool.',
         'Different tools have different durations: catalog search '
         'finishes in 1 second; compiling a production may take 30. '
         'A stuck tool must not block the entire agent turn.'],
        ['Idempotent',
         'Framework (proposed)',
         'UC1',
         'Boolean flag. True means safe to retry on transient failure.',
         'ListProductions is safe to retry; CreateProduction is not '
         '(would create a duplicate). The orchestrator needs to know.'],
        ['Category',
         'Framework (proposed)',
         'All UCs',
         'Grouping label (e.g., "production", "transform", "testing"). '
         'Used by the UI and analytics.',
         'With 40+ tools the UI and audit reports need grouping by '
         'functional area.'],
        ['AllowedNamespaces',
         'Framework (proposed)',
         'UC1',
         'List of namespaces where this tool may run. Empty = all '
         'namespaces the user has access to.',
         'DeleteProduction should only run in development namespaces. '
         'ListProductions is safe everywhere.'],
    ]
)


doc.add_heading('4.3 Tool Execution Lifecycle', level=2)

add_para(
    'Every tool invocation should traverse a deterministic state '
    'machine. State transitions are emitted as streaming events '
    'so the UI can render tool cards in real time. This lifecycle '
    'is Framework (proposed). All three UCs need it for transparency.'
)

add_table(
    ['State', 'Meaning', 'Next States'],
    [
        ['Planned',
         'Agent has selected this tool but not yet invoked it.',
         'AwaitingApproval, Queued, Cancelled'],
        ['AwaitingApproval',
         'Tool has RequiresConfirmation=true. Execution paused '
         'for user approval.',
         'Queued, Cancelled'],
        ['Queued',
         'Approved (or auto-approved). Waiting for execution slot.',
         'Running, Cancelled'],
        ['Running',
         'Tool dispatched. Execution in progress.',
         'Succeeded, Failed, Retrying, Cancelled'],
        ['Succeeded',
         'Completed. Output matches expectations.',
         'Terminal'],
        ['Failed',
         'Unrecoverable error.',
         'Terminal'],
        ['Retrying',
         'Transient error on an Idempotent tool. Re-invoking.',
         'Running, Failed'],
        ['Cancelled',
         'User rejected, timeout exceeded, or agent cancelled.',
         'Terminal'],
    ]
)


doc.add_heading('4.4 Skill Sub-Agents', level=2)

add_para(
    'The %AI.Agent.Skill class (Framework, exists) provides declarative '
    'sub-agents. All three MVP use cases depend on skills for domain '
    'knowledge. Two blocking issues must be resolved.'
)

add_table(
    ['Issue', 'Status', 'Driven By', 'What Is Needed', 'Why'],
    [
        ['%OnNew JSON marshaling bug',
         'Framework (bug fix)',
         'All UCs',
         '%AI.Agent.Skill.%OnNew passes a %DynamicObject to $ZF(-6) '
         'instead of calling %ToJSON() first. The Rust bridge receives '
         'the OREF integer instead of a JSON string. Compare with '
         '%AI.Provider which correctly calls .%ToJSON() before $ZF.',
         'Without this fix, no Skill subclass can be instantiated. '
         'The entire declarative-skill mechanism is unusable. The '
         'reference implementation workaround (a Base class overriding '
         '%OnNew) is fragile and relies on inlined macro values.'],
        ['%AI.INC include file visibility',
         'Framework (bug fix)',
         'All UCs',
         'The %AI.INC include file (macros like $$$IrisLLMLibrary = '
         '1042, $$$LLMBUILDSKILLFROMJSON = 66) is only accessible from '
         '%SYS. Classes in user namespaces cannot use '
         '[ IncludeCode = %AI ].',
         'Developers authoring skills in application namespaces must '
         'inline magic numbers extracted from %SYS. This is fragile '
         'and undocumented.'],
    ]
)

add_para(
    'Once fixed, the MVP needs skills covering at minimum these domains '
    '(application-level, shipped with the health interop IPM module):'
)

add_table(
    ['Skill Domain', 'Needed By', 'What It Teaches the Agent'],
    [
        ['Productions', 'UC1, UC2',
         'Production anatomy, BusinessService/BusinessProcess/'
         'BusinessOperation patterns, lifecycle management, '
         'setting-level recommendations.'],
        ['DTL', 'UC1, UC3',
         'DTL syntax, foreach blocks for repeating segments, '
         'subtransform references, lookup table integration, '
         'virtual document path navigation.'],
        ['BPL', 'UC1',
         'BPL activities (assign, call, code, transform, rule), '
         'compensation handlers, async patterns.'],
        ['Routing Rules', 'UC1',
         'Constraint-based routing, transform chaining, '
         'dead-letter handling.'],
        ['HL7v2', 'UC1, UC2, UC3',
         'Message types, segment structure, ACK semantics, '
         'schema path navigation (flat vs nested groups).'],
        ['FHIR R4', 'UC1',
         'Resources, references, search parameters, R4 bundles.'],
        ['SDA', 'UC1',
         'SDA3 model as transformation hub. HL7-to-SDA-to-FHIR '
         'pipeline. IRIS never converts directly between wire '
         'formats; SDA is always the intermediary.'],
    ]
)


doc.add_heading('4.5 Catalog and Grounding', level=2)

add_para(
    'The %AI.RAG.KnowledgeBase class (Framework, exists) provides '
    'chunk, embed, store, and auto-expose-as-tool capabilities. The '
    'MVP requires these additional catalog behaviors.'
)

add_table(
    ['Requirement', 'Status', 'Driven By', 'What Is Needed', 'Why'],
    [
        ['IRIS class catalog',
         'Application-level',
         'UC1, UC2',
         'A searchable vector catalog of Ens.*, HS.*, and adapter '
         'classes. Each entry includes: class name, description, '
         'key settings, when-to-use guidance.',
         'When a user says "receive ADT via MLLP", the agent must '
         'find EnsLib.HL7.Service.TCPService. LLM training data '
         'does not know what is installed on this IRIS instance.'],
        ['Catalog auto-growth',
         'Framework (proposed)',
         'UC1, UC3',
         'After a successful build, new DTL/BPL/routing rule classes '
         'are automatically embedded and indexed.',
         'Without auto-growth, the catalog stales after the first '
         'build. The next user cannot discover what was just created.'],
        ['Staleness detection',
         'Framework (proposed)',
         'UC2',
         'Track when entries were last re-embedded. Flag entries '
         'whose source classes changed since last indexing.',
         'UC2 recommends based on catalog data. Stale entries cause '
         'incorrect recommendations.'],
        ['Class introspection tool',
         'Application-level',
         'UC1, UC2',
         'A tool that reads %Dictionary.CompiledClass to return '
         'properties, methods, parameters, and inheritance.',
         'Catalog search returns candidates. Introspection verifies '
         'the class exists and reveals its capabilities.'],
    ]
)


doc.add_heading('4.6 Streaming and Real-Time Feedback', level=2)

add_table(
    ['Requirement', 'Status', 'Driven By', 'What Is Needed', 'Why'],
    [
        ['Token streaming',
         'Framework (proposed)',
         'All UCs',
         'LLM output streams token-by-token via SSE (Server-Sent '
         'Events, a protocol where the server pushes text to the '
         'browser over a single long-lived HTTP connection).',
         'A 10-second wait then a wall of text feels broken. '
         'Streaming gives immediate feedback.'],
        ['Tool lifecycle events',
         'Framework (proposed)',
         'All UCs',
         'Each tool state transition (Planned, Running, Succeeded, '
         'Failed) emits a structured SSE event.',
         'UC1 has 5-10 tool calls over 90 seconds. Without lifecycle '
         'events the UI shows a loading spinner, not progress.'],
        ['Approval gate events',
         'Framework (proposed)',
         'UC1, UC3',
         'When a tool reaches AwaitingApproval, a structured event '
         'tells the UI to render an Approve/Reject prompt.',
         'The user must see what will happen and consent before any '
         'mutation executes.'],
        ['No silent execution',
         'Framework (proposed)',
         'All UCs',
         'Every server-side action (tool running, LLM thinking, '
         'retry) produces a visible event.',
         'Silent execution erodes trust. Users who cannot see what '
         'the agent is doing will not trust its results.'],
    ]
)


doc.add_heading('4.7 Security and Audit', level=2)

doc.add_heading('4.7.1 Authorization Layers', level=3)

add_para(
    'The %AI.Policy.Authorization class (Framework, exists) provides '
    'the extension point. The MVP requires four authorization layers:'
)

add_table(
    ['Layer', 'What It Enforces', 'Status'],
    [
        ['Agent',
         'Which users may converse with this agent.',
         'Framework (proposed)'],
        ['MCP Server',
         'Caller authentication and namespace validation at the '
         'server boundary.',
         'Framework (exists via %CSP.REST authentication)'],
        ['ToolSet',
         'Group-level access control (coarse permission to use a '
         'category of tools).',
         'Framework (proposed)'],
        ['Tool',
         'RequiresConfirmation gating, AllowedNamespaces enforcement, '
         'RBAC role check.',
         'Framework (proposed for confirmation/namespaces; exists for '
         'role-based via Policy)'],
    ]
)

doc.add_heading('4.7.2 Audit Record', level=3)

add_para(
    'Every tool invocation must produce an immutable audit record. '
    'The %AI.Policy.Audit class (Framework, exists) provides the '
    'abstract hook. Each record must capture:'
)

add_table(
    ['Field', 'Description'],
    [
        ['toolName', 'Name of the tool invoked'],
        ['inputArguments', 'Full input payload after schema validation'],
        ['outputResult', 'Full output, or structured error on failure'],
        ['executionDuration', 'Wall-clock time from Running to terminal state'],
        ['state', 'Terminal state: Succeeded, Failed, or Cancelled'],
        ['namespace', 'IRIS namespace in which the tool ran'],
        ['userIdentity', 'Authenticated user identity ($username)'],
        ['timestamp', 'UTC timestamp at terminal-state transition'],
    ]
)

doc.add_heading('4.7.3 Anti-Fabrication', level=3)

add_para(
    'The agent must not claim an operation succeeded unless it has '
    'the corresponding tool confirmation. "I have rebuilt the '
    'production" is only valid after the build tool returned Succeeded '
    'AND PostBuildValidation returned Succeeded. This applies to all '
    'three MVP use cases and is enforced by the agent system prompt '
    'and validated via the audit trail.'
)


doc.add_heading('4.8 Namespace Execution', level=2)

add_para(
    'Every tool runs in the namespace of the authenticated user session. '
    'The namespace is established at session start and validated before '
    'every tool invocation. An invocation against a namespace the user '
    'cannot access is rejected before the tool runs. Cross-namespace '
    'operations are never implicit.'
)


doc.add_heading('4.9 Error Contract', level=2)

add_para(
    'All tools should return errors in a standard structured shape so '
    'the orchestrator can make consistent decisions:'
)

add_table(
    ['Field', 'Type', 'Purpose'],
    [
        ['errorCode', 'string', 'Machine-readable code for retry/routing decisions'],
        ['message', 'string', 'Human-readable description for the UI'],
        ['retryable', 'boolean', 'Whether the orchestrator should retry'],
        ['severity', 'enum', 'info / warning / error / critical'],
        ['tool', 'string', 'Name of the tool that produced the error'],
        ['namespace', 'string', 'Namespace where the error occurred'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# 5. MVP EXPERIENCE REQUIREMENTS
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('5. MVP Experience Requirements', level=1)

doc.add_heading('5.1 Builder Experience (End-User UX)', level=2)

add_para(
    'Builders work inside the IRIS Management Portal. The framework '
    'must provide these configuration surfaces for the MVP.'
)

add_table(
    ['Capability', 'What the Builder Does', 'Driven By'],
    [
        ['Agent configuration',
         'Create or clone an agent. Edit system prompt. Set iteration '
         'limits and token budgets. Select MCP Servers and Skills. '
         'Choose LLM provider.',
         'All UCs'],
        ['MCP Server assembly',
         'Create MCP Servers by selecting from available ToolSets. '
         'Each MCP scopes a functional domain.',
         'All UCs'],
        ['ToolSet browsing',
         'Browse available ToolSets and their tools. See tool names, '
         'descriptions, and parameter signatures.',
         'All UCs'],
        ['Skill attachment',
         'View available Skills and attach them to Agents. See '
         'name, description, and estimated token usage.',
         'All UCs'],
        ['Provider configuration',
         'Configure LLM connections. Paste API keys (stored in IRIS '
         'Secure Wallet via %AI.Utils.WalletStore, never in '
         'plaintext). Click Test. See green/red status.',
         'All UCs'],
        ['Chat experience',
         'Open chatbot. See streaming responses with tool-call cards. '
         'Approve or reject mutations. Resume past conversations. '
         'See active namespace.',
         'All UCs'],
        ['Audit visibility',
         'View audit trail: tools called, arguments, results, '
         'duration, user, namespace.',
         'All UCs'],
    ]
)


doc.add_heading('5.2 Developer Experience (DX)', level=2)

add_table(
    ['Capability', 'What the Developer Does',
     'What the Framework Must Provide'],
    [
        ['Tool authoring',
         'Write %AI.Tool subclasses. Each public ClassMethod becomes '
         'a tool. The Description comment is the tool description the '
         'LLM sees. The FormalSpec defines parameters.',
         'Auto-introspection of ClassMethods into tool schemas. '
         'Compile-time validation.'],
        ['Skill authoring',
         'Write XData INSTRUCTIONS (markdown) and XData SUMMARY (YAML) '
         'in %AI.Agent.Skill subclasses. Set Parameter TOOLS for '
         'ToolSet access.',
         'Skill class must accept XData blocks without errors '
         '(requires %OnNew bug fix).'],
        ['Catalog seeding',
         'Walk %Dictionary for relevant classes. Curate descriptions. '
         'Generate embeddings via %AI.RAG.KnowledgeBase.',
         'Programmatic ingest API. Batch embedding. '
         'Idempotent re-indexing.'],
        ['Testing and dry-run',
         'Test individual tools by providing input and seeing output. '
         'No agent or LLM required.',
         'A tool execution path that bypasses the agent and calls '
         'the tool directly.'],
        ['IPM packaging',
         'All classes, skills, seed data ship as a single IPM module.',
         'Namespace-agnostic installation. CSP app creation at '
         'install time.'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# 6. MVP PERFORMANCE TARGETS
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('6. MVP Performance Targets', level=1)

add_table(
    ['Metric', 'Target', 'Driven By', 'Why'],
    [
        ['Time to first token', 'Under 3 seconds', 'All UCs',
         'Users lose confidence if chat feels unresponsive.'],
        ['Full production build (UC1)', 'Under 90 seconds', 'UC1',
         'Plan + approve + build + test must complete within '
         'LLM provider rate limits.'],
        ['Catalog search', 'Under 500ms per query', 'UC1, UC2',
         'Searches run in the plan phase; slow search delays '
         'everything.'],
        ['Tool execution', 'Under 5s typical, 30s max', 'All UCs',
         'Most tools complete in 1-5 seconds. Compilation may '
         'take longer.'],
        ['SSE stream gaps', 'No gaps > 2 seconds', 'All UCs',
         'Gaps make the UI feel frozen.'],
        ['Concurrent sessions', '10+ simultaneous users', 'All UCs',
         'Each session is independent. Process pool must handle load.'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# 7. MVP FRAMEWORK GAPS (BLOCKING)
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('7. Framework Gaps Blocking the MVP', level=1)

add_para(
    'These are the specific items that must be resolved before the '
    'MVP use cases can be delivered. Each gap was discovered during '
    'the reference implementation build and has a documented workaround '
    'that validates the proposed fix.'
)

add_table(
    ['Gap', 'Type', 'Blocks',
     'What the Reference Implementation Did', 'Proposed Framework Fix'],
    [
        ['Skill %OnNew marshaling',
         'Bug fix',
         'All UCs',
         'AgenticInterop.Skill.Base overrides %OnNew, calls '
         '%ToJSON().%ToJSON() to produce a JSON string before $ZF.',
         'Fix %AI.Agent.Skill.%OnNew to call .%ToJSON() on the '
         'DynamicObject before passing to $ZF, matching '
         '%AI.Provider behavior.'],
        ['%AI.INC not visible outside %SYS',
         'Bug fix',
         'All UCs',
         'Inlined macro values (IrisLLMLibrary=1042, '
         'LLMBUILDSKILLFROMJSON=66) extracted from %SYS.',
         'Make %AI.INC include file accessible from all namespaces.'],
        ['No RequiresConfirmation flag',
         'New feature',
         'UC1, UC3',
         'Custom %AI.Policy.Authorization subclass '
         '(ConfirmationGate) checks a class-level annotation.',
         'Add RequiresConfirmation as a first-class tool metadata '
         'property, with orchestrator support for pausing and '
         'emitting an approval event.'],
        ['No iteration limit / token budget',
         'New feature',
         'All UCs',
         'AgenticInterop.Agent.Monitor wraps agent.Run() with a '
         '60-second deadline and 50,000-token budget, killing the '
         'process if exceeded.',
         'Add MaxIterations and TokenBudget as %AI.Agent properties '
         'with graceful stop behavior.'],
        ['No tool lifecycle streaming',
         'New feature',
         'All UCs',
         'Application-level SSE event dispatch in the REST layer, '
         'manually emitting events at each tool state transition.',
         'The agent execution engine should emit structured lifecycle '
         'events (Planned, Running, Succeeded, Failed) that any '
         'UI can consume.'],
        ['No catalog auto-reindex',
         'New feature',
         'UC1, UC3',
         'Application-level post-build hook calls '
         'KnowledgeBase.AddDocument() for new artifacts.',
         'KnowledgeBase should support a watch/trigger mechanism '
         'that auto-indexes new classes matching a pattern.'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# APPENDIX A: FRAMEWORK FOUNDATION
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('Appendix A: %AI Framework Foundation', level=1)

add_para(
    'Complete catalog of %AI Framework classes shipped in IRIS for '
    'Health 2026.2. All classes listed here are Framework (exists). '
    'Refer to this appendix when evaluating whether a proposed '
    'requirement is already covered.'
)

add_table(
    ['Category', 'Class', 'Purpose'],
    [
        ['Core', '%AI.Agent', 'Execution engine; declarative via Parameter + XData'],
        ['Core', '%AI.Agent.Session', 'Persistent conversation state; Export/Import as JSON'],
        ['Core', '%AI.Provider', 'LLM provider abstraction (Anthropic, Bedrock, Azure, OpenAI, Gemini, NIM)'],
        ['Core', '%AI.System', 'AI-Core Rust bridge utilities'],
        ['LLM Types', '%AI.LLM.Context', 'Conversation context wrapper'],
        ['LLM Types', '%AI.LLM.Response', 'Response wrapper'],
        ['LLM Types', '%AI.LLM.ContentPart', 'Multi-modal content parts'],
        ['Tools', '%AI.Tool', 'Base tool class; auto-introspects public ClassMethods'],
        ['Tools', '%AI.ToolMgr', 'Unified registry (ObjectScript + Rust + MCP tools)'],
        ['Tools', '%AI.ToolSet', 'XML-driven tool grouping; compile-time composition'],
        ['Tools', '%AI.Tool.Resolver', 'URI-based resolver (rust: / iris: / mcp:stdio: / mcp:remote:)'],
        ['Tools', '%AI.Tool.Schema', 'Schema-building helpers for code generators'],
        ['Built-in', '%AI.Tools.FileSystem', 'File/directory operations'],
        ['Built-in', '%AI.Tools.SQL', 'SQL query execution'],
        ['Built-in', '%AI.Tools.ShellTools', 'Bundled ToolSet (filesystem + bash + SQL + web search)'],
        ['Sub-Agents', '%AI.Agent.SubAgent', 'Abstract base; tool that spawns a child agent'],
        ['Sub-Agents', '%AI.Agent.Skill', 'Declarative sub-agent via XData SUMMARY + INSTRUCTIONS'],
        ['MCP', '%AI.MCP.Service', 'REST + WebSocket MCP host for external clients'],
        ['Policies', '%AI.Policy.Authorization', 'Abstract auth gate; called before every tool invocation'],
        ['Policies', '%AI.Policy.Audit', 'Abstract audit logger; called after every invocation'],
        ['Policies', '%AI.Policy.Discovery', 'Abstract tool catalog filter; called each agent turn'],
        ['Policies', '%AI.Policy.ConsoleAuth', 'Built-in console confirmation'],
        ['Policies', '%AI.Policy.ConsoleAudit', 'Built-in console logger'],
        ['RAG', '%AI.RAG.Embedding', 'Abstract embedding provider'],
        ['RAG', '%AI.RAG.Embedding.FastEmbed', 'Local ONNX embedding (AllMiniLML6V2, 384d, no API key)'],
        ['RAG', '%AI.RAG.Embedding.OpenAI', 'OpenAI embedding via %AI.Provider'],
        ['RAG', '%AI.RAG.VectorStore.IRIS', 'SQL-backed vector table'],
        ['RAG', '%AI.RAG.KnowledgeBase', 'Chunk + embed + store + auto-expose search tool'],
        ['Settings', '%AI.Utils.SettingStore', 'Abstract setting resolver (@{prefix.key} syntax)'],
        ['Settings', '%AI.Utils.ConfigStore', 'ConfigStore-backed settings'],
        ['Settings', '%AI.Utils.WalletStore', 'IRIS Secure Wallet-backed secrets'],
        ['Shell', '%AI.Shell.Console', 'Abstract REPL for terminal agent interaction'],
        ['Shell', '%AI.Shell.ConsoleAgent', 'Pre-configured terminal agent'],
        ['Shell', '%AI.Shell.StreamRenderer', 'Stream rendering for terminal'],
        ['ToolSet Spec', 'Specification.ToolSet', 'Root XML element for ToolSet XData'],
        ['ToolSet Spec', 'Specification.Tool / Include / Exclude / Filter',
         'Tool declaration, inclusion, exclusion, conditional filtering'],
        ['ToolSet Spec', 'Specification.Query', 'Inline SQL exposed as a tool (compile-time validated)'],
        ['ToolSet Spec', 'Specification.MCP / .Stdio / .Remote', 'MCP server references'],
        ['ToolSet Spec', 'Specification.Policies / .PolicyElement', 'Policy attachment'],
        ['ToolSet Spec', 'Specification.Config / .EnvVar / .Requirement',
         'Static config, env vars, runtime dependencies'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# APPENDIX B: GENERAL REQUIREMENTS BEYOND MVP
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('Appendix B: General Requirements Beyond MVP', level=1)

add_para(
    'These requirements are important for the full product but are '
    'not blocking the three MVP use cases. They should be addressed '
    'after the MVP ships.'
)

add_table(
    ['Requirement', 'What Is Needed', 'Why'],
    [
        ['Context strategy',
         'Agent-level property controlling how prior turns are handled '
         'when approaching context limits. Options: full, summarized, '
         'hybrid. Invariants: tool results never dropped, user decisions '
         'kept verbatim, catalog grounding persists, recent turns in full.',
         'Long conversations exceed context limits. The MVP use cases '
         'typically complete in 5-15 turns, within most provider limits. '
         'Longer-running sessions need this.'],
        ['Skill loading strategies',
         'Static (always present), retrieval-based (loaded when tags '
         'match current turn), hybrid (summary always, full on demand).',
         'The MVP loads all skills statically. With 9+ skills, '
         'retrieval-based loading improves token efficiency.'],
        ['Skill priority and conflict resolution',
         'When multiple skills provide overlapping guidance, a priority '
         'value determines precedence.',
         'Not blocking MVP with 9 curated skills, but matters as '
         'third-party skills are added.'],
        ['Skill token budget',
         'Per-skill cap on context tokens consumed by INSTRUCTIONS.',
         'Prevents one verbose skill from crowding out others.'],
        ['Conversation search and history UI',
         'Queryable index of past conversations. Search by title, '
         'content, and date.',
         'MVP works with basic session list. History search improves '
         'the experience for heavy users.'],
        ['Conversation export',
         'Export conversations in structured format for compliance.',
         'Framework already supports Session.Export as JSON.'],
        ['Deterministic retries',
         'Retries follow declared policy (Idempotent flag + retry '
         'count). Agent does not invent retry behavior.',
         'Important for reliability but MVP can function without '
         'automatic retries.'],
        ['Streaming continuity',
         'SSE events resume on reconnection via correlation ID.',
         'Nice-to-have for MVP; essential for production deployments.'],
        ['Graceful degradation',
         'When catalog re-indexing or monitoring is unavailable, agent '
         'operates in reduced-capability mode.',
         'MVP assumes all subsystems are available.'],
        ['OutputSchema auto-generation',
         'Auto-generate JSON Schema for tool return values, matching '
         'how InputSchema is auto-generated from FormalSpec.',
         'Enables output validation. Not blocking MVP.'],
        ['Multiple agents per deployment',
         'Data model supports multi-agent from day one. Practical '
         'multi-agent orchestration (routing, handoff) is post-MVP.',
         'MVP ships a single agent. The data model is ready.'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# APPENDIX C: REFERENCE IMPLEMENTATION
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('Appendix C: Reference Implementation', level=1)

add_para(
    'The Health Interop Copilot (codebase: agentic_interop) is a '
    'working application built on the %AI Framework that validates '
    'every requirement in this document. Everything in this appendix '
    'is Application-level: examples of what the framework enables.'
)

doc.add_heading('C.1 Project Summary', level=2)

add_table(
    ['Metric', 'Value'],
    [
        ['ObjectScript classes', '58'],
        ['Tool classes (%AI.Tool)', '5'],
        ['Tools (public ClassMethods)', '42'],
        ['ToolSets (%AI.ToolSet)', '5'],
        ['MCP Servers', '5 (Production, Transform, Testing, Catalog, Monitoring)'],
        ['Skills (%AI.Agent.Skill)', '9'],
        ['IPM module version', '1.0.0'],
        ['Repository', 'github.com/dfrancoisc/agentic_interop'],
    ]
)

doc.add_heading('C.2 Tools by Domain', level=2)

add_table(
    ['Tool Class', 'Count', 'Example Tools'],
    [
        ['Tool.Production', '10',
         'ListProductions, GetProduction, CreateProduction, '
         'AddBusinessHost, UpdateBusinessHostSettings, '
         'StartProduction, StopProduction, PostBuildValidation'],
        ['Tool.Transform', '14',
         'CreateDTL, CompileDTL, DryRunDTL, CreateBPL, '
         'CreateRoutingRule, GetHL7SchemaMap, ListLookupTables, '
         'ListSDAFHIRDTLs, DescribeTransformationPipeline'],
        ['Tool.Testing', '6',
         'SendHL7, SendFHIR, ValidateHL7Structure, '
         'ValidateHL7Semantics, ValidateFHIRResource, CompareMessages'],
        ['Tool.Catalog', '7',
         'DescribeClass, GetUserNamespace, SearchApiIndex, '
         'LookupErrorCode, LookupGlossaryTerm, ExplainStatus'],
        ['Tool.Monitoring', '5',
         'QueryEventLog, TopErrors, QueryMessageStatus, '
         'MessageSummary, QueueStatus'],
    ]
)

doc.add_heading('C.3 Skills', level=2)

add_table(
    ['Skill', 'Domain', 'ToolSet Access'],
    [
        ['Productions', 'Production anatomy, lifecycle, best practices', 'ToolSet.Production'],
        ['DTL', 'DTL syntax, foreach, subtransforms, lookup tables', 'ToolSet.Transform'],
        ['BPL', 'BPL activities, compensation, async patterns', 'ToolSet.Transform'],
        ['RoutingRules', 'Constraint-based routing, dead-letter handling', 'ToolSet.Production'],
        ['HL7v2', 'Message types, segments, ACK semantics', 'ToolSet.Testing'],
        ['FHIRR4', 'Resources, references, search parameters', 'ToolSet.Testing'],
        ['SDA', 'SDA3 as transformation hub', 'ToolSet.Testing'],
        ['RestInProductions', 'REST services/operations in productions', 'ToolSet.Production'],
        ['ESBPattern', 'Enterprise Service Bus patterns', 'Production + Transform'],
    ]
)

doc.add_heading('C.4 Framework Gap Workarounds', level=2)

add_table(
    ['Workaround Class', 'Framework Gap', 'Section'],
    [
        ['AgenticInterop.Skill.Base',
         '%AI.Agent.Skill.%OnNew OREF-to-JSON marshaling',
         'Section 4.4, Section 7'],
        ['AgenticInterop.Agent.Monitor',
         'No iteration limit / token budget',
         'Section 4.1, Section 7'],
        ['AgenticInterop.Policy.ConfirmationGate',
         'No RequiresConfirmation flag on tools',
         'Section 4.2, Section 7'],
        ['AgenticInterop.Policy.ToolFilter',
         'No AllowedNamespaces / Category on tools',
         'Section 4.2, Section 7'],
        ['Inlined macro values (1042, 66)',
         '%AI.INC not visible outside %SYS',
         'Section 4.4, Section 7'],
    ]
)


# ═════════════════════════════════════════════════════════════════════
# APPENDIX D: GLOSSARY
# ═════════════════════════════════════════════════════════════════════

doc.add_heading('Appendix D: Glossary', level=1)

add_table(
    ['Term', 'Definition'],
    [
        ['Agent',
         'Orchestration runtime. Backed by %AI.Agent (Framework, exists). '
         'Plans and selects tools. Does not mutate IRIS directly.'],
        ['MCP Server',
         'Runtime boundary exposing ToolSets. Backed by %AI.MCP.Service '
         '(Framework, exists). Handles auth and namespace context.'],
        ['ToolSet',
         'Logical tool grouping. Backed by %AI.ToolSet (Framework, exists). '
         'Permission and discovery unit.'],
        ['Tool',
         'Atomic execution unit. Backed by %AI.Tool (Framework, exists). '
         'The only component that mutates IRIS state.'],
        ['Skill',
         'Domain knowledge sub-agent. Backed by %AI.Agent.Skill (Framework, '
         'exists). Declarative XData-driven.'],
        ['Catalog',
         'Vector + introspection store. Backed by %AI.RAG.KnowledgeBase '
         '(Framework, exists). Grounds agent references.'],
        ['Provider',
         'LLM backend. Backed by %AI.Provider (Framework, exists). Supports '
         'Anthropic, Bedrock, Azure OpenAI, OpenAI, Gemini, NIM.'],
        ['Session',
         'Persistent conversation state. Backed by %AI.Agent.Session '
         '(Framework, exists). Export/Import as JSON.'],
        ['Policy',
         'Lifecycle hook: Authorization (gates), Audit (logs), Discovery '
         '(filters). Framework (exists).'],
        ['SSE',
         'Server-Sent Events. Protocol where the server pushes text events '
         'to the browser over a single long-lived HTTP connection.'],
        ['DTL',
         'Data Transformation Language. IRIS mapping format for transforming '
         'messages between schemas.'],
        ['BPL',
         'Business Process Language. IRIS workflow language for '
         'orchestrating business logic.'],
        ['SDA',
         'Summary Document Architecture. IRIS clinical data model; '
         'intermediary hub between HL7, FHIR, and other formats.'],
        ['Builder',
         'End-user persona. Works in IRIS Management Portal. Configures '
         'agents without code.'],
        ['Developer (DX)',
         'Creator persona. Works in IDE. Authors tools, skills, catalogs '
         'as compiled classes.'],
        ['RequiresConfirmation',
         'Proposed tool flag that pauses execution for user approval. '
         'Framework (proposed).'],
        ['MutatingOperation',
         'Proposed tool flag indicating IRIS state changes. '
         'Framework (proposed).'],
        ['PostBuildValidation',
         'Validation pattern: agent confirms build succeeded before '
         'claiming success. Application-level.'],
        ['IPM',
         'InterSystems Package Manager. Namespace-aware deployment.'],
        ['FastEmbed',
         'Local ONNX embedding (AllMiniLML6V2, 384d). Bundled with IRIS. '
         'No API key. Framework (exists).'],
    ]
)


# ── footer ───────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Health Interoperability Agentic Framework - '
    'Product Requirements Specification - May 2026\n'
    'Review document: DF_Agentic_Interop_AI_Executive_Briefing.pptx'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── save ─────────────────────────────────────────────────────────────

out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, 'Health_Interop_Agentic_Framework_Dev_Spec.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
