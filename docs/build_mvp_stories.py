#!/usr/bin/env python3
"""Generate MVP_User_Stories.docx — product-level features for the
Health Interop Agentic Framework MVP.

Structure:
  Part 1 — Generic features that apply across all use cases
  Part 2 — Use-case-specific features (UC1, UC2, UC3)

Each feature is a high-level user story with acceptance criteria.
No implementation details, no class names, no framework internals.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)

HEADER_BG = RGBColor(0x1a, 0x1a, 0x2e)
HEADER_FG = RGBColor(0xff, 0xff, 0xff)


# ── helpers ─────────────────────────────────────────────────────────

def add_para(text, bold_prefix=None):
    p = doc.add_paragraph(style='Normal')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = HEADER_FG
        sh = c._element.get_or_add_tcPr()
        bg = sh.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto',
            qn('w:fill'): '1a1a2e',
        })
        sh.append(bg)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 1:
                sh = c._element.get_or_add_tcPr()
                bg = sh.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear', qn('w:color'): 'auto',
                    qn('w:fill'): 'f5f5fa',
                })
                sh.append(bg)
    doc.add_paragraph()
    return t


def feature_block(feature_id, user_story, acceptance):
    """Render a single feature as a compact block."""
    p = doc.add_paragraph()
    r = p.add_run(f'{feature_id}')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    p2 = doc.add_paragraph()
    r3 = p2.add_run(user_story)
    r3.italic = True
    r3.font.size = Pt(11)

    p3 = doc.add_paragraph()
    r4 = p3.add_run('Acceptance Criteria:')
    r4.bold = True
    r4.font.size = Pt(10)

    for ac in acceptance:
        add_bullet(ac)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Health Interoperability\nAgentic Framework')
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MVP Features and User Stories')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Product-level features organized by epic.\n'
    'Part 1 covers features that apply to every use case.\n'
    'Part 2 covers features specific to each MVP use case.\n\n'
    'May 2026'
)
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# HOW TO READ
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('How to Read This Document', level=1)

add_para(
    'This document defines the MVP at the product feature level. '
    'Each feature is written as a user story: '
    '"As a [Developer or Interface Engineer], I want to [action] so that '
    'I can [goal]." Features are grouped into epics.'
)

doc.add_heading('Personas', level=2)

add_table(
    ['Persona', 'Who They Are', 'Where They Work'],
    [
        ['Developer',
         'Writes ObjectScript or Python code. Builds tools, '
         'skills, and integrations. Ships IPM packages.',
         'IDE (VS Code, Studio), terminal, source control'],
        ['Int. Eng.',
         'Configures and operates the AI copilot. No code. '
         'Assembles agents from available components, tests '
         'connections, runs conversations.',
         'IRIS Management Portal (browser)'],
    ]
)

doc.add_heading('Document Structure', level=2)

add_para(
    'Part 1 (Generic Features) covers capabilities that every '
    'use case needs: connecting an LLM, configuring AI settings, '
    'the developer authoring experience, the chatbot conversation '
    'experience, and trust and audit. These are the foundation.'
)

add_para(
    'Part 2 (Use Case Features) covers capabilities specific to '
    'each of the three MVP use cases. These build on top of '
    'Part 1.'
)


# ═══════════════════════════════════════════════════════════════════
# MVP USE CASES
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('MVP Use Cases', level=1)

add_para(
    'The MVP is defined by three use cases. Every feature in this '
    'document exists to enable one or more of these use cases.'
)

add_table(
    ['ID', 'Name', 'What the User Says'],
    [
        ['UC1',
         'Guided Production Build',
         '"I need to receive ADT admission messages from our HIS '
         'and send observation reports to the downstream LIS."'],
        ['UC2',
         'Production Review',
         '"Review the LAB.Production and tell me what it does '
         'and how to improve it."'],
        ['UC3',
         'Complex HL7 Transformations',
         '"Transform ADT^A01 into ORU^R01 with cross-segment '
         'field mappings."'],
    ]
)

add_para(
    'UC1 is the most demanding use case. It requires the agent to '
    'have a multi-turn conversation, search a catalog, present a '
    'plan, build multiple artifacts, test them, and validate. '
    'If UC1 works end to end, UC2 and UC3 are subsets of the same '
    'capabilities.'
)


doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#  PART 1 — GENERIC FEATURES
# ═══════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PART 1')
r.font.size = Pt(14)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Generic Features')
r.font.size = Pt(20)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Capabilities that apply across all use cases.\n'
    'These are the foundation every AI copilot needs.'
)
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════
# EPIC A — LLM PROVIDER CONNECTION
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Epic A: LLM Provider Connection', level=1)

add_para(
    'Before the copilot can do anything, it needs a working '
    'connection to a large language model. Both Interface Engineers (via the '
    'Management Portal) and Developers (via code) need to be able '
    'to configure, test, and monitor this connection.'
)

feature_block(
    'A.1  Configure an LLM provider via the Management Portal',
    'As an Interface Engineer, I want to configure an LLM provider connection '
    'by selecting a provider type, entering the required credentials, '
    'and saving the configuration so that the AI copilot has a '
    'working LLM to use.',
    [
        'The configuration form adapts to the selected provider type '
        '(e.g., AWS Bedrock needs a region and bearer token; a direct '
        'API provider needs an API key and model name).',
        'Credentials are stored securely in the platform credential '
        'store, never in plaintext in a database table or config file.',
        'The form validates required fields before saving.',
        'Multiple providers can be configured, but each agent uses '
        'one at a time.',
    ]
)

feature_block(
    'A.2  Configure an LLM provider via code',
    'As a Developer, I want to configure an LLM provider '
    'programmatically so that I can automate setup, include provider '
    'configuration in deployment scripts, and test against different '
    'models without using the UI.',
    [
        'A programmatic API allows creating, reading, updating, and '
        'deleting provider configurations.',
        'Credentials written via the API go to the same secure '
        'credential store as the UI path.',
        'Provider configuration can be exported and imported as part '
        'of a deployment package (minus secrets).',
    ]
)

feature_block(
    'A.3  Test the LLM connection and see a status indicator',
    'As an Interface Engineer, I want to click a Test button after configuring '
    'a provider and immediately see a green or red status indicator '
    'so that I know the connection works before anyone tries to '
    'use the chatbot.',
    [
        'Clicking Test sends a minimal request to the LLM and shows '
        'the result within a few seconds.',
        'Success shows: green status, model name, and response latency.',
        'Failure shows: red status and the verbatim error message '
        'from the provider (not a generic error).',
        'The status indicator is visible everywhere the provider is '
        'referenced: the provider list, the agent configuration, '
        'and the chatbot header.',
        'If the provider is red, the chatbot input is disabled with '
        'a clear explanation.',
    ]
)

feature_block(
    'A.4  Secure credential management',
    'As an Interface Engineer, I want my API keys and tokens to be stored '
    'securely and never visible in plaintext so that credentials '
    'cannot be leaked through the UI, logs, database queries, or '
    'API responses.',
    [
        'Credentials are stored in the platform credential store.',
        'The UI shows masked values (dots or stars) for saved '
        'credentials.',
        'API responses never include credential values.',
        'Logs and audit records never contain credential values.',
        'A credential can be rotated (overwritten) without deleting '
        'and recreating the provider.',
    ]
)


# ═══════════════════════════════════════════════════════════════════
# EPIC B — BUILDER CONFIGURATION EXPERIENCE
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Epic B: AI Configuration for Interface Engineers', level=1)

add_para(
    'Interface Engineers work in the IRIS Management Portal. They need to '
    'assemble and configure AI agents without writing code: choosing '
    'which tools and skills an agent has, setting safety limits, and '
    'understanding what each component does.'
)

feature_block(
    'B.1  Create and configure an agent via UI',
    'As an Interface Engineer, I want to create and configure an AI agent '
    'through a form in the Management Portal so that I can define '
    'its behavior, set safety limits, and choose its LLM provider '
    'without writing code.',
    [
        'The form allows setting: agent name, system prompt (free '
        'text), LLM provider (selected from configured providers), '
        'maximum number of LLM turns per request, and total token '
        'budget per request.',
        'The system prompt can be edited in a text area large enough '
        'to see the full content.',
        'Changes take effect on the next new conversation (active '
        'conversations are not disrupted).',
        'The agent list shows all configured agents with their '
        'provider status indicator.',
    ]
)

feature_block(
    'B.2  Assemble tools and skills for an agent',
    'As an Interface Engineer, I want to select which tool groups and domain '
    'skills are available to my agent so that I can control what '
    'the agent can do and what it knows about.',
    [
        'The agent configuration shows a list of available tool '
        'groups (organized by functional area) and domain skills.',
        'The Interface Engineer adds or removes tool groups and skills using '
        'a multi-select interface.',
        'Each tool group shows its name, a plain-language '
        'description, and how many individual tools it contains.',
        'Each skill shows its name and a summary of the domain '
        'knowledge it provides.',
        'The agent only has access to tools and skills explicitly '
        'assigned to it.',
    ]
)

feature_block(
    'B.3  Browse and understand available tools',
    'As an Interface Engineer, I want to browse all available tools, read their '
    'descriptions, and see their parameters so that I understand '
    'what each tool does before including it in an agent.',
    [
        'A tool browser lists all tools, grouped by functional area.',
        'Each tool shows: name, plain-language description, input '
        'parameters with types and descriptions, and whether it '
        'changes state (read-only vs. mutating).',
        'Tools are read-only in the Interface Engineer UI. Developers create '
        'tools in code.',
        'The tool browser supports search and filtering.',
    ]
)

feature_block(
    'B.4  Set approval policies for state-changing actions',
    'As an Interface Engineer, I want to configure whether the agent asks for '
    'my approval before making changes so that I can control the '
    'level of autonomy the agent has.',
    [
        'The agent configuration offers approval modes: '
        'always ask (every state-changing tool pauses for approval), '
        'never ask (for dev/test environments), or '
        'per-tool (each tool has its own setting).',
        'The default for new agents is always ask.',
        'Changing the approval mode takes effect on the next '
        'conversation.',
    ]
)


# ═══════════════════════════════════════════════════════════════════
# EPIC C — DEVELOPER EXPERIENCE
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Epic C: AI Development for Developers', level=1)

add_para(
    'Developers create the tools and skills that agents use. They '
    'work in an IDE, write code, test independently of the LLM, '
    'and package everything for distribution.'
)

feature_block(
    'C.1  Create tools by writing code',
    'As a Developer, I want to create a new tool by writing a '
    'method in my IDE so that the framework automatically discovers '
    'it and makes it available to agents without any manual '
    'registration step.',
    [
        'A tool is a public method on a tool class. The method '
        'description becomes the tool description the LLM sees.',
        'The method parameters are auto-converted to a schema the '
        'LLM understands.',
        'No XML editing, no manual registration, no restart. '
        'Compile the class and the tool is available.',
        'The tool description is a contract with the LLM: imperative '
        'verb, scope, side effects, expected inputs. The developer '
        'writes it as carefully as an API summary.',
    ]
)

feature_block(
    'C.2  Annotate tool behavior',
    'As a Developer, I want to mark a tool with metadata (requires '
    'user confirmation, mutating vs. read-only, timeout, category) '
    'so that the framework, UI, and audit system can handle it '
    'appropriately without per-tool special-casing.',
    [
        'A tool can be marked as requiring user confirmation before '
        'execution.',
        'A tool can be marked as mutating (changes state) or '
        'read-only.',
        'A tool can specify a timeout so a stuck tool does not '
        'block the agent.',
        'A tool can specify a category for grouping in the UI and '
        'audit reports.',
        'These annotations are declared in the tool code, not in '
        'a separate configuration file.',
    ]
)

feature_block(
    'C.3  Create domain skills by writing code',
    'As a Developer, I want to create a domain skill by writing '
    'instructions in markdown and metadata in a structured block '
    'so that the agent gains domain-specific knowledge without '
    'changes to the agent class.',
    [
        'A skill is a class containing a markdown instruction block '
        'and a metadata block (summary, description, domains).',
        'The instructions are injected into the agent context when '
        'the skill is attached.',
        'A skill can declare which tool groups it works with, so '
        'the agent knows which tools to use for that domain.',
        'Skills compile and instantiate without errors in any '
        'application namespace.',
    ]
)

feature_block(
    'C.4  Test tools without an LLM',
    'As a Developer, I want to invoke a tool directly with input '
    'arguments and see the raw output so that I can test and debug '
    'during development without needing an LLM connection or '
    'running the full agent.',
    [
        'A dry-run mechanism accepts a tool name and input, '
        'executes the tool, and returns the raw result.',
        'Dry-run bypasses the agent loop and the LLM entirely.',
        'The admin UI provides a dry-run panel for each tool.',
        'The developer can also dry-run from the terminal or IDE.',
    ]
)

feature_block(
    'C.5  Package and distribute as a module',
    'As a Developer, I want to package all my tools, skills, '
    'catalog seeds, and UI assets as a single installable module '
    'so that another developer can install the entire copilot '
    'with one command.',
    [
        'The module installs into any compatible IRIS namespace.',
        'Installation compiles classes, creates web applications, '
        'deploys UI assets, and seeds reference data.',
        'Uninstalling cleanly removes all artifacts.',
        'No machine-specific paths or hardcoded namespace names '
        'in the module definition.',
    ]
)


# ═══════════════════════════════════════════════════════════════════
# EPIC D — CHATBOT EXPERIENCE
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Epic D: Chatbot Experience', level=1)

add_para(
    'The chatbot is where Interface Engineers interact with the agent. It must '
    'feel responsive, show what the agent is doing at every moment, '
    'give the user control over state-changing actions, and maintain '
    'conversation context across sessions.'
)

feature_block(
    'D.1  Streaming responses',
    'As an Interface Engineer, I want to see the agent\'s response appear '
    'word by word as it is generated so that the chat feels '
    'responsive and I can start reading before the full answer '
    'is ready.',
    [
        'The response streams to the browser in real time using '
        'Server-Sent Events (SSE).',
        'The first visible word appears within a few seconds of '
        'sending a message.',
        'There is no visible freeze or gap during generation under '
        'normal conditions.',
        'If the connection drops, the UI shows a clear reconnection '
        'status, not a blank screen.',
    ]
)

feature_block(
    'D.2  Visible tool activity',
    'As an Interface Engineer, I want to see a card in the chat every time the '
    'agent uses a tool so that I can follow each step of a '
    'multi-step operation and understand what the agent is doing.',
    [
        'Each tool invocation renders as an inline card showing: '
        'tool name, a plain-language description of the action, '
        'and live status (running, succeeded, failed).',
        'The card updates in real time as the tool progresses.',
        'The card is collapsible: clicking it shows the full input '
        'and output.',
        'Multiple tool cards appear in sequence for multi-step '
        'operations, creating a visible trace of the agent\'s work.',
    ]
)

feature_block(
    'D.3  Approval gate for state-changing actions',
    'As an Interface Engineer, I want the agent to pause and show me what it '
    'is about to do whenever it wants to make a change, and let me '
    'Approve or Reject before anything happens.',
    [
        'When the agent reaches a state-changing tool, the chat '
        'shows an approval card with: what will change, the key '
        'arguments, and Approve / Reject buttons.',
        'Nothing changes in IRIS until the Interface Engineer clicks Approve.',
        'Clicking Reject cancels the action and the agent is '
        'informed so it can adjust its plan.',
        'The approval card clearly distinguishes read-only actions '
        '(no approval needed) from state-changing actions.',
    ]
)

feature_block(
    'D.4  Conversation history and context',
    'As an Interface Engineer, I want to see my past conversations, reopen any '
    'of them, and continue where I left off so that I do not lose '
    'work between sessions.',
    [
        'A side panel lists past conversations with a title or '
        'summary for each.',
        'Clicking a past conversation reopens it with the full '
        'message history.',
        'The agent has context of the prior messages when the '
        'conversation is resumed.',
        'Starting a new conversation clears the context.',
        'Conversations are scoped to the authenticated user.',
    ]
)

feature_block(
    'D.5  Namespace awareness',
    'As an Interface Engineer, I want to see which IRIS namespace I am working '
    'in so that I always know where the agent will make changes.',
    [
        'The chatbot header shows the active namespace.',
        'All tools execute in the namespace of the authenticated '
        'session.',
        'If the user does not have access to a namespace, the '
        'agent cannot operate in it.',
        'The namespace cannot be changed mid-conversation; '
        'starting a new conversation picks up the current session '
        'namespace.',
    ]
)

feature_block(
    'D.6  Guided onboarding',
    'As an Interface Engineer, I want to see example prompts when I start a '
    'new conversation so that I can discover what the agent can '
    'do without guessing.',
    [
        'A new conversation shows 3 to 5 clickable starter prompts '
        'aligned to the MVP use cases.',
        'Clicking a prompt inserts it as the first message.',
        'Examples cover all three use cases: building a production, '
        'reviewing a production, and creating a transformation.',
    ]
)

feature_block(
    'D.7  Continuous visible feedback',
    'As an Interface Engineer, I want the chat to never appear frozen or stuck '
    'so that I always know the system is working.',
    [
        'When the LLM is thinking, a visible indicator is shown.',
        'When a tool is running, its card shows a running status.',
        'When a retry occurs, it is visible.',
        'No execution path produces zero visible output for more '
        'than a few seconds.',
        'Long operations break into multiple visible steps rather '
        'than one long silent wait.',
    ]
)


# ═══════════════════════════════════════════════════════════════════
# EPIC E — TRUST, SAFETY, AND AUDIT
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Epic E: Trust, Safety, and Audit', level=1)

add_para(
    'The agent operates on production data. Interface Engineers need to trust '
    'that it does what it says, that every action is recorded, and '
    'that safety limits prevent runaway behavior.'
)

feature_block(
    'E.1  Complete audit trail',
    'As an Interface Engineer, I want every tool invocation to produce an '
    'immutable audit record so that I have a complete history of '
    'what the agent did, when, who requested it, and in which '
    'namespace.',
    [
        'Each record captures: tool name, input arguments, output '
        'result, execution duration, final status (succeeded, '
        'failed, cancelled), namespace, user identity, and '
        'timestamp.',
        'Records are append-only. No modification after creation.',
        'The admin UI shows audit records with filtering by time, '
        'user, tool, status, and namespace.',
    ]
)

feature_block(
    'E.2  Agent honesty — no unverified claims',
    'As an Interface Engineer, I want the agent to never claim an operation '
    'succeeded unless it has tool confirmation so that I can trust '
    'every statement the agent makes.',
    [
        'The agent does not say "I have created the production" '
        'unless the build tool returned success AND a validation '
        'tool confirmed the result.',
        'The agent cites tool results when making factual claims.',
        'The audit trail allows independent verification of every '
        'success claim.',
    ]
)

feature_block(
    'E.3  Safety limits',
    'As an Interface Engineer, I want the agent to stop gracefully when it '
    'reaches its configured limits (turns or tokens) so that a '
    'confused model cannot loop indefinitely or consume unlimited '
    'resources.',
    [
        'The agent respects the maximum turn limit set in its '
        'configuration.',
        'The agent respects the token budget set in its '
        'configuration.',
        'When a limit is reached, the agent stops and returns what '
        'it has produced so far with a clear explanation.',
        'The response indicates the limit was reached, not an error.',
    ]
)

feature_block(
    'E.4  Post-build validation',
    'As an Interface Engineer, I want the agent to validate the results of '
    'every build operation before claiming success so that I know '
    'the production or transformation actually works.',
    [
        'After building, the agent runs validation: artifact exists, '
        'compiles without errors, dependencies are resolved, and a '
        'test message flows through.',
        'The agent reports validation results as a checklist.',
        'Failed checks include specific errors and suggested '
        'corrective actions.',
    ]
)


# ═══════════════════════════════════════════════════════════════════
# EPIC F — CATALOG AND GROUNDING
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Epic F: Catalog and Grounding', level=1)

add_para(
    'The agent must recommend real IRIS classes and components, not '
    'hallucinated ones. A searchable catalog of available classes '
    'grounds the agent in what actually exists in the target namespace.'
)

feature_block(
    'F.1  Catalog search before building',
    'As an Interface Engineer, I want the agent to search a catalog of '
    'available IRIS classes before building anything so that it '
    'uses real class names from my instance, not guessed ones.',
    [
        'A searchable catalog of IRIS classes (business services, '
        'processes, operations, adapters, transformations) is '
        'available as a tool the agent can call.',
        'Search results include: class name, description, key '
        'settings, and when-to-use guidance.',
        'The agent cites catalog results in its plan so the '
        'Interface Engineer can verify the choices.',
    ]
)

feature_block(
    'F.2  Catalog seeding and refresh',
    'As a Developer, I want to seed the catalog from the classes '
    'available in the target namespace and refresh it when classes '
    'change so that the search index stays current.',
    [
        'The catalog can be built from the classes available in '
        'the target namespace.',
        'The catalog can also be seeded from a curated source '
        '(spreadsheet or export).',
        'A rebuild can be triggered from the admin UI or '
        'programmatically.',
        'Rebuild is idempotent: re-running does not create '
        'duplicate entries.',
    ]
)

feature_block(
    'F.3  Auto-index new artifacts',
    'As an Interface Engineer, I want new artifacts created during a build '
    'session (transformations, routing rules, productions) to '
    'automatically appear in the catalog so that the next user '
    'can discover and reuse them.',
    [
        'After a successful build, newly created classes are '
        'indexed into the catalog without manual action.',
        'The next search in a new conversation returns the newly '
        'created artifact.',
    ]
)


doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#  PART 2 — USE CASE FEATURES
# ═══════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PART 2')
r.font.size = Pt(14)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Use Case Features')
r.font.size = Pt(20)
r.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Capabilities specific to each MVP use case.\n'
    'These build on the generic features from Part 1.'
)
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════
# UC1 — GUIDED PRODUCTION BUILD
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('UC1: Guided Production Build', level=1)

add_para(
    'The Interface Engineer says: "I need to receive ADT admission messages '
    'from our HIS and send observation reports to the downstream '
    'LIS." The agent guides them through the entire build process.'
)

add_para(
    'This is the most demanding use case. It exercises every '
    'generic feature: streaming, tool cards, approval gates, '
    'catalog search, audit, and validation. If UC1 works end to '
    'end, the foundation is solid.'
)

feature_block(
    'UC1.1  Guided requirements gathering',
    'As an Interface Engineer, I want the agent to ask me clarifying questions '
    'about my integration needs before building anything so that '
    'the resulting production matches my actual requirements.',
    [
        'The agent asks about: message types and versions, '
        'transport protocols, source and destination systems, '
        'field mappings or transformation needs, error handling '
        'preferences, and throughput expectations.',
        'Each question explains why it matters in plain language.',
        'The agent does not proceed to planning until the '
        'requirements are clear.',
        'The agent summarizes the requirements back to the Interface Engineer '
        'for confirmation before moving on.',
    ]
)

feature_block(
    'UC1.2  Structured plan with catalog grounding',
    'As an Interface Engineer, I want the agent to present a structured plan '
    'listing every component it will create, which IRIS classes it '
    'will use, and the data flow before building anything so that '
    'I can review and approve the design.',
    [
        'The plan lists each component: name, what it does, the '
        'IRIS class it will use (from the catalog), and key '
        'settings.',
        'The plan shows the data flow from source to destination.',
        'Every IRIS class in the plan comes from a catalog search '
        'result, not from the LLM guessing.',
        'The plan ends with a clear "Shall I proceed?" question.',
        'No mutating tool runs until the Interface Engineer approves the plan.',
    ]
)

feature_block(
    'UC1.3  End-to-end build',
    'As an Interface Engineer, I want the agent to build the complete '
    'production in a single session so that I do not have to '
    'manually wire up each component.',
    [
        'The agent creates all required artifacts: production '
        'definition, business hosts, transformations, routing '
        'rules, and lookup tables.',
        'Artifacts are created in the correct dependency order.',
        'Cross-artifact references are correct: routing rules '
        'point to real transformations, transformations reference '
        'real lookup tables.',
        'All artifacts compile without errors.',
        'Each creation step shows a tool card so the Interface Engineer can '
        'follow progress.',
    ]
)

feature_block(
    'UC1.4  Automated testing and validation',
    'As an Interface Engineer, I want the agent to send a test message through '
    'the production after building it and show me the results so '
    'that I know it actually works.',
    [
        'The agent generates a sample message matching the '
        'production\'s expected input.',
        'The test message is sent through the production.',
        'The report shows: message sent, transformation output, '
        'routing result, processing time, and any errors.',
        'If errors occur, the agent diagnoses and offers to fix '
        'them.',
    ]
)


# ═══════════════════════════════════════════════════════════════════
# UC2 — PRODUCTION REVIEW
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('UC2: Production Review and Optimization', level=1)

add_para(
    'The Interface Engineer says: "Review the LAB.Production and tell me what '
    'it does and how to improve it." The agent inspects the '
    'production, explains it in plain language, and gives '
    'actionable recommendations.'
)

feature_block(
    'UC2.1  Plain-language production explanation',
    'As an Interface Engineer, I want to ask the agent to review an existing '
    'production and explain what it does so that I understand the '
    'data flow without reading code.',
    [
        'The agent reads the production definition and describes '
        'each component: what it does, what it connects to, and '
        'its role in the data flow.',
        'The explanation uses plain language, not just class names.',
        'The agent distinguishes between different types of '
        'settings (host-level vs. adapter-level vs. connection '
        'settings).',
        'The explanation follows the data flow from source to '
        'destination.',
    ]
)

feature_block(
    'UC2.2  Actionable recommendations',
    'As an Interface Engineer, I want the agent to give me specific, '
    'setting-level optimization recommendations so that each '
    'suggestion tells me exactly what to change and why.',
    [
        'Each recommendation cites a specific setting by name.',
        'Each recommendation includes: current value, recommended '
        'value, and the expected impact of the change.',
        'Recommendations are grounded in the catalog and the '
        'agent\'s domain knowledge, not generic advice.',
        'Recommendations are prioritized by impact.',
    ]
)

feature_block(
    'UC2.3  Error pattern analysis',
    'As an Interface Engineer, I want the agent to check the Event Log for '
    'recent errors and include error patterns in its review so '
    'that I see both configuration issues and runtime problems.',
    [
        'The agent queries the Event Log for errors in a '
        'configurable time window.',
        'Error patterns are grouped by frequency and connected to '
        'specific components.',
        'The review connects runtime errors to configuration '
        'recommendations where applicable.',
    ]
)


# ═══════════════════════════════════════════════════════════════════
# UC3 — COMPLEX HL7 TRANSFORMATIONS
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('UC3: Complex HL7 Transformations', level=1)

add_para(
    'The Interface Engineer says: "Transform ADT^A01 into ORU^R01 with '
    'cross-segment field mappings." The agent builds the complete '
    'transformation, handling the schema differences between the '
    'two message types.'
)

feature_block(
    'UC3.1  Schema-aware mapping',
    'As an Interface Engineer, I want the agent to look up the exact message '
    'structure for both my source and target message types before '
    'writing any transformation so that the mapping uses verified '
    'paths, not guessed ones.',
    [
        'The agent introspects the schema for the source message '
        'type (e.g., ADT_A01) and the target message type '
        '(e.g., ORU_R01).',
        'The introspection returns exact segment and field paths, '
        'including nested groups.',
        'The agent shows the user which fields will be mapped and '
        'how before creating the transformation.',
    ]
)

feature_block(
    'UC3.2  Complete transformation in one pass',
    'As an Interface Engineer, I want the agent to build the complete '
    'transformation with all field mappings efficiently so that '
    'the transformation is created in a single step, not through '
    'many incremental updates.',
    [
        'The transformation is built in one creation step, not '
        'through iterative additions.',
        'The transformation includes: field mappings, logic for '
        'repeating segments, conditional mappings, references to '
        'sub-transformations, and lookup table entries.',
        'The transformation compiles without errors.',
    ]
)

feature_block(
    'UC3.3  Test and verify field-level output',
    'As an Interface Engineer, I want the agent to test the transformation with '
    'a sample message and verify that every mapped field actually '
    'produces data so that I know the mapping is complete.',
    [
        'The agent runs the transformation with a sample message.',
        'The test output shows every mapped field and its resulting '
        'value.',
        'If any field is empty or incorrect, the agent checks the '
        'paths and offers to correct them.',
        'The Interface Engineer can provide their own sample message or use '
        'an agent-generated one.',
    ]
)


doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# SUMMARY
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('Summary', level=1)

add_table(
    ['Epic', 'Features', 'Scope'],
    [
        ['A. LLM Provider Connection', '4', 'Generic'],
        ['B. AI Configuration for Interface Engineers', '4', 'Generic'],
        ['C. AI Development for Developers', '5', 'Generic'],
        ['D. Chatbot Experience', '7', 'Generic'],
        ['E. Trust, Safety, and Audit', '4', 'Generic'],
        ['F. Catalog and Grounding', '3', 'Generic'],
        ['UC1. Guided Production Build', '4', 'Use Case'],
        ['UC2. Production Review', '3', 'Use Case'],
        ['UC3. HL7 Transformations', '3', 'Use Case'],
        ['TOTAL', '37', ''],
    ]
)

add_para(
    '27 generic features provide the foundation that any AI copilot '
    'built on this framework needs. 10 use-case features define the '
    'specific capabilities for the three MVP scenarios. The generic '
    'features should be built first, as every use case depends on them.'
)

doc.add_paragraph()

add_para(
    'Build sequence recommendation: ', bold_prefix=''
)
add_bullet(
    'Epic A (LLM connection) and Epic C (developer tooling) first '
    '— nothing works without an LLM and tools.')
add_bullet(
    'Epic D (chatbot experience) and Epic E (trust and audit) '
    'next — the interaction layer.')
add_bullet(
    'Epic B (builder configuration) and Epic F (catalog) in '
    'parallel — the configuration and grounding layer.')
add_bullet(
    'UC1 features, then UC2, then UC3 — each builds on the '
    'foundation and proves it works.')


# ── footer ──────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Health Interoperability Agentic Framework - '
    'MVP Features and User Stories - May 2026'
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── save ────────────────────────────────────────────────────────────

out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, 'MVP_User_Stories.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
