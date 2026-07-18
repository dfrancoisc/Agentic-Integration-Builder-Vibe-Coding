#!/usr/bin/env python3
"""
Build the Product Requirements Document (Word + Markdown) for the
Integration Agentic Builder.

Structure (MVP, in build-priority order):
  Overview
  Epic 1. The Agent        (most important: tools, skills, catalogs, policies, connection)
  Epic 2. Chat Experience
  Epic 3. AI Setting Experience
  4. Walkthrough  5. Summaries  6. Definition of Done  7. Non functional requirements

Style: plain English, short sentences, no em-dashes. Implementation agnostic,
except two facts the business owner asked to keep: the product is built on the
AI Hub Framework, and artifacts must be ObjectScript-based so they open in the
Health Connect Cloud graphical tools.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX_OUT = os.path.join(HERE, "Product_Requirements_Integration_Agentic_Builder.docx")
MD_OUT = os.path.join(HERE, "Product_Requirements_Integration_Agentic_Builder.md")

NAVY = RGBColor(0x15, 0x3E, 0x6B)
SLATE = RGBColor(0x33, 0x3A, 0x44)
INK = RGBColor(0x1A, 0x1D, 0x22)
HEADER_FILL = "153E6B"
ZEBRA_FILL = "EEF2F7"

IE = "an Integration Engineer"
AD = "an AI Admin"

BLOCKS = []
def B(kind, payload=None): BLOCKS.append((kind, payload))


B("title", "Integration Agentic Builder")
B("subtitle", "Product Requirements for the MVP. Health Connect Cloud.")
B("meta", [
    ("Product", "Integration Agentic Builder. An AI agent that builds healthcare interfaces by chat."),
    ("Built on", "InterSystems AI Hub Framework."),
    ("Platform (MVP)", "Health Connect Cloud."),
    ("Status", "Draft for review."),
    ("Date", "June 2026."),
])

# ---------------------------------------------------------------- OVERVIEW
B("h1", "Overview")
B("p", "The Integration Agentic Builder is an AI agent. It is built on the InterSystems AI Hub Framework. It helps an integration engineer build healthcare interfaces in Health Connect Cloud by talking to it in plain language. In this document, the agent means the Integration Agentic Builder.")
B("p", "The MVP has three parts. Build them in this order.")
B("numbers", [
    "The Agent. Its tools, skills, catalogs, policies, and LLM connection. This is the most important part. Build it first.",
    "The Chat Experience. Where the integration engineer works with the agent. Build it second.",
    "The AI Setting Experience. Where an AI Admin assembles the agent. It is internal to InterSystems. Customers do not see it in the MVP. Build it third.",
])
B("p", "More agents come later. They are not in the MVP.")
B("p", "Two people use the product. The Integration Engineer is a non developer. They build interfaces by chatting with the agent and by using the Health Connect Cloud graphical tools. The AI Admin assembles the agent. For this product the AI Admin is an InterSystems person.")
B("callout", "Driving use case (Epic to Quest). Epic sends HL7 v2 ADT over a TCP port, and FHIR over a web endpoint. Quest must receive HL7 over a TCP port, and FHIR at a secured FHIR server URL. A business process transforms Epic's data into the format Quest requires. This one use case drives the tools and skills we must build.")
B("lead", "Two hard requirements apply everywhere")
B("bullets", [
    "Inherit the security context. The agent runs inside the calling namespace and the signed in user. It respects both before its own capabilities. A namespace that does not allow FHIR blocks FHIR, even if the agent can do FHIR. A user who cannot set up OAuth blocks OAuth, even if the tool exists. Capability never overrides permission.",
    "Integrate with the pipeline. Every artifact the agent builds enters the CI/CD and change control pipeline that Health Connect Cloud already uses. The agent never writes around it.",
])

B("pagebreak")

# ---------------------------------------------------------------- EPIC 1
B("h1", "Epic 1. The Agent")
B("p", "Scope: MVP. Highest priority. The Agent is the core of the product. Build it first.")

B("h2", "1.1 How the agent works")
B("p", "The agent always follows the same safe loop. It researches, proposes, waits for approval, builds, tests, and reports.")
B("table", {
    "headers": ["Step", "What the agent does"],
    "rows": [
        ["1. Research", "Understand the request. Read the uploaded spec. Search the catalogs. Pick the right building blocks and patterns."],
        ["2. Propose", "Show a plan. List what it will create and change."],
        ["3. Approval gate", "Stop and wait. The engineer must approve. The agent builds nothing without approval."],
        ["4. Build", "Create the interface and its parts. Validate as it goes."],
        ["5. Test", "Send test messages. Validate them. Compare input to output."],
        ["6. Report", "Say what it built. Show the results. Link to the graphical tools and the visual trace."],
    ],
})
B("story", {"id": "US-1.01", "title": "The agent works in a safe loop",
    "persona": IE, "need": "have the agent research, propose, wait for my approval, build, test, and then report",
    "reason": "I must see and approve the plan before anything changes, and I need proof at the end",
    "ac": [
        "The agent never changes anything before I approve the plan.",
        "The agent tests what it built and reports the result.",
        "If a step fails, the agent stops and tells me plainly.",
    ]})

B("h2", "1.2 Tools to build")
B("p", "To fulfill the Epic to Quest use case, we must build this set of tools. We already proved them in the project we built. We bring across the ones the use case needs. We do not bring FHIR server creation. The MVP does not create FHIR servers.")
B("table", {
    "headers": ["Tool group", "What the agent can do with it"],
    "rows": [
        ["Productions and interfaces", "Create, edit, update, set up, start, stop, and validate a production."],
        ["Business hosts", "Create, edit, update, and set up inbound services and outbound operations."],
        ["HL7 connections over TCP", "Set up inbound and outbound HL7 v2 over a TCP port."],
        ["FHIR connections", "Set up inbound FHIR at a web endpoint, and outbound FHIR to a secured external server."],
        ["DTL transformations", "Create, edit, update, compile, and test data transformations."],
        ["BPL business processes", "Create, edit, and update business processes."],
        ["Routing rules", "Create, edit, and update routing rules."],
        ["Lookup and code tables", "Create, edit, and update lookup tables."],
        ["OAuth security", "Set up OAuth 2.0 for a secured connection."],
        ["File folders", "Set up the folders an interface reads from or writes to."],
        ["Test messages", "Build and send HL7 and FHIR test messages. Validate them. Compare input to output."],
        ["Monitoring and trace", "Read the event log, message status, queues, and throughput. Open the visual trace."],
        ["Catalog search", "Search the catalogs for the right building block or pattern."],
    ],
})
B("p", "Not in the MVP: creating or administering a FHIR server. Every artifact the agent builds is ObjectScript based, so it opens in the Health Connect Cloud graphical tools.")

B("h2", "1.3 Skills to build")
B("p", "The agent needs two kinds of skills. The first kind teaches it how to build things with best practices. The second kind teaches it the healthcare standards. The MVP centers on HL7 v2 and FHIR.")
B("lead", "Build skills (best practices)")
B("table", {
    "headers": ["Skill", "What it teaches"],
    "rows": [
        ["Productions", "How to build and run a production with best practices."],
        ["Business hosts", "How to choose and set up services and operations."],
        ["DTL", "How to build data transformations."],
        ["BPL", "How to build business processes."],
        ["Routing rules", "How to build routing rules."],
        ["OAuth", "How to set up OAuth 2.0 security."],
        ["Connections and adapters", "How to choose and set up inbound and outbound connections, including TCP."],
    ],
})
B("lead", "Healthcare standards")
B("table", {
    "headers": ["Standard", "In the MVP", "What it teaches"],
    "rows": [
        ["HL7 v2", "Core", "What HL7 v2 is and how to use it in Health Connect Cloud."],
        ["FHIR", "Core", "What FHIR is and how to use it in Health Connect Cloud."],
        ["SDA", "Supporting", "The clinical model used to bridge HL7 v2 and FHIR."],
        ["CDA", "Later", "What CDA is and how to use it in Health Connect Cloud."],
        ["X12", "Later", "What X12 is and how to use it in Health Connect Cloud."],
    ],
})
B("lead", "Customer knowledge")
B("table", {
    "headers": ["Skill", "What it teaches"],
    "rows": [
        ["Organization interface template", "How to read the customer's own interface specification. The customer authors this skill. See US-B09."],
    ],
})

B("h2", "1.4 Catalogs to build")
B("p", "The agent searches catalogs to find the right parts. Build two.")
B("table", {
    "headers": ["Catalog", "What it holds", "Why the agent needs it"],
    "rows": [
        ["Interface building blocks", "The services, operations, and adapters available in Health Connect Cloud.", "To find the right inbound or outbound component."],
        ["Transformation patterns", "Known transformation mappings and patterns.", "To reuse proven mappings instead of guessing."],
    ],
})
B("p", "The AI Admin builds and refreshes the catalogs. See Epic 3, story US-B10.")

B("h2", "1.5 Policies to enforce")
B("p", "The agent must enforce these policies on every action.")
B("table", {
    "headers": ["Policy", "What it means"],
    "rows": [
        ["Approval", "The agent pauses before any change. The user must approve it. Nothing is applied without approval."],
        ["Acting as the user", "The agent acts within the signed in user's permissions. It can never do something the user could not do by hand."],
        ["Tool and knowledge visibility", "The agent uses only the tools and knowledge the AI Admin enabled. It cannot use anything excluded, even if a user asks for it."],
        ["Respect the context (namespace and user)", "The agent inherits the security of the namespace and the signed in user it is called in. It respects both before its own capabilities. A namespace that does not allow FHIR blocks FHIR, even if the agent can do FHIR. A user who cannot set up OAuth blocks OAuth, even if the tool exists. Capability never overrides permission."],
        ["Integrate with CI/CD (hard requirement)", "Every artifact the agent builds enters the same CI/CD and change control pipeline Health Connect Cloud already uses. The agent never writes around the pipeline and never skips a required review."],
    ],
})

B("h2", "1.6 LLM connection")
B("p", "The agent talks to an AI model. The AI Admin sets up the connection. See Epic 3, stories US-B01 to US-B03.")
B("bullets", [
    "Choose the provider and the model.",
    "Store the credential securely. Never in plain text.",
    "Test the connection. Show available or unavailable.",
])

B("pagebreak")

# ---------------------------------------------------------------- EPIC 2 (CHAT)
B("h1", "Epic 2. Chat Experience")
B("p", "Scope: MVP. Customer facing. This is the Integration Engineer's main surface. Build it second.")

B("h2", "2.1 Building by conversation")
B("story", {"id": "US-A01", "title": "Build an interface by describing it",
    "persona": IE, "need": "describe the interface I want in plain language and have the agent plan and build it",
    "reason": "I am not a developer, and I should not have to assemble an interface screen by screen",
    "ac": [
        "I state the goal in business terms. I get back a clear, ordered plan of what will be built.",
        "The agent recommends the right inbound and outbound components for each leg. It does this before it builds anything.",
        "The agent shows the plan. It waits for my approval before it creates anything.",
        "After I approve, the agent builds the interface. It reports in plain language what it created.",
    ]})
B("story", {"id": "US-A02", "title": "Approve or reject every change",
    "persona": IE, "need": "approve or reject each change before it is applied",
    "reason": "I am accountable for what gets created in my environment, and I must stay in control",
    "ac": [
        "Every change pauses for my approval. This covers create, modify, start, and stop.",
        "The agent shows what will happen before it happens.",
        "Nothing is applied until I approve.",
        "If I reject, the agent acknowledges it. It asks how I want to proceed.",
    ]})
B("story", {"id": "US-A03", "title": "See what the agent is doing",
    "persona": IE, "need": "see, in real time, what the agent is doing and the result of each step",
    "reason": "I need to trust and follow its work, not watch a spinner",
    "ac": [
        "The response appears as it is produced. I never wait at a blank screen.",
        "Each action is shown in plain form. I see what it did, with what inputs, and the result. I can expand it for detail.",
        "Errors are explained in plain language. I do not see technical traces.",
    ]})
B("story", {"id": "US-A04", "title": "Remember the conversation (context memory)",
    "persona": IE, "need": "rely on the agent remembering the context of our conversation",
    "reason": "building an interface is step by step, and I must refer back without repeating myself",
    "ac": [
        "The agent keeps the full context of the current conversation. It builds on earlier turns.",
        "I can say 'now add the outbound side'. The agent knows which interface I mean.",
        "I can leave and come back to the same conversation. The context is still there.",
    ]})
B("story", {"id": "US-A05", "title": "Stay responsive, never hang",
    "persona": IE, "need": "count on the agent staying responsive at all times",
    "reason": "a frozen build is worse than a slow one, and I need to know where I stand",
    "ac": [
        "Long tasks are split into short steps. I see progress at each step.",
        "The agent never looks frozen. If a step runs long, it summarizes progress and offers to continue or retry.",
    ]})

B("h2", "2.2 Working from the organization's specification")
B("story", {"id": "US-A06", "title": "Upload our interface specification",
    "persona": IE, "need": "attach our organization's interface specification documents to the conversation and have the agent build from them",
    "reason": "retyping a multi page specification into a chat box is slow and error prone",
    "ac": [
        "I attach common document and spreadsheet files. Each one shows in a list. I can remove any of them.",
        "The agent reads the documents. It produces a clear, structured summary of what it understands.",
        "The agent does not build anything until I confirm the summary. I can edit the summary. The agent revises it.",
    ]})
B("story", {"id": "US-A07", "title": "Have the agent understand our template",
    "persona": IE, "need": "rely on the agent reading our organization's own interface template correctly",
    "reason": "a generic reading will misread our conventions and produce the wrong mappings",
    "ac": [
        "The AI Admin provides organization specific knowledge. See Epic 3. The agent then reads our template the way we intend.",
        "The summary reflects our template's real meaning. It is not a generic guess.",
        "This is the hand off between the two people. The AI Admin authors the knowledge. The engineer benefits from it on upload.",
    ]})

B("h2", "2.3 Refining the result, by prompt or in the graphical tools")
B("p", "This behavior defines the product for a non developer. The engineer builds with the agent. The engineer then keeps working in the graphical tools. The engineer goes back to the agent when useful. There is no code at any point.")
B("story", {"id": "US-A08", "title": "Refine by prompting",
    "persona": IE, "need": "ask the agent to change something it just built",
    "reason": "for most changes, talking is faster than hand editing",
    "ac": [
        "The agent finds the existing artifact. It proposes the change. It applies the change after I approve. It re checks the result.",
    ]})
B("story", {"id": "US-A09", "title": "Open any result in the graphical editors",
    "persona": IE, "need": "open any artifact the agent created directly in the matching Health Connect Cloud graphical editor with one action",
    "reason": "for detailed work I prefer the visual tools, and I must never read or edit source code",
    "ac": [
        "One action opens the right editor. This works for data transformations, business process logic, routing rules, production configuration, lookup tables, security configuration, and host settings.",
        "The link opens the exact artifact. It does not open a generic landing page.",
        "I never see or edit source code.",
        "Every artifact opens because the agent builds it in ObjectScript.",
    ]})
B("story", {"id": "US-A10", "title": "Edit in a graphical tool, then keep talking to the agent",
    "persona": IE, "need": "edit an artifact in a graphical editor and then keep working with the agent using my current changes",
    "reason": "I move between chat and visual editing, and the agent must never use an old copy",
    "ac": [
        "I edit an artifact in the visual tool. I then ask the agent about it. The agent reflects my current changes.",
        "The agent can summarize, validate, or keep building on my edits.",
    ]})
B("story", {"id": "US-A11", "title": "Inspect runtime behavior through the existing tools",
    "persona": IE, "need": "jump from the chat to message trace, the event log, and message search for the interface I built",
    "reason": "checking real message flow is a visual, click through task",
    "ac": [
        "After a test, the agent links me to the visual message trace. It links me to the event log for this interface. It links me to message search for the message type.",
        "The agent can also answer 'did my message reach the destination'. It then offers the visual link so I can confirm.",
    ]})

B("h2", "2.4 Testing and validating from the chat")
B("story", {"id": "US-A12", "title": "Send a test message",
    "persona": IE, "need": "ask the agent to send representative test messages through the interface",
    "reason": "I must prove the path works before I tell the destination we are live",
    "ac": [
        "The agent generates and sends a representative HL7 message. It also sends a representative FHIR resource.",
        "The agent reports whether each one reached the destination. It surfaces any error.",
    ]})
B("story", {"id": "US-A13", "title": "Validate the output against the destination's specification",
    "persona": IE, "need": "have the agent validate the outbound result and compare input to output",
    "reason": "the destination rejects anything that does not match its format, and I must catch that now",
    "ac": [
        "The agent validates the structure and meaning of the outbound message.",
        "The agent compares the inbound message to the outbound message. It flags fields the destination expects that are missing or wrong.",
    ]})

B("h2", "2.5 Managing conversations")
B("story", {"id": "US-A14", "title": "Keep, search, resume, and label conversations",
    "persona": IE, "need": "review, search, resume, and rename my past conversations",
    "reason": "building an interface spans several sessions, and I need continuity",
    "ac": [
        "A history view lists past conversations. I can search, resume, and rename them.",
        "A clear control starts a new conversation.",
        "A clear indicator shows whether the agent is available.",
    ]})
B("story", {"id": "US-A15", "title": "Start from a guided prompt",
    "persona": IE, "need": "start from guided example prompts",
    "reason": "a blank box is intimidating, and good examples show me what is possible",
    "ac": [
        "Example prompts appear on a new conversation. They fill the composer when I pick one.",
        "At least one example shows a source to destination build like the Epic to Quest case.",
    ]})

B("pagebreak")

# ---------------------------------------------------------------- EPIC 3 (AI SETTING)
B("h1", "Epic 3. AI Setting Experience")
B("p", "Scope: MVP. Built and operated by InterSystems. Not shown to customers in the MVP. The AI Admin assembles the agent here. Build it third. It does not need to be fully featured or customer grade for the MVP. A later release can open it to customers.")

B("h2", "3.1 Connecting the product to an AI model")
B("story", {"id": "US-B01", "title": "Connect to the chosen AI provider",
    "persona": AD, "need": "connect the product to the chosen AI model provider in the interface",
    "reason": "the organization decides which model powers the agent, and I must set it up without code or config files",
    "ac": [
        "I choose a provider and a model in a form. The choice is saved as configuration.",
        "The organization controls which provider is used. It therefore controls where its data goes.",
    ]})
B("story", {"id": "US-B02", "title": "Store credentials securely",
    "persona": AD, "need": "enter the provider credential in a masked field and have it stored securely",
    "reason": "secrets must never be visible, exported, or stored in plain text, and I am accountable for that",
    "ac": [
        "The credential is masked when I enter it. It is stored in an encrypted secret store.",
        "It is never shown, returned, or written to any log or record after I save it.",
    ]})
B("story", {"id": "US-B03", "title": "Test the connection and see its status",
    "persona": AD, "need": "test the connection and see a clear available or unavailable status",
    "reason": "I must know the model is reachable before engineers depend on it",
    "ac": [
        "A clear status shows available, unavailable, or not yet tested.",
        "On failure, the agent shows the provider's exact error message.",
        "The chatbot shows this status to end users. If the model is down, engineers are told plainly.",
    ]})

B("h2", "3.2 Assembling the agent")
B("story", {"id": "US-B04", "title": "Choose which tools and skills the agent has",
    "persona": AD, "need": "choose exactly which tools and which skills an agent can use",
    "reason": "the agent's skill and its safety both depend on the right scope and nothing more",
    "ac": [
        "I see the full pool of tools and skills. I include or exclude each one per agent.",
        "An excluded tool cannot be used, even if a user asks for it.",
        "Changes take effect for new conversations. No redeployment is needed.",
    ]})
B("story", {"id": "US-B05", "title": "Tune the agent's instructions",
    "persona": AD, "need": "write and tune the agent's instructions in the interface",
    "reason": "instruction quality is the biggest lever on output quality, and I must tune it without a developer",
    "ac": [
        "I edit the agent's instructions, persona, planning rules, and guardrails.",
        "I save my changes. I can revert to the default.",
        "My changes persist across product updates.",
    ]})
B("story", {"id": "US-B06", "title": "Group tools and toggle whole areas",
    "persona": AD, "need": "group related tools and turn whole areas on or off for an agent",
    "reason": "I want to enable or disable a whole area with one switch",
    "ac": [
        "I enable or disable each tool area. I can give it a description.",
    ]})

B("h2", "3.3 Authoring skills")
B("story", {"id": "US-B07", "title": "Author a skill from a document",
    "persona": AD, "need": "create a skill by writing or uploading a plain language document",
    "reason": "my expertise lives in documents, and I should turn it into agent knowledge without code",
    "ac": [
        "I create a skill with a name, a description, and plain language content. I paste it or upload it.",
        "The skill becomes available to attach to any agent.",
        "No coding is required.",
    ]})
B("story", {"id": "US-B08", "title": "Register a supplied skill",
    "persona": AD, "need": "register skills supplied by the vendor or a developer, in the same interface",
    "reason": "some skills come pre built, and I must make them available without rebuilding them",
    "ac": [
        "Self authored skills and supplied skills appear in one pool.",
        "I attach both kinds the same way.",
    ]})
B("story", {"id": "US-B09", "title": "Teach the agent our interface template",
    "persona": AD, "need": "author a skill that explains how to read our own interface template, and attach it to the agent",
    "reason": "when engineers upload that template, the agent must read our conventions correctly to produce the right mappings",
    "ac": [
        "I write the skill in plain language. It describes our template's structure and conventions.",
        "I attach it to the engineers' agent.",
        "After I attach it, uploaded templates are read our way.",
        "This is the worked example of a skill that customers build for themselves.",
    ]})

B("h2", "3.4 Managing the catalogs")
B("story", {"id": "US-B10", "title": "Provide and refresh the catalogs",
    "persona": AD, "need": "provide and refresh the two catalogs the agent uses, the interface building blocks and the transformation patterns, and choose which catalogs each agent can use",
    "reason": "the agent can only recommend the right part if the catalogs are present, current, and relevant",
    "ac": [
        "I see each catalog's size and last update time. I refresh it on demand.",
        "I run a test search to check quality before I expose a catalog.",
        "I bind specific catalogs to specific agents.",
    ]})

B("h2", "3.5 Oversight")
B("story", {"id": "US-B11", "title": "Review the audit trail",
    "persona": AD, "need": "review a complete record of every action the agent took for a user",
    "reason": "I am accountable for what was done, by whom, and where, for compliance and troubleshooting",
    "ac": [
        "The record shows who, when, which environment, what action, and the result.",
        "I can filter to errors only.",
        "The record captures successful actions and denied attempts.",
    ]})
B("story", {"id": "US-B12", "title": "Verify a tool before exposing it",
    "persona": AD, "need": "inspect and safely try a tool before I expose it to engineers",
    "reason": "I want to confirm it behaves as intended",
    "ac": [
        "I browse each tool with its description and inputs.",
        "I safely try a non destructive tool with sample input.",
    ]})

B("pagebreak")

# ---------------------------------------------------------------- WALKTHROUGH
B("h1", "4. End to end walkthrough")
B("p", "This is the acceptance narrative. It shows both people at work.")
B("lead", "AI Admin. One time set up in the AI Setting Experience. Internal to InterSystems.")
B("numbers", {"start": 1, "items": [
    "Connects the product to the chosen AI model. Stores the credential securely. Tests the connection until it shows available.",
    "Provides and refreshes the catalogs of building blocks and transformation patterns. Confirms relevance with a test search.",
    "Authors the organization specific template skill. Attaches it to the engineers' agent.",
    "Assembles the agent. Selects the tools and skills it needs. Tunes its instructions. Binds the catalogs.",
]})
B("lead", "Integration Engineer. The daily work in the chat. Customer facing.")
B("numbers", {"start": 5, "items": [
    "Opens the chatbot. Attaches the Epic to Quest interface specification.",
    "The agent reads it. It applies the template skill. It shows a structured summary. The engineer edits one line and confirms.",
    "The engineer says 'build it'. The agent proposes a plan: the two inbound legs, the business process and transformations, and the two outbound legs. It asks for approval.",
    "The engineer approves step by step. The agent builds. It validates as it goes. It keeps full context throughout.",
    "The engineer opens the new transformation in the graphical editor. The engineer adjusts one mapping and saves it. The engineer asks the agent to re validate. The agent re validates against the current version.",
    "The agent sends a test HL7 message and a test FHIR resource. It validates them against Quest's requirements. It compares input to output. It links to the visual trace.",
    "Everything passes. The engineer starts the interface.",
]})
B("lead", "AI Admin. Oversight.")
B("numbers", {"start": 12, "items": [
    "Reviews the audit trail. Sees every action, by this engineer, in this environment.",
]})

# ---------------------------------------------------------------- SUMMARIES
B("h1", "5. What to build, summaries")

B("h2", "5.1 The Agent (Epic 1). MVP. Build first.")
B("bullets", [
    "Tools: the tool groups in section 1.2. The key ones for the use case are inbound and outbound HL7 over TCP, inbound and outbound FHIR, the business process, and the data transformations.",
    "Skills: how to build productions, business hosts, DTL, BPL, routing rules, OAuth, and connections. Plus HL7 v2 and FHIR for the MVP. Plus the customer template skill.",
    "Catalogs: interface building blocks, and transformation patterns.",
    "Policies: approval, acting as the user, and tool and knowledge visibility.",
    "LLM connection: provider, model, secured credential, and health test.",
    "Artifact form: every artifact is ObjectScript based, so it opens in the graphical tools.",
])

B("h2", "5.2 Chat Experience (Epic 2). MVP. Customer facing.")
B("table", {
    "headers": ["Capability", "Stories"],
    "rows": [
        ["Research, propose, approve, build, test, report", "US-1.01, US-A01, US-A02"],
        ["Real time view of actions and results", "US-A03"],
        ["Conversation context memory", "US-A04"],
        ["Always responsive, no hanging", "US-A05"],
        ["Upload a specification, confirm understanding before building", "US-A06, US-A07"],
        ["Refine by prompt", "US-A08"],
        ["Open any result in the graphical editors, one action, no code", "US-A09"],
        ["Round trip back from visual edits, no stale state", "US-A10"],
        ["Links to message trace, event log, message search", "US-A11"],
        ["Test, validate, and compare in the chat", "US-A12, US-A13"],
        ["Conversation history and availability indicator", "US-A14"],
        ["Guided starter prompts", "US-A15"],
    ],
})

B("h2", "5.3 AI Setting Experience (Epic 3). MVP. Internal to InterSystems.")
B("table", {
    "headers": ["Area", "Stories"],
    "rows": [
        ["Connection: provider, model, secure credential, status", "US-B01, US-B02, US-B03"],
        ["Agent: tool and skill selection, instruction tuning", "US-B04, US-B05, US-B06"],
        ["Skills: author from a document, register supplied, organization template", "US-B07, US-B08, US-B09"],
        ["Catalogs: provide, refresh, test, bind", "US-B10"],
        ["Oversight: audit trail, tool check", "US-B11, US-B12"],
    ],
})

# ---------------------------------------------------------------- DoD
B("h1", "6. Definition of Done (MVP)")
B("p", "The MVP is done when all of these are true on Health Connect Cloud.")
B("numbers", {"start": 1, "items": [
    "The agent follows the safe loop: research, propose, wait for approval, build, test, report.",
    "An Integration Engineer uploads the Epic to Quest spec. The agent reads it through the organization template skill and shows a correct summary.",
    "From the chat, the agent builds the full interface: both inbound legs, the business process and transformations to Quest's format, and both outbound legs, including the secured external FHIR. The engineer approves every change.",
    "The engineer opens any result in the matching graphical editor with one action and edits it with no code. The agent then works from those edits.",
    "The agent sends test HL7 and FHIR messages, validates them, and the engineer confirms flow in the visual trace and the event log.",
    "Every artifact is ObjectScript based and opens in its graphical editor.",
    "The agent respects the calling namespace and the signed in user. It never exceeds what they allow, even when it has the capability.",
    "Every artifact the agent builds enters the Health Connect Cloud CI/CD and change control pipeline. Nothing is written around the pipeline. Each change is attributed to the signed in user.",
    "The agent is connected to the chosen AI model. The credential is secured. The status shows available.",
    "An InterSystems AI Admin can assemble the agent in the AI Setting Experience: connect the LLM, build the catalogs, pick the tools and skills, set the policies, and write the prompt.",
    "Every action is recorded in the audit trail.",
    "The solution uses only what Health Connect Cloud provides. No CLI. No file server access. No developer tools.",
]})
B("callout", "Scope reminder. The MVP has three parts. The Agent and the Chat Experience are customer facing. The AI Setting Experience is built and operated by InterSystems, and it is not shown to customers in the MVP. The AI Admin for this product is an InterSystems person. More agents come later.")

# ---------------------------------------------------------------- NFR
B("h1", "7. Non functional requirements")
B("table", {
    "headers": ["Requirement", "Target"],
    "rows": [
        ["Responsiveness", "Visible output starts within a couple of seconds. The agent stays responsive during a task."],
        ["Reliability", "The agent never fails silently. It explains every error in plain language."],
        ["Security, acting as the user", "The agent acts within the signed in user's permissions. It can never do something the user could not do by hand."],
        ["Security, change gate", "Every change needs the user's explicit approval before it is applied."],
        ["Security, secrets", "Credentials are stored encrypted. They are never shown, exported, or logged."],
        ["Security, inherit the context", "The agent runs inside the calling namespace and the signed in user. It cannot exceed what that namespace allows or what that user is permitted to do. It refuses in plain language when a restriction applies."],
        ["Auditability", "Every action taken for a user is recorded and can be reviewed."],
        ["Data control", "The organization chooses the AI provider. It therefore chooses where its data is processed."],
        ["Usability", "A non developer can use the product. No source code is visible in the engineer's workflow."],
        ["Platform constraints", "The product uses only what Health Connect Cloud provides. No command line. No file server access. No local developer tools."],
        ["Artifact form", "The agent builds ObjectScript based artifacts. Each one opens in a Health Connect Cloud graphical tool. Python based artifacts are not used where a graphical view is needed."],
        ["Change control and CI/CD (hard requirement)", "Every artifact the agent builds enters the same CI/CD and change control pipeline Health Connect Cloud already uses. Each change is versioned, tracked, attributed to the signed in user, and promoted through the normal pipeline. The agent never writes around the pipeline or skips a required review."],
    ],
})


# ----------------------------------------------------------------------------
# RENDER HELPERS
# ----------------------------------------------------------------------------

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
    tcPr.append(sh)

def set_cell_text(cell, text, bold=False, color=INK, white=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text); run.font.size = Pt(size); run.font.bold = bold
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    run.font.name = "Calibri"

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, HEADER_FILL); set_cell_text(c, h, bold=True, white=True)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if r_i % 2 == 1:
                shade(cells[i], ZEBRA_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13.5, NAVY), ("Heading 3", 11.5, SLATE)]:
        st = doc.styles[name]
        st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color

def num_parts(payload):
    if isinstance(payload, dict):
        return payload.get("start", 1), payload["items"]
    return 1, payload

def render_docx():
    doc = Document()
    style_doc(doc)
    for sec in doc.sections:
        sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
        sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)

    for kind, payload in BLOCKS:
        if kind == "title":
            p = doc.add_paragraph(); r = p.add_run(payload)
            r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY; r.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(2)
        elif kind == "subtitle":
            p = doc.add_paragraph(); r = p.add_run(payload)
            r.font.size = Pt(13); r.font.color.rgb = SLATE; r.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(10)
        elif kind == "meta":
            t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
            for k, v in payload:
                cells = t.add_row().cells
                shade(cells[0], ZEBRA_FILL); set_cell_text(cells[0], k, bold=True)
                set_cell_text(cells[1], v)
            doc.add_paragraph()
        elif kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "h3":
            doc.add_heading(payload, level=3)
        elif kind == "p":
            p = doc.add_paragraph(payload); p.paragraph_format.space_after = Pt(6)
        elif kind == "lead":
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(payload); r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)
        elif kind == "callout":
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
            r = p.add_run(payload); r.font.italic = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:space'), '10'); left.set(qn('w:color'), '153E6B')
            pbdr.append(left); pPr.append(pbdr)
        elif kind == "bullets":
            for item in payload:
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "numbers":
            start, items = num_parts(payload)
            for i, item in enumerate(items, start):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(f"{i}.  {item}"); r.font.size = Pt(10.5)
        elif kind == "table":
            add_table(doc, payload["headers"], payload["rows"])
        elif kind == "story":
            s = payload
            doc.add_heading(f"{s['id']}: {s['title']}", level=3)
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"As {s['persona']}, I need to be able to {s['need']}, because {s['reason']}.")
            r.font.italic = True; r.font.size = Pt(10.5)
            lbl = doc.add_paragraph(); lbl.paragraph_format.space_after = Pt(0)
            lr = lbl.add_run("Acceptance criteria"); lr.font.bold = True; lr.font.size = Pt(10)
            for item in s["ac"]:
                doc.add_paragraph(item, style="List Bullet")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "pagebreak":
            doc.add_page_break()

    doc.save(DOCX_OUT)


def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in rows:
        out.append("| " + " | ".join(c.replace("\n", " ") for c in row) + " |")
    return "\n".join(out)

def render_md():
    L = []
    for kind, payload in BLOCKS:
        if kind == "title": L.append(f"# {payload}\n")
        elif kind == "subtitle": L.append(f"> {payload}\n")
        elif kind == "meta":
            for k, v in payload: L.append(f"- **{k}:** {v}")
            L.append("")
        elif kind == "h1": L.append(f"\n## {payload}\n")
        elif kind == "h2": L.append(f"\n### {payload}\n")
        elif kind == "h3": L.append(f"\n#### {payload}\n")
        elif kind == "p": L.append(payload + "\n")
        elif kind == "lead": L.append(f"**{payload}**\n")
        elif kind == "callout": L.append(f"> {payload}\n")
        elif kind == "bullets":
            for item in payload: L.append(f"- {item}")
            L.append("")
        elif kind == "numbers":
            start, items = num_parts(payload)
            for i, item in enumerate(items, start): L.append(f"{i}. {item}")
            L.append("")
        elif kind == "table":
            L.append(md_table(payload["headers"], payload["rows"]) + "\n")
        elif kind == "story":
            s = payload
            L.append(f"\n#### {s['id']}: {s['title']}\n")
            L.append(f"*As {s['persona']}, I need to be able to {s['need']}, because {s['reason']}.*\n")
            L.append("Acceptance criteria:")
            for item in s["ac"]: L.append(f"- {item}")
            L.append("")
        elif kind == "pagebreak": L.append("\n---\n")
    with open(MD_OUT, "w") as f:
        f.write("\n".join(L).strip() + "\n")


if __name__ == "__main__":
    render_docx()
    render_md()
    print("Wrote:")
    print(" ", DOCX_OUT)
    print(" ", MD_OUT)
