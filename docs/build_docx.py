#!/usr/bin/env python3
"""Generate PRD.docx from the PRD content."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    doc.add_paragraph('')

def bold_para(label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(' ' + text)

def bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Health Interop AI Copilot', level=1)
subtitle = doc.add_paragraph('Product Requirements Document')
subtitle.style = doc.styles['Normal']
run = subtitle.runs[0]
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

# ============================================================
# 1. PRODUCT OVERVIEW
# ============================================================
doc.add_heading('1. Product Overview', level=2)

doc.add_heading('What it is', level=3)
doc.add_paragraph(
    'Health Interop AI Copilot is an AI-powered assistant embedded inside InterSystems IRIS for Health '
    'that helps integration engineers build, review, and optimize healthcare interoperability workflows '
    'through natural conversation. Instead of navigating dozens of Management Portal screens and writing '
    'ObjectScript by hand, engineers describe what they need in plain English and the copilot builds it '
    'using real IRIS APIs.'
)

doc.add_heading('Who it serves', level=3)
doc.add_paragraph(
    'System integrators and integration engineers working with IRIS for Health who need to stand up HL7, '
    'FHIR, and SDA interoperability pipelines. The primary persona has healthcare data experience but '
    'limited InterSystems platform knowledge. The copilot bridges that gap by encoding IRIS best practices '
    'into its behavior.'
)

doc.add_heading('How it works', level=3)
doc.add_paragraph(
    'A single orchestrator agent (Health Interop) receives the user\'s request, searches a vector catalog '
    'of IRIS classes to find the right components, presents a plan, and on approval executes the build using '
    'tools organized across toolsets. Domain skills (e.g., Productions, DTL, BPL, Routing Rules, HL7v2, '
    'FHIR R4, SDA, REST) give the agent deep knowledge of IRIS-specific concepts. The agent operates through '
    'MCP servers (e.g., Production, Transform, Testing, Catalog) that scope tool access by domain.'
)

doc.add_heading('Two personas, two experiences', level=3)
doc.add_paragraph('The product separates two distinct user journeys:')

bold_para('Developer Experience (DX) —',
    'InterSystems engineers and partners who author the underlying capabilities: writing Tool classes in '
    'ObjectScript/Python, authoring Skill documents, building catalog embeddings. This work happens in '
    'VS Code (or any IDE) and ships as compiled classes inside an IPM package. Developers define what the '
    'copilot can do.')

bold_para('Builder Experience (End User) —',
    'Integration engineers inside IRIS for Health and Health Connect who configure and use the copilot: '
    'creating Agents with custom system prompts, assembling MCP Servers from available Toolsets, linking '
    'Skills to Agents, and chatting with the copilot to build productions. This work happens entirely in '
    'the IRIS Management Portal UI. Builders decide how the copilot behaves for their use case.')

# ============================================================
# 2. USE CASES
# ============================================================
doc.add_heading('2. Use Cases', level=2)

doc.add_heading('Use Case 1: Guided Production Build (Simple to Complex)', level=3)
doc.add_paragraph(
    'A user types: "I need to receive ADT admission messages from our HIS and send observation reports '
    'to the downstream LIS."'
).italic = False

bold_para('Step 1 — Prompt Refinement.',
    'The agent identifies missing information and asks targeted questions: What HL7 version? What transport '
    '(MLLP port, file drop, TCP)? What fields need to map from ADT to ORU? What happens to messages that fail '
    'transformation? What is the throughput expected? Do you need OAuth integration? The agent guides the user '
    'to a complete specification without requiring them to know IRIS terminology.')

bold_para('Step 2 — Catalog Search and Plan.',
    'The agent calls a catalog (e.g., Ens.* and HS.* classes) — a vector database, for example — '
    'to find the right business host classes, adapter classes, and existing transformation templates. It presents '
    'a structured plan: every component name, the class it will use (citing which catalog result and why), key '
    'settings, data flow, field mappings, and the complete test message. Ends with "Ready to build. Shall I proceed?"')

bold_para('Step 3 — Build.',
    'On approval, the agent executes in a strict sequence: Call MCP Servers and Tools to build the production '
    'end to end. Number of tool calls must be defined (less than 10 based on experience). For complex '
    'productions this includes:')

bullet('DTL (Data Transformation Language) with complete field mappings, foreach blocks for repeating segments, '
       'conditional logic, and subtransform references')
bullet('BPL (Business Process Language) for orchestration workflows with compensation handlers, async callbacks, '
       'and code activities')
bullet('Routing Rules with constraint-based message routing, transform chaining, and dead-letter handling')
bullet('Lookup Tables for code mappings (e.g., facility codes, provider identifiers, insurance plan maps)')
bullet('HL7 Schema validation ensuring source and target structure paths are correct before any mapping')
bullet('FHIR resource construction for R4 bundles, references, and search parameter configuration')
bullet('SDA (Summary Document Architecture) as the intermediary hub for cross-format transformations '
       '(HL7 → SDA → FHIR)')

doc.add_paragraph(
    'The tools must understand these artifacts as a connected system, not isolated components. A routing rule '
    'references a DTL by class name, a DTL references lookup tables by name, a BPL calls sub-DTLs and routes '
    'to specific business hosts. The build sequence enforces these dependencies.'
)

bold_para('Step 4 — Test and Validate.',
    'The agent creates sample messages (1 to 10 based on customer request) and sends test HL7 messages to the '
    'Business Service indicated in the production to check if: production exists and is running, all hosts are '
    'enabled, no event log errors, messages flowed through the pipeline. Failed checks are fixed silently and '
    'validation re-run until all pass. The completion report shows the actual test message sent, transformation '
    'output, and other metrics like transformation time, end-to-end time, and errors.')

bold_para('Step 5 — Catalog Update.',
    'After a successful build, any new DTL, BPL, or routing rule classes created during the session are '
    'automatically indexed into the HS.* catalog. This ensures the next user who asks for a similar '
    'transformation can discover and reuse what was just built, rather than creating it from scratch. '
    'The catalog grows with every successful build.')

doc.add_heading('Use Case 2: Production Review and Optimization', level=3)
doc.add_paragraph(
    'A user asks: "Review the LAB.Production and tell me what it does and how to improve it."'
)
doc.add_paragraph(
    'The agent calls MCP Server and Tools (e.g., Production and GetProduction) to inspect every business host, '
    'its class, settings, and connections. It calls MCP Server and Tools (e.g., QueryEventLog and MessageSummary) '
    'for recent error patterns and throughput data. It then explains in plain language: what each host does, the '
    'data flow from source to destination, which settings are at defaults versus customized, and which features '
    'are unused (pool sizing, retry intervals, alert triggers, archive flags). Recommendations are grounded in '
    'the Ens.* catalog descriptions and the Skills knowledge base. Each recommendation cites the specific setting, '
    'its current value, the recommended value, the expected impact, and possibly the link to the documentation.'
)

doc.add_heading('Use Case 3: Complex HL7-to-HL7 Transformations', level=3)
doc.add_paragraph(
    'A user needs to transform HL7 v2.5 ADT^A01 messages into ORU^R01 observation reports with field mappings '
    'that span multiple segments.'
)
doc.add_paragraph(
    'The agent calls MCP Server and Tools (e.g., GetHL7SchemaMap) for both source (ADT_A01) and target (ORU_R01) '
    'structures before writing any DTL. This is critical because ADT_A01 has flat PID paths while ORU_R01 has '
    'nested paths like PIDgrpgrp(1).PIDgrp.PID. Guessed paths produce silent empty output.'
)
doc.add_paragraph(
    'The agent builds the complete DTL XML (e.g., CreateDTL) with all field mappings. By experience, do not use '
    'iterative UpdateDTL calls which burn rate-limited API tokens. It then runs testing tools with a sample message '
    'to verify every mapped field produces output. If any field is empty, the agent checks the schema paths and '
    'corrects them. For complex scenarios (conditional logic, repeating segments, lookup tables), the agent uses '
    'the DTL skill knowledge to generate foreach blocks, subtransform references, and lookup table entries.'
)

# ============================================================
# 3. PRODUCT REQUIREMENTS
# ============================================================
doc.add_heading('3. Product Requirements', level=2)

# 3.1 Performance
doc.add_heading('3.1 Performance', level=3)
add_table(
    ['Requirement', 'Target', 'Rationale'],
    [
        ['Time to first token', 'Under 3 seconds', 'Users lose confidence if the chat feels unresponsive after sending a message'],
        ['Full production build', 'Under 90 seconds wall clock', 'The 7-iteration target with parallel tool batching must complete within the LLM provider\'s rate limits'],
        ['Catalog vector search latency', 'Under 500ms per query', 'search_ens and search_hs run in the first iteration; slow search delays the entire plan phase'],
        ['Tool execution (non-streaming)', 'Under 5 seconds per tool', 'Each tool has a 5–30 second timeout; most should complete well under that'],
        ['SSE stream delivery', 'No gaps longer than 2 seconds', 'Token-by-token streaming with tool-call lifecycle events; gaps make the UI feel frozen'],
        ['Concurrent chat sessions', '10+ simultaneous users', 'Each session is an independent agent instance; IRIS process pooling must handle the load'],
    ]
)

# 3.2 Scalability
doc.add_heading('3.2 Scalability', level=3)
add_table(
    ['Requirement', 'Target', 'Rationale'],
    [
        ['Catalog size', '5,000+ classes indexed', 'Current catalogs cover Ens.* and HS.* classes; must scale as customers add custom classes'],
        ['Catalog auto-growth', 'New artifacts indexed after every successful build', 'DTLs, BPLs, and routing rules created by the copilot are added to the HS.* catalog automatically'],
        ['Tool count per agent', '50+ tools without degradation', 'Currently 48 tools across 5 classes; adding new tool classes must not degrade prompt token budget'],
        ['Namespace independence', 'Any namespace on the instance', 'Tools execute in the user\'s selected namespace via X-IRIS-Namespace header'],
        ['LLM provider flexibility', 'Bedrock, Anthropic, Azure OpenAI', 'Provider is user-configurable with connection health checks; switching providers requires no code changes'],
        ['Skill extensibility', 'Add new skills without redeployment', 'Skills are loaded dynamically from class parameters; new Skill.Base subclasses register automatically'],
    ]
)

# 3.3 Quality
doc.add_heading('3.3 Quality of Outputs', level=3)
add_table(
    ['Requirement', 'Target', 'Rationale'],
    [
        ['Catalog grounding', 'Every component choice cites a catalog search result', 'Prevents the LLM from hallucinating class names'],
        ['Schema accuracy', 'Zero guessed HL7/FHIR paths', 'GetHL7SchemaMap is mandatory before any DTL involving EnsLib.HL7.Message'],
        ['Build correctness', 'PostBuildValidation passes on first run', 'Validation checklist must catch build issues before the user discovers them'],
        ['DTL completeness', 'One-shot DTL creation with all field mappings', 'Iterative UpdateDTL calls cause provider rate limiting and incomplete transforms'],
        ['Cross-artifact integrity', 'Routing rules reference compiled DTLs', 'Dangling references cause silent runtime failures'],
        ['Anti-fabrication', 'No claimed success without tool confirmation', 'Every claim of success must be backed by a tool result'],
        ['Error transparency', 'Failures reported with actionable detail', 'The agent reports the specific error, not a generic "something went wrong"'],
    ]
)

# 3.4 User Experience
doc.add_heading('3.4 User Experience', level=3)

doc.add_heading('3.4.1 Builder Experience (End Users in IRIS for Health)', level=3)
doc.add_paragraph(
    'Builders work inside the IRIS Management Portal. They do not write code. Their experience is entirely '
    'configuration-driven through the admin UI.'
)
add_table(
    ['Requirement', 'Description'],
    [
        ['Agent configuration', 'Create and configure Agents: write or edit system prompts, set temperature and iteration limits, select which MCP Servers and Skills to attach'],
        ['MCP Server assembly', 'Create MCP Servers by selecting from available Toolsets. Each MCP scopes a domain. The builder decides which tool groups the agent can access'],
        ['Toolset browsing', 'Browse available Toolsets and their individual Tools. See tool names, descriptions, and parameter signatures'],
        ['Skill attachment', 'View available Skills (read-only) and attach them to Agents. The builder selects which knowledge domains are relevant'],
        ['Connection management', 'Configure LLM provider connections with a visual health semaphore. Paste API keys, click Test, see green/red status'],
        ['Chat experience', 'Streaming chat with tool-call cards, confirmation gates for mutating operations, guided example prompts'],
        ['Audit visibility', 'View the audit trail: which tools were called, arguments, results, and duration'],
    ]
)

doc.add_heading('3.4.2 Developer Experience (DX — VS Code / IDE)', level=3)
doc.add_paragraph(
    'Developers author the building blocks that builders configure. Their work produces compiled classes '
    'that ship in an IPM package.'
)
add_table(
    ['Requirement', 'Description'],
    [
        ['Tool authoring', 'Write %AI.Tool subclasses in ObjectScript or Embedded Python. Each public ClassMethod becomes a tool the LLM can call. The method Description is the contract with the LLM'],
        ['Skill authoring', 'Write markdown documents as XData blocks in Skill.Base subclasses. Skills are versioned in source control alongside the code'],
        ['Catalog seeding', 'Build and maintain vector catalogs. Walk %Dictionary, curate descriptions, generate embeddings, persist to vector tables'],
        ['Testing and dry-run', 'Test individual tools via the dry-run panel or programmatically. Every tool gets at least one happy-path test'],
        ['IPM packaging', 'All classes, skills, reference data, and install hooks ship as a single IPM module installable into any IRIS for Health namespace'],
    ]
)

# 3.5 Tool Depth
doc.add_heading('3.5 Tool Depth: What the Tools Must Understand', level=3)
doc.add_paragraph(
    'The agent is only as capable as its tools. For end-to-end production building, the tools must deeply '
    'understand the interconnected IRIS artifacts:'
)
add_table(
    ['Domain', 'What the tools must do', 'Why it matters'],
    [
        ['Productions', 'Create production classes, add/remove/configure business hosts with correct settings targets (Host vs Adapter level), manage lifecycle', 'A misconfigured setting target is silently ignored. The tool must know the difference'],
        ['DTL', 'Build complete DTL XML with field mappings, foreach blocks, subtransform references, conditional assign actions, virtual document paths', 'A partial DTL produces silent empty output that the user won\'t catch until production'],
        ['BPL', 'Generate BPL XML with activities (assign, call, code, transform, rule), compensation handlers, async patterns', 'Without BPL tools, the agent can only build simple point-to-point flows'],
        ['Routing Rules', 'Create routing rule XML with constraint-based routing, transform chaining, multi-destination routing, dead-letter fallback', 'Missing a compiled rule class or BusinessRuleName setting causes an OnProcessInput crash'],
        ['Lookup Tables', 'Create and populate lookup tables for code mappings (facility codes, insurance plans, provider identifiers)', 'Without them, DTLs can only do structural mapping, not semantic mapping'],
        ['HL7 Schemas', 'Introspect schema structures for any message type and version, return exact segment paths with nested group navigation', 'ADT_A01 and ORU_R01 have fundamentally different path structures. Guessing paths guarantees empty output'],
        ['FHIR R4', 'Construct FHIR resources, validate against profiles, handle references and contained resources', 'FHIR is the target format for modern interoperability'],
        ['SDA', 'Know the SDA3 model as the transformation hub, understand the HL7→SDA→FHIR pipeline', 'IRIS never converts directly between wire formats. SDA is always the intermediary'],
        ['Catalog', 'Vector search across 2,500+ IRIS classes, class introspection, auto-index new artifacts after successful builds', 'The catalog prevents hallucination and enables reuse'],
    ]
)

# 3.6 Architecture
doc.add_heading('3.6 Architecture: Agents, MCPs, Tools, and Skills', level=3)

bold_para('Agent: Health Interop (Orchestrator) —',
    'The single router agent that receives every user message. It does not answer domain questions directly. '
    'Instead it dispatches to the right combination of tools and skills.')

doc.add_paragraph('Behavior rules:')
bullet('Research before planning (mandatory catalog search in iteration 0)')
bullet('Plan before building (present component list, wait for approval)')
bullet('Build silently on approval (call tools without narrating each step)')
bullet('Validate after building (PostBuildValidation is mandatory, never skipped)')
bullet('Parallel tool batching (independent tools in one round-trip to minimize iterations)')
bullet('Update catalog after building (new DTLs, BPLs, routing rules are indexed for future reuse)')

doc.add_paragraph('')
bold_para('MCP Servers', '')
add_table(
    ['MCP', 'Domain', 'Key Tools'],
    [
        ['Production', 'Production lifecycle, business host CRUD, monitoring', 'CreateProduction, AddBusinessHost, StartProduction, StopProduction, GetProduction, PostBuildValidation'],
        ['Transform', 'DTL/BPL authoring, routing rules, HL7 schema introspection', 'CreateDTL, CreateBPL, CreateRoutingRule, CompileDTL, DryRunDTL, GetHL7SchemaMap, ListLookupTables'],
        ['Testing', 'Message send and validation', 'SendHL7, SendFHIR, ValidateHL7, ValidateFHIR, CompareMessages'],
        ['Catalog', 'Vector search, class introspection, catalog maintenance', 'search_ens, search_hs, DescribeClass, GetUserNamespace, IndexNewArtifacts'],
    ]
)

bold_para('Skills (domain knowledge documents) —',
    'Productions, DTL, BPL, Routing Rules, HL7v2, FHIR R4, SDA, REST in Productions, ESB Patterns. '
    'Each skill is a markdown document injected into the agent\'s system prompt. '
    'Skills are authored by developers (DX) and attached to agents by builders (End User UI).')

doc.add_paragraph('')
bold_para('Catalog Vector Search', '')
add_table(
    ['Catalog', 'Content', 'Embedding'],
    [
        ['Ens Catalog (search_ens)', 'Business hosts, adapters, services, processes, operations', 'FastEmbed AllMiniLML6V2, 384 dimensions'],
        ['HS Catalog (search_hs)', 'Health-specific transformations, FHIR/SDA/HL7 DTLs, mappers', 'FastEmbed AllMiniLML6V2, 384 dimensions'],
    ]
)
doc.add_paragraph(
    'Search is mandatory before any build. The catalog auto-grows: new DTLs, BPLs, and routing rules '
    'created during a build session are embedded and indexed so the next session can discover them.'
)

# Save
out = os.path.join(os.path.dirname(__file__), 'PRD.docx')
doc.save(out)
print(f'Saved: {out}')
