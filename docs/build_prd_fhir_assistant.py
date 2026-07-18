#!/usr/bin/env python3
"""
Build the Product Requirements Document (Word + Markdown) for the
FHIR Assistant.

The FHIR Assistant is the FHIR-platform sibling of the Integration Agentic
Builder. It is the AgenticInterop.Agent.FHIRSpecialist %AI.Agent, surfaced as a
launcher button ("FHIR Assistant") injected into the header of the shipped IRIS
for Health FHIR Server Management portal (/csp/fhir-management). It helps a FHIR
engineer stand up and run a FHIR platform by talking to it: the FHIR R4 server
(repository), loading data, SDA<->FHIR transformation, Bulk FHIR export, and the
FHIR SQL Builder.

Structure mirrors Product_Requirements_Integration_Agentic_Builder (the exemplar):
  Meta -> 1. Introduction -> 2. Personas -> 3. Use case example
  Epic 1. The Agent -> Epic 2. Chat Experience -> Epic 3. AI Setting Experience
  4. Walkthrough -> 5. Definition of Done -> 6. Non functional requirements

Style: plain English, short sentences, no em-dashes. Implementation agnostic,
except the facts the business owner keeps: built on the AI Hub Framework, and it
works through the FHIR Server Management portal and its graphical tools so a
non-developer never touches IRIS internals or source code.

Reuses the exact render helpers and house style from build_prd_aiadmin.py.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX_OUT = os.path.join(HERE, "Product_Requirements_FHIR_Assistant.docx")
MD_OUT = os.path.join(HERE, "Product_Requirements_FHIR_Assistant.md")

NAVY = RGBColor(0x15, 0x3E, 0x6B)
SLATE = RGBColor(0x33, 0x3A, 0x44)
INK = RGBColor(0x1A, 0x1D, 0x22)
HEADER_FILL = "153E6B"
ZEBRA_FILL = "EEF2F7"

FE = "a FHIR Engineer"   # the daily user
AD = "an AI Admin"       # assembles and owns the agent

BLOCKS = []
def B(kind, payload=None): BLOCKS.append((kind, payload))


B("title", "FHIR Assistant")
B("subtitle", "Product Requirements for the MVP. IRIS for Health FHIR Server Management portal.")
B("meta", [
    ("Product", "FHIR Assistant. An AI agent that helps build and run a FHIR platform through conversation."),
    ("Built on", "InterSystems AI Hub Framework."),
    ("Where it lives", "Inside the shipped FHIR Server Management portal (/csp/fhir-management), as a launcher in the header."),
    ("Target platform (MVP)", "IRIS for Health and Health Connect, including Health Connect Cloud."),
    ("Primary personas", "AI Admin. FHIR Engineer / Data Engineer."),
    ("Status", "Draft for review."),
])

# ============================================================ 1. INTRODUCTION
B("h1", "1. Introduction")
B("p", "This document specifies the FHIR Assistant. The FHIR Assistant is an AI agent. It helps a healthcare team build and run their FHIR platform by talking to it in plain language, instead of clicking through the FHIR Server Management screens and reading InterSystems documentation. In this document, the “AGENT” refers to the FHIR Assistant.")
B("p", "The FHIR Assistant is an implementation of the InterSystems AI Hub Framework. The AI Hub Framework is the foundation. It provides the core building blocks. These include agents, connections to AI models, capabilities, knowledge modules, and catalogs. The FHIR Assistant uses those building blocks for a specific job. It helps a FHIR engineer stand up a FHIR R4 server, load data into it, query it, transform between FHIR and the clinical model, and share it, all by describing what they want.")
B("p", "The AI Hub on its own is a foundation. The FHIR Assistant is what turns that foundation into value for FHIR work. It applies the framework to a real job, and it does that job from inside the tool the user already has open: the FHIR Server Management portal. The framework makes the agent possible. The FHIR Assistant makes it useful.")
B("p", "The FHIR Assistant is a sibling of the Integration Agentic Builder. The Integration Agentic Builder builds interfaces (productions, transformations, routing). The FHIR Assistant runs the FHIR platform (the repository, data loading, transformation, bulk export, and SQL projection). They share the same framework, the same safe loop, the same policies, and the same AI Setting Experience. This document states what the FHIR Assistant must do. It does not state how it is engineered.")

B("h2", "1.1 How to read this document")
B("p", "The document is built around two personas. It is grounded in one realistic example. Every story is checked against that example.")
B("bullets", [
    "The AI Admin configures and owns the agent. The AI Admin connects it to an AI model (LLM), chooses the capabilities and knowledge the agent has access to, and tailors it to the organization, for example by capturing FHIR profiling and coding conventions in Skills. All of this happens in the user interface. None of it requires direct access to code.",
    "The FHIR Engineer or Data Engineer uses the agent to build and run the FHIR platform. This person describes what they want in plain language. The agent proposes an implementation, waits for approval, and then acts. This person then refines the result, either by prompting the agent again or by using the graphical tools in the FHIR Server Management portal. This person never touches code and never sees InterSystems internals.",
])
B("p", "Three summaries appear at the end. Each covers a deliverable area. The areas are the Chat Experience, the AI Setting Experience, and the Agent itself, meaning the capabilities, knowledge, and policies the agent needs to finish the example.")

B("h2", "1.2 What the MVP includes")
B("p", "The MVP centers on the FHIR platform lifecycle: create a server, load data, transform, query, and share. One part is internal to InterSystems and is not shown to customers in the MVP.")
B("table", {
    "headers": ["Part", "What it is", "In the MVP", "Shown to customers"],
    "rows": [
        ["The Agent", "The FHIR Assistant itself. Its tools, skills, policies, catalogs, and LLM connection are defined and built.", "Yes. This is the core.", "Yes, through the chat. The customer experiences the agent, not its internals."],
        ["The Chat Experience", "The launcher and slide-in chat panel inside the FHIR Server Management portal, plus the Load FHIR Data and FHIR Server Audit helpers.", "Yes.", "Yes. This is the customer facing surface, embedded in the FHIR portal the engineer already uses."],
        ["The AI Setting Experience", "The screens where an AI Admin assembles the agent. For this product the AI Admin is an InterSystems person.", "Partially.", "No. InterSystems operates it. Customers do not see it in the MVP."],
        ["Other agents", "More specialized agents beyond the FHIR Assistant, such as the Integration Agentic Builder or monitoring agents.", "No. Those are separate deliverables.", "No."],
    ],
})

B("h2", "1.3 Platform constraints")
B("p", "The solution must use only what the managed platform provides. On Health Connect Cloud the user works in a browser. They do not get server access.")
B("bullets", [
    "Do not rely on command line access (CLI). The user does not get a shell.",
    "Do not rely on direct file system access. The user cannot reach the server file system. This matters for FHIR because the bulk directory loader reads files from the IRIS server itself, not from a user's desktop. The Load FHIR Data helper exists precisely to bridge this: it stages uploaded files into a server-readable folder so the agent can load them.",
    "Do not rely on a local development environment. The user does not install or run developer tools.",
    "Do not require anything the managed service does not offer. If the platform does not provide it, the product must not depend on it.",
])
B("p", "These constraints shape the whole design. The agent must do its work through the same surfaces a FHIR Server Management user already has.")

B("h2", "1.4 Working through the FHIR Server Management portal and its graphical tools")
B("p", "The FHIR Engineer is not an InterSystems developer. This person wants to keep using the graphical tools the FHIR Server Management portal already provides: the endpoint configuration screens, the capability statement viewer, the resource repository, profile and package management, and the FHIR SQL Builder wizard.")
B("p", "This sets a hard requirement for the agent. The agent must operate through those same surfaces, and it must hand the user off to the matching graphical tool for detailed work. When the agent creates or configures an endpoint, the user can open that endpoint in the portal. When the user wants to project FHIR data to SQL, the agent walks them through the FHIR SQL Builder, because the FHIR SQL Builder has no programmatic interface and is used through its own wizard (Analysis, then Specification, then Projection).")
B("p", "The user never reads or edits InterSystems source code, and never sees internal class names. The agent speaks only in FHIR and platform terms: endpoint paths, namespaces, resource types and counts, storage strategy, profiles, and plain outcomes.")

# ================================================================ 2. PERSONAS
B("h1", "2. Personas")

B("h2", "2.1 AI Admin")
B("table", {
    "headers": ["Attribute", "Detail"],
    "rows": [
        ["Who they are", "An engineer or a system administrator the organization picks to configure and own the agent. This person does not have to be a software developer. For the MVP, this is an InterSystems developer or product manager."],
        ["What they do", "They connect the product to an AI model. They assemble the FHIR Assistant from the available capabilities (tools, skills, policies). They author organization specific skills, for example the organization's FHIR profiling and coding conventions. They tune the agent's prompt. They choose which catalogs the agent can use."],
        ["Primary interface", "The product's AI Hub configuration screens. User interface only. No coding. No developer tools."],
        ["Hard rule", "Everything is point and click. They configure the agent from the tools, skills, and policies available in the system. They never write or compile code."],
        ["Deliverable", "A ready to use FHIR Assistant, scoped to the exact capabilities, knowledge, and catalogs the organization wants."],
    ],
})
B("p", "For the FHIR Assistant, the AI Admin is an InterSystems person. InterSystems assembles and operates the agent for the MVP. The AI Setting Experience is built, but it is internal and does not need to be exposed to customers in the MVP. A later release can open this role to a customer's own AI Admin.")
B("story", {"id": "Anchor story", "title": "Set up the FHIR Assistant",
    "persona": AD, "need": "set up a FHIR Assistant that helps FHIR engineers build and run their FHIR platform from the capabilities and knowledge already built into the system",
    "reason": "my engineers need an assistant that understands both the FHIR platform and our own conventions, and I must deliver it without writing code",
    "ac": [
        "I connect the agent to an AI model, pick its tools and skills, tune its prompt, and bind its catalogs, all in the interface.",
        "I can author organization specific skills, for example how to read our FHIR data-sharing specification, and attach them to the agent.",
        "I deliver the agent without writing or compiling any code.",
    ]})

B("h2", "2.2 FHIR Engineer / Data Engineer")
B("table", {
    "headers": ["Attribute", "Detail"],
    "rows": [
        ["Who they are", "A healthcare FHIR engineer, systems integrator, or data engineer. An experienced systems integrator who knows FHIR and clinical data, but does not know InterSystems internals and is not an ObjectScript developer."],
        ["What they do", "They stand up and run FHIR servers, load data, project FHIR data to SQL for analytics, transform between FHIR and the clinical model, and share data through bulk export. They describe what they want to the agent, review and approve what it proposes, and refine the result by prompting again or by using the FHIR portal's graphical tools."],
        ["Primary interface", "The FHIR Assistant chat panel inside the FHIR Server Management portal, plus the graphical tools that portal already provides."],
        ["Hard rule", "They never touch code and never see InterSystems internals. They work only with the chat and the FHIR portal's graphical tools: endpoint configuration, the capability statement, the resource repository, profile and package management, and the FHIR SQL Builder wizard."],
        ["Deliverable", "A working, populated, queryable FHIR platform: a running FHIR R4 server with data loaded, projected to SQL where needed, transformed where needed, and shared through a validated bulk export."],
    ],
})

# ========================================================= 3. USE CASE EXAMPLE
B("h1", "3. The use case example")
B("p", "The FHIR Engineer is asked to build it.")
B("p", "Stand up a FHIR R4 data platform for a hospital and put it to work. The hospital wants a FHIR R4 repository they can load their existing patient data into, query with standard SQL for analytics, and share with an external partner through a bulk export. The engineer must create the server, load a batch of FHIR resource files into it, project selected resources to SQL so analysts can query without knowing FHIR, and configure a secured Bulk FHIR export that hands the partner a clean extract. Throughout, the engineer validates resources, checks what the server supports, counts what was loaded, and confirms storage and ingestion are healthy.")
B("table", {
    "headers": ["Stage", "What it does"],
    "rows": [
        ["Create the server", "Create a FHIR R4 endpoint in a FHIR foundation namespace, with the right storage strategy and the R4 (and US Core) profile packages."],
        ["Load the data", "Load a batch of FHIR resource files into the server. Because the user has no file system access, files are staged server-side through the Load FHIR Data helper, then loaded and monitored to completion."],
        ["Validate and inspect", "Read the capability statement, validate representative resources, and count resources by type to confirm the load landed as expected."],
        ["Query with SQL", "Project selected resources (for example Patient, Observation, Condition) to SQL tables with the FHIR SQL Builder, so analysts can query the data over SQL, JDBC, or ODBC without FHIR knowledge."],
        ["Transform where needed", "Where the source is not FHIR (HL7 v2, CDA), bridge it to FHIR through the clinical model (SDA) so it lands in the repository in the right shape."],
        ["Share by bulk export", "Configure and run a Bulk FHIR export ($export) to hand the partner an extract, secured with SMART backend OAuth 2.0, and confirm the export completed."],
        ["Audit", "Confirm the endpoint's storage size, resource counts, and ingestion performance are healthy."],
    ],
})
B("p", "This one example covers a lot. It creates and configures a FHIR server, loads and validates data, projects to SQL, bridges non-FHIR sources through the clinical model, and shares through a secured bulk export. That breadth makes it a good acceptance test for the MVP.")

B("pagebreak")

# ================================================================ EPIC 1
B("h1", "Epic 1. The Agent")
B("p", "Scope: MVP. Highest priority. The Agent is the core of the product. Build it first.")

B("h2", "Security and CI/CD premise")
B("p", "The agent inherits the security model of the context it runs in. It never works outside that context. A chat is always open in a specific namespace and under a specific signed in user. The agent respects both, and it respects them first, before its own capabilities.")
B("p", "The namespace comes first. A FHIR endpoint can only be created in a FHIR foundation namespace, one that already holds FHIR metadata, never an application namespace. If the namespace does not allow a capability, the agent does not use it there, even though it knows how. Second, the user. The agent can only do what the signed in user is allowed to do. If the user cannot stand up an OAuth server, the agent does not stand one up for them, even though the tool exists.")
B("p", "Capability never overrides permission. When the agent must refuse for one of these reasons, it says so in plain language and explains what is missing, instead of failing silently. This is achieved through properly scoped and secured tools.")
B("bullets", [
    "All artifacts and configuration the agent creates must be compatible with the CI/CD and change control pipeline the platform already uses. The agent is used in a DEV or LOCAL environment, not in TEST or PRODUCTION (at least for the MVP). This connection is out of scope for this document, but it is a hard requirement that must be respected when building the agent experience.",
])

B("h2", "1.1 How the agent works")
B("p", "The agent always follows the same safe loop. It researches, proposes, waits for approval, builds, tests, and reports.")
B("table", {
    "headers": ["Step", "What the agent does"],
    "rows": [
        ["1. Research", "Understand the request. Read any uploaded spec or data. Discover FHIR namespaces and existing endpoints, read the capability statement, count resources, check status and metrics. Search the catalogs. All read-only, no permission needed just to look."],
        ["2. Propose", "Show a plan in plain language. State the target namespace, endpoint path, FHIR version and profiles, storage strategy, and every choice it rests on. Recommend a sensible default for anything the user left open, and ask them to confirm or change it."],
        ["3. Approval gate", "Stop and wait. The engineer must approve. The agent creates or changes nothing without approval. Having all the details is not authorization."],
        ["4. Build", "Create or configure the server, load data, project to SQL, or set up the export. Validate as it goes."],
        ["5. Test", "Validate resources, count what loaded, and confirm the export or projection worked. Compare what was asked to what exists."],
        ["6. Report", "Say what now exists and where, in one or two plain sentences. Link to the matching FHIR portal graphical tool where useful."],
    ],
})
B("story", {"id": "US-1.01", "title": "The agent works in a safe loop",
    "persona": FE, "need": "have the agent research, propose, wait for my approval, act, test, and then report",
    "reason": "I must see and approve the plan before anything changes in my FHIR platform, and I need proof at the end",
    "ac": [
        "The agent never creates, configures, loads, resets, or deletes anything before I approve the plan, which is expressed in plain language including the target namespace, endpoint, profiles, and storage strategy.",
        "The agent tests what it did and reports the result, including where the result lives and a link to the matching portal tool where applicable.",
        "If a step fails, the agent stops and tells me plainly what is missing.",
    ]})

B("h2", "1.2 Tools to build")
B("p", "To fulfill the use case, we must build this set of tools. They come from three capability groups: the FHIR Server (repository) tools, the Bulk FHIR export tools, and the Catalog and introspection tools. The FHIR SQL Builder has no programmatic interface, so the agent guides the user through its wizard instead of calling a tool.")
B("table", {
    "headers": ["Tool group", "What the agent can do with it"],
    "rows": [
        ["FHIR namespaces and endpoints", "Discover FHIR-enabled foundation namespaces. List, inspect, create, and configure FHIR R4 endpoints. Read the capability statement. Manage profile and metadata packages."],
        ["FHIR resources", "Search, read, create, update, delete, and validate resources. Count resources by type."],
        ["Loading data", "Load a server-side folder of FHIR files into an endpoint (ordered, asynchronous), monitor ingestion status and metrics, keep a durable run history, bulk load, and reset data."],
        ["Storage and performance audit", "Report an endpoint's storage size, resource counts, and ingestion performance. Probe query performance."],
        ["Bulk FHIR export", "Configure and run Bulk FHIR Coordinator exports that pull resources from a source FHIR endpoint and write them to files or into a target server. List, create, configure, activate, and delete export configurations, start exports, and monitor sessions. Provision the prerequisites end to end: a storage directory, an SSL/TLS configuration, an interop credential, and a SMART backend OAuth server and client."],
        ["FHIR SQL Builder (guided, no API)", "Walk the user through the FHIR SQL Builder wizard: Analysis, then Specification, then Projection. There is no automation interface, so the agent teaches and guides rather than acting."],
        ["Catalog search and introspection", "Search the catalogs for the right FHIR or transformation class. Introspect namespaces and classes. Look up glossary terms."],
    ],
})
B("p", "The Load FHIR Data helper (a menu in the FHIR portal) stages uploaded files into a server-readable folder so the loading tools can reach them. The user never needs file system access.")

B("h2", "1.3 Skills to build")
B("p", "The agent needs two kinds of skills. The first kind teaches it how to run the FHIR platform with best practices. The second kind teaches it the healthcare standards. The MVP centers on FHIR and the clinical model that bridges to it.")
B("lead", "Platform skills (best practices)")
B("table", {
    "headers": ["Skill", "What it teaches"],
    "rows": [
        ["FHIR Server", "How to build and administer the FHIR R4 server: discover the foundation namespace, create and configure endpoints, manage metadata and profile packages, do resource CRUD, search, and validation, read the capability statement, and provision safely."],
        ["Bulk FHIR", "How to configure the Bulk FHIR Coordinator and run $export (system, Patient, Group), using SMART Backend Services, fetch and storage adapters, and the asynchronous REST flow."],
        ["FHIR SQL Builder", "How to project FHIR data into relational SQL: Analysis, then Specification, then Projection, with columns, subtables, and filters, and how to query the result over SQL, JDBC, or ODBC."],
    ],
})
B("lead", "Healthcare standards")
B("table", {
    "headers": ["Standard", "In the MVP", "What it teaches"],
    "rows": [
        ["FHIR R4", "Core", "Resources, references, search parameters, bundles, and SDA/FHIR transformation and SMART/interop."],
        ["SDA", "Core", "The clinical model used as the canonical hub behind any HL7 v2, CDA, or X12 to FHIR conversion."],
        ["HL7 v2", "Supporting", "What HL7 v2 is and how it bridges to FHIR through the clinical model."],
        ["CDA", "Later", "What CDA is and how it bridges to FHIR."],
        ["X12", "Later", "What X12 is and how it relates to the platform."],
    ],
})
B("lead", "Customer knowledge")
B("table", {
    "headers": ["Skill", "What it teaches"],
    "rows": [
        ["Organization data-sharing specification", "How to read the customer's own FHIR data-sharing or profiling specification. The customer authors this skill. See US-B09."],
    ],
})

B("h2", "1.4 Catalogs to build")
B("p", "The agent searches catalogs to find the right parts and mappings. Build two.")
B("table", {
    "headers": ["Catalog", "What it holds", "Why the agent needs it"],
    "rows": [
        ["FHIR and platform classes", "The FHIR server, transformation, and clinical-model classes available in IRIS for Health, searchable by natural language.", "To find the right FHIR or transformation building block instead of guessing at internal names."],
        ["Transformation and mapping patterns", "Field-level mappings between HL7 v2, SDA, FHIR R4, CDA, and X12, through the clinical model.", "To reuse proven mappings when bridging a non-FHIR source into the FHIR repository."],
    ],
})
B("p", "The AI Admin builds and refreshes the catalogs. See Epic 3, story US-B10.")

B("h2", "1.5 Policies to enforce")
B("p", "The agent must enforce these policies on every action.")
B("table", {
    "headers": ["Policy", "What it means"],
    "rows": [
        ["Approval", "The agent pauses before any change. The user must approve it. Nothing is applied without approval."],
        ["Acting as the user", "The agent acts within the signed in user's permissions. It can never do something the user could not do by hand."],
        ["Respect the context (namespace and user)", "The agent inherits the security of the namespace and the signed in user. A FHIR endpoint is created only in a FHIR foundation namespace. Capability never overrides permission."],
        ["Tool and knowledge visibility", "The agent uses only the tools and knowledge the AI Admin enabled. It cannot use anything excluded, even if a user asks for it."],
        ["Irreversible actions", "Deleting an endpoint with its data, resetting data, disabling an endpoint, or weakening validation cannot be undone. The agent states the consequence plainly and proceeds only on the user's clear go-ahead for that exact action."],
        ["Stay on scope", "The agent stays strictly on FHIR platform work. It never introduces productions, business hosts, HL7 v2 routing, or any topic the user did not raise."],
        ["Integrate with CI/CD (hard requirement)", "Everything the agent creates is compatible with the same change control pipeline the platform already uses. The agent only operates in DEV or LOCAL, not TEST or PROD, for the MVP."],
    ],
})

B("h2", "1.6 LLM connection")
B("p", "The agent talks to an AI model. The AI Admin sets up the connection. See Epic 3, stories US-B01 to US-B03.")
B("bullets", [
    "Choose the provider and the model.",
    "Store the credential securely. Never in plain text.",
    "Test the connection. Show available or unavailable.",
])

B("pagebreak")

# ================================================================ EPIC 2 (CHAT)
B("h1", "Epic 2. Chat Experience")
B("p", "Scope: MVP. Customer facing. This is the FHIR Engineer's main surface, embedded in the FHIR Server Management portal. Build it second.")

B("h2", "2.1 Building by conversation")
B("story", {"id": "US-A01", "title": "Build the FHIR platform by describing it",
    "persona": FE, "need": "describe what I want in plain language and have the agent plan and do it",
    "reason": "I know FHIR, but I should not have to learn InterSystems internals or click through every screen",
    "ac": [
        "I state the goal in FHIR terms (“create an R4 server at this path”, “load these files”, “export this Group”). I get back a clear plan of what will be created or changed.",
        "The agent recommends sensible defaults for anything I left open (storage strategy, profiles, where export files land) and states them before acting.",
        "The agent shows the plan and waits for my approval before it creates anything.",
        "After I approve, the agent acts and reports, in plain language, what now exists and where.",
    ]})
B("story", {"id": "US-A02", "title": "Approve or reject every change",
    "persona": FE, "need": "approve or reject each change before it is applied",
    "reason": "I am accountable for what gets created in my FHIR platform, and I must stay in control",
    "ac": [
        "Every change pauses for my approval. This covers create, configure, load, reset, delete, activate, and start.",
        "The agent shows what will happen before it happens, including irreversible actions called out plainly.",
        "Nothing is applied until I approve.",
        "If I reject, the agent acknowledges it and asks how I want to proceed.",
    ]})
B("story", {"id": "US-A03", "title": "See what the agent is doing",
    "persona": FE, "need": "see, in real time, what the agent is doing and the result of each step",
    "reason": "I need to trust and follow its work, not watch a spinner",
    "ac": [
        "The response appears as it is produced. I never wait at a blank screen.",
        "Each action is shown in plain form. I see what it did and the outcome. I can expand it for detail.",
        "Errors are explained in plain language. I do not see technical traces unless I ask for them.",
    ]})
B("story", {"id": "US-A04", "title": "Remember the conversation (context memory)",
    "persona": FE, "need": "rely on the agent remembering the context of our conversation",
    "reason": "building a FHIR platform is step by step, and I must refer back without repeating myself",
    "ac": [
        "The agent keeps the full context of the current conversation. It builds on earlier turns.",
        "I can say “now load last month's data into that server” and the agent knows which server I mean.",
        "I can leave and come back to the same conversation. The context is still there.",
        "I can have multiple conversations open at once, with contexts isolated from one another.",
    ]})
B("story", {"id": "US-A05", "title": "Stay responsive, never hang",
    "persona": FE, "need": "count on the agent staying responsive at all times",
    "reason": "loading data or running an export can take a while, and I need to know where I stand",
    "ac": [
        "Long tasks such as a directory load or a bulk export are started asynchronously. The agent reports progress and lets me check status, rather than blocking.",
        "The agent never looks frozen. If a step runs long, it summarizes progress and offers to continue, check status, or retry.",
    ]})

B("h2", "2.2 Working from uploaded data and specifications")
B("story", {"id": "US-A06", "title": "Upload FHIR data and specifications",
    "persona": FE, "need": "attach FHIR data files and specification documents to the conversation and have the agent work from them",
    "reason": "retyping data or a multi-page specification into a chat box is slow and error prone, and I have no file system access to the server",
    "ac": [
        "I upload FHIR resource files through the Load FHIR Data helper. They are staged into a server-readable folder so the agent can load them.",
        "I upload common document and spreadsheet files (PDF, Word, Markdown, spreadsheet) into the chat.",
        "The agent reads them and produces a clear, structured summary of what it understands. It does not act until I confirm the summary. I can edit it, and the agent revises it.",
    ]})
B("story", {"id": "US-A07", "title": "Have the agent understand our specification",
    "persona": FE, "need": "rely on the agent reading our organization's own FHIR data-sharing or profiling specification correctly",
    "reason": "a generic reading will misread our conventions and produce the wrong projection or export",
    "ac": [
        "The AI Admin provides organization specific knowledge. See Epic 3. The agent then reads our specification the way we intend.",
        "The summary reflects our specification's real meaning. It is not a generic guess.",
        "This is the hand off between the two people. The AI Admin authors the knowledge. The engineer benefits from it on upload.",
    ]})

B("h2", "2.3 Refining the result, by prompt or in the graphical tools")
B("p", "This behavior defines the product for a non-developer. The engineer builds with the agent, then keeps working in the FHIR portal's graphical tools, and comes back to the agent when useful. There is no code at any point.")
B("story", {"id": "US-A08", "title": "Refine by prompting",
    "persona": FE, "need": "ask the agent to change something it just did",
    "reason": "for most changes, talking is faster than clicking through screens",
    "ac": [
        "The agent finds the existing endpoint, configuration, or export. It proposes the change, applies it after I approve, and re-checks the result.",
    ]})
B("story", {"id": "US-A09", "title": "Open any result in the FHIR portal graphical tools",
    "persona": FE, "need": "open what the agent created directly in the matching FHIR Server Management screen",
    "reason": "for detailed work I prefer the visual tools, and I must never read or edit source code",
    "ac": [
        "The agent's report links me to the right screen: endpoint configuration, the capability statement, the resource repository, profile and package management, and the FHIR SQL Builder wizard.",
        "The link opens the exact item, not a generic landing page.",
        "I never see or edit source code or internal class names.",
    ]})
B("story", {"id": "US-A10", "title": "Edit in a graphical tool, then keep talking to the agent",
    "persona": FE, "need": "change something in a FHIR portal screen and then keep working with the agent using my current state",
    "reason": "I move between chat and the visual tools, and the agent must never use a stale copy",
    "ac": [
        "I change an endpoint or a projection in the portal. I then ask the agent about it. The agent reflects my current state.",
        "The agent can summarize, validate, or keep building on my changes.",
    ]})
B("story", {"id": "US-A11", "title": "Inspect the server through the existing tools",
    "persona": FE, "need": "jump from the chat to the capability statement, the resource repository, and the server audit for the server I built",
    "reason": "checking what is really in the server is a visual, click-through task",
    "ac": [
        "The agent links me to the capability statement, to the resource repository for a given type, and to the FHIR Server Audit view for storage, counts, and ingestion.",
        "The agent can also answer “how many Patients did I load” and offer the visual link so I can confirm.",
    ]})

B("h2", "2.4 Testing and validating from the chat")
B("story", {"id": "US-A12", "title": "Validate resources and confirm the load",
    "persona": FE, "need": "ask the agent to validate representative resources and confirm what loaded",
    "reason": "I must prove the data is well formed and complete before analysts or partners depend on it",
    "ac": [
        "The agent validates representative resources against the server's profiles and reports the outcome.",
        "The agent counts resources by type and reports whether the load landed as expected. It surfaces any error.",
    ]})
B("story", {"id": "US-A13", "title": "Confirm the export or projection against the requirement",
    "persona": FE, "need": "have the agent confirm the outbound export or the SQL projection matches what the requirement asked for",
    "reason": "the partner rejects anything that does not match, and analysts need the right columns, so I must catch that now",
    "ac": [
        "For a bulk export, the agent confirms the export completed and reports what was written and where.",
        "For a SQL projection, the agent confirms the expected resources and fields are queryable, and flags anything missing.",
    ]})

B("h2", "2.5 Managing conversations")
B("story", {"id": "US-A14", "title": "Keep, search, resume, and label conversations",
    "persona": FE, "need": "review, search, resume, and rename my past conversations",
    "reason": "standing up a FHIR platform spans several sessions, and I need continuity",
    "ac": [
        "A history view lists past conversations. I can search, resume, rename, and delete them.",
        "When I resume a past conversation, the agent reviews the context before engaging.",
        "A clear indicator shows whether the agent is available.",
    ]})
B("story", {"id": "US-A15", "title": "Start from a guided prompt",
    "persona": FE, "need": "start from guided example prompts",
    "reason": "a blank box is intimidating, and good examples show me what is possible",
    "ac": [
        "Example prompts appear on a new conversation. They fill the composer when I pick one.",
        "At least one example shows an end-to-end flow like create a server, load data, and export it.",
    ]})

B("pagebreak")

# =========================================================== EPIC 3 (AI SETTING)
B("h1", "Epic 3. AI Setting Experience")
B("p", "Scope: MVP. Built and operated by InterSystems. Not shown to customers in the MVP. The AI Admin assembles the agent here. Build it third. It does not need to be fully featured or customer grade for the MVP. A later release can open it to customers.")

B("h2", "3.1 Connecting the product to an AI model")
B("story", {"id": "US-B01", "title": "Connect to the chosen AI provider",
    "persona": AD, "need": "connect the product to the chosen AI model provider in the interface",
    "reason": "the organization decides which model powers the agent, and I must set it up without code or config files",
    "ac": [
        "I choose a provider and a model in a form. The choice is saved as configuration.",
        "The organization controls which provider is used. It therefore controls where its data goes.",
    ]})
B("story", {"id": "US-B02", "title": "Store credentials securely",
    "persona": AD, "need": "enter the provider credential in a masked field and have it stored securely",
    "reason": "secrets must never be visible, exported, or stored in plain text, and I am accountable for that",
    "ac": [
        "The credential is masked when I enter it. It is stored in an encrypted secret store.",
        "It is never shown, returned, or written to any log or record after I save it.",
    ]})
B("story", {"id": "US-B03", "title": "Test the connection and see its status",
    "persona": AD, "need": "test the connection and see a clear available or unavailable status",
    "reason": "I must know the model is reachable before engineers depend on it",
    "ac": [
        "A clear status shows available, unavailable, or not yet tested.",
        "On failure, the interface shows the provider's exact error message.",
        "The chat shows this status to end users. If the model is down, engineers are told plainly.",
    ]})

B("h2", "3.2 Assembling the agent")
B("story", {"id": "US-B04", "title": "Choose which tools and skills the agent has",
    "persona": AD, "need": "choose exactly which tools and which skills the FHIR Assistant can use",
    "reason": "the agent's skill and its safety both depend on the right scope and nothing more",
    "ac": [
        "I see the full pool of tools and skills. I include or exclude each one per agent.",
        "An excluded tool cannot be used, even if a user asks for it.",
        "Changes take effect for new conversations. No redeployment is needed.",
    ]})
B("story", {"id": "US-B05", "title": "Tune the agent's instructions",
    "persona": AD, "need": "write and tune the agent's instructions in the interface",
    "reason": "instruction quality is the biggest lever on output quality, and I must tune it without a developer",
    "ac": [
        "I edit the agent's instructions, persona, planning rules, and guardrails.",
        "I save my changes. I can revert to the default.",
        "My changes persist across product updates.",
    ]})
B("story", {"id": "US-B06", "title": "Group tools and toggle whole areas",
    "persona": AD, "need": "group related tools and turn whole areas on or off for the agent",
    "reason": "I want to enable or disable a whole capability area with one switch",
    "ac": [
        "I enable or disable each tool area (FHIR server, bulk export, catalog). I can give it a description.",
    ]})

B("h2", "3.3 Authoring skills")
B("story", {"id": "US-B07", "title": "Author a skill from a document",
    "persona": AD, "need": "create a skill by writing or uploading a plain language document",
    "reason": "my expertise lives in documents, and I should turn it into agent knowledge without code",
    "ac": [
        "I create a skill with a name, a description, and plain language content. I paste it or upload it.",
        "The skill becomes available to attach to any agent.",
        "No coding is required.",
    ]})
B("story", {"id": "US-B08", "title": "Register a supplied skill",
    "persona": AD, "need": "register skills supplied by the vendor or a developer, in the same interface",
    "reason": "some skills come pre-built, and I must make them available without rebuilding them",
    "ac": [
        "Self-authored skills and supplied skills appear in one pool.",
        "I attach both kinds the same way.",
    ]})
B("story", {"id": "US-B09", "title": "Teach the agent our data-sharing specification",
    "persona": AD, "need": "author a skill that explains how to read our own FHIR data-sharing or profiling specification, and attach it to the agent",
    "reason": "when engineers upload that specification, the agent must read our conventions correctly to produce the right projection or export",
    "ac": [
        "I write the skill in plain language. It describes our specification's structure and conventions.",
        "I attach it to the engineers' agent.",
        "After I attach it, uploaded specifications are read our way.",
        "This is the worked example of a skill that customers build for themselves.",
    ]})

B("h2", "3.4 Managing the catalogs")
B("story", {"id": "US-B10", "title": "Provide and refresh the catalogs",
    "persona": AD, "need": "provide and refresh the two catalogs the agent uses, the FHIR and platform classes and the transformation and mapping patterns, and choose which catalogs each agent can use",
    "reason": "the agent can only recommend the right part if the catalogs are present, current, and relevant",
    "ac": [
        "I see each catalog's size and last update time. I refresh it on demand.",
        "I run a test search to check quality before I expose a catalog.",
        "I bind specific catalogs to specific agents.",
    ]})

B("h2", "3.5 Oversight")
B("story", {"id": "US-B11", "title": "Review the audit trail",
    "persona": AD, "need": "review a complete record of every action the agent took for a user",
    "reason": "I am accountable for what was done, by whom, and where, for compliance and troubleshooting",
    "ac": [
        "The record shows who, when, which environment, what action, and the result.",
        "I can filter to errors only.",
        "The record captures successful actions and denied attempts.",
    ]})
B("story", {"id": "US-B12", "title": "Verify a tool before exposing it",
    "persona": AD, "need": "inspect and safely try a tool before I expose it to engineers",
    "reason": "I want to confirm it behaves as intended",
    "ac": [
        "I browse each tool with its description and inputs.",
        "I safely try a non-destructive tool (for example, count resources) with sample input.",
    ]})

B("pagebreak")

# ================================================================ WALKTHROUGH
B("h1", "4. End to end walkthrough")
B("p", "This is the acceptance narrative. It shows both people at work.")
B("lead", "AI Admin. One time set up in the AI Setting Experience. Internal to InterSystems.")
B("numbers", {"start": 1, "items": [
    "Connects the product to the chosen AI model. Stores the credential securely. Tests the connection until it shows available.",
    "Provides and refreshes the catalogs of FHIR and platform classes and transformation patterns. Confirms relevance with a test search.",
    "Authors the organization specific data-sharing specification skill. Attaches it to the engineers' agent.",
    "Assembles the FHIR Assistant. Selects the tools and skills it needs. Tunes its instructions. Binds the catalogs.",
]})
B("lead", "FHIR Engineer. The daily work in the chat, inside the FHIR Server Management portal. Customer facing.")
B("numbers", {"start": 5, "items": [
    "Opens the FHIR Assistant from the portal header. Says “create an R4 server at /acme/fhir with Advanced JSON and US Core”.",
    "The agent confirms a suitable FHIR foundation namespace and that the path is free, presents the plan, and waits. The engineer approves. The agent creates the server and reports the endpoint.",
    "The engineer uploads a batch of FHIR files through Load FHIR Data, then says “load these”. The agent stages, loads asynchronously, monitors to completion, and reports the counts by type.",
    "The engineer asks the agent to validate a few resources and confirm the load. The agent validates and reports.",
    "The engineer says “I want analysts to query Patients and Observations in SQL”. The agent walks them through the FHIR SQL Builder wizard: Analysis, Specification, Projection, and confirms the tables are queryable.",
    "The engineer says “now export the diabetic cohort to our partner”. The agent proposes a Bulk FHIR export with SMART backend OAuth, provisions the prerequisites after approval, runs the export, and confirms it completed.",
    "The engineer opens the endpoint configuration and the FHIR Server Audit in the portal to confirm storage, counts, and ingestion are healthy.",
]})
B("lead", "AI Admin. Oversight.")
B("numbers", {"start": 12, "items": [
    "Reviews the audit trail. Sees every action, by this engineer, in this environment.",
]})

# ================================================================ DoD
B("h1", "5. Definition of Done (MVP)")
B("p", "The MVP is done when all of these are true, from inside the FHIR Server Management portal.")
B("numbers", {"start": 1, "items": [
    "The agent follows the safe loop: research, propose, wait for approval, act, test, report.",
    "From the chat, the agent creates a FHIR R4 server in a FHIR foundation namespace, with the storage strategy and profiles the engineer approved.",
    "The agent loads a batch of FHIR files (staged server-side through Load FHIR Data), monitors ingestion to completion, and reports the counts by type.",
    "The agent validates representative resources and confirms the load against what was expected.",
    "The agent walks the engineer through the FHIR SQL Builder to project selected resources to SQL, and confirms the tables are queryable.",
    "The agent configures and runs a secured Bulk FHIR export, provisioning the OAuth, SSL/TLS, storage, and credential prerequisites, and confirms it completed.",
    "The engineer opens any result in the matching FHIR portal graphical tool with one action and never sees or edits source code.",
    "The agent respects the calling namespace and the signed in user. It never exceeds what they allow, even when it has the capability, and it refuses in plain language when a restriction applies.",
    "Irreversible actions (reset, delete with data, disable, weaken validation) happen only on the user's explicit go-ahead for that exact action.",
    "Everything the agent creates is compatible with the platform's change control pipeline. The agent operates only in DEV or LOCAL for the MVP.",
    "The agent is connected to the chosen AI model. The credential is secured. The status shows available.",
    "An InterSystems AI Admin can assemble the agent in the AI Setting Experience: connect the LLM, build the catalogs, pick the tools and skills, set the policies, and write the prompt.",
    "Every action is recorded in the audit trail.",
    "The solution uses only what the platform provides. No CLI. No file server access for the user. No developer tools.",
]})
B("callout", "Scope reminder. The MVP has three parts. The Agent and the Chat Experience are customer facing, embedded in the FHIR Server Management portal. The AI Setting Experience is built and operated by InterSystems, and it is not shown to customers in the MVP. The AI Admin for this product is an InterSystems person. More agents, including the Integration Agentic Builder, are separate deliverables.")

# ================================================================ NFR
B("h1", "6. Non functional requirements")
B("table", {
    "headers": ["Requirement", "Target"],
    "rows": [
        ["Responsiveness", "Visible output starts within a couple of seconds. Long tasks (directory load, bulk export) run asynchronously with progress and status, and the agent stays responsive throughout."],
        ["Reliability", "The agent never fails silently. It explains every error in plain language, and never claims something was created, loaded, or run unless a tool confirmed it."],
        ["Security, acting as the user", "The agent acts within the signed in user's permissions. It can never do something the user could not do by hand."],
        ["Security, change gate", "Every change needs the user's explicit approval before it is applied. Irreversible actions are called out plainly and confirmed for the exact action."],
        ["Security, secrets", "Credentials are stored encrypted. They are never shown, exported, or logged."],
        ["Security, inherit the context", "The agent runs inside the calling namespace and the signed in user. A FHIR endpoint is created only in a FHIR foundation namespace. It refuses in plain language when a restriction applies."],
        ["No internals exposed", "The agent never shows tool names, raw tool output, JSON, or internal IRIS or HealthShare class names. It speaks in FHIR and platform terms only."],
        ["Auditability", "Every action taken for a user is recorded and can be reviewed."],
        ["Data control", "The organization chooses the AI provider. It therefore chooses where its data is processed."],
        ["Usability", "A non-developer FHIR engineer can use the product from the FHIR portal. No source code is visible in the workflow. Acronyms specific to InterSystems are expanded on first use."],
        ["Platform constraints", "The product uses only what the managed platform provides. No command line. No file server access for the user. No local developer tools. Uploaded files are staged server-side."],
        ["Change control and CI/CD (hard requirement)", "Everything the agent creates is compatible with the same change control pipeline the platform already uses. The agent operates only in DEV or LOCAL for the MVP, never TEST or PROD."],
    ],
})


# ----------------------------------------------------------------------------
# RENDER HELPERS  (identical house style to build_prd_aiadmin.py)
# ----------------------------------------------------------------------------

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
    tcPr.append(sh)

def set_cell_text(cell, text, bold=False, color=INK, white=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text); run.font.size = Pt(size); run.font.bold = bold
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    run.font.name = "Calibri"

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, HEADER_FILL); set_cell_text(c, h, bold=True, white=True)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if r_i % 2 == 1:
                shade(cells[i], ZEBRA_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13.5, NAVY), ("Heading 3", 11.5, SLATE)]:
        st = doc.styles[name]
        st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color

def num_parts(payload):
    if isinstance(payload, dict):
        return payload.get("start", 1), payload["items"]
    return 1, payload

def render_docx():
    doc = Document()
    style_doc(doc)
    for sec in doc.sections:
        sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
        sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)

    for kind, payload in BLOCKS:
        if kind == "title":
            p = doc.add_paragraph(); r = p.add_run(payload)
            r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY; r.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(2)
        elif kind == "subtitle":
            p = doc.add_paragraph(); r = p.add_run(payload)
            r.font.size = Pt(13); r.font.color.rgb = SLATE; r.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(10)
        elif kind == "meta":
            t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
            for k, v in payload:
                cells = t.add_row().cells
                shade(cells[0], ZEBRA_FILL); set_cell_text(cells[0], k, bold=True)
                set_cell_text(cells[1], v)
            doc.add_paragraph()
        elif kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "h3":
            doc.add_heading(payload, level=3)
        elif kind == "p":
            p = doc.add_paragraph(payload); p.paragraph_format.space_after = Pt(6)
        elif kind == "lead":
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(payload); r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)
        elif kind == "callout":
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
            r = p.add_run(payload); r.font.italic = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:space'), '10'); left.set(qn('w:color'), '153E6B')
            pbdr.append(left); pPr.append(pbdr)
        elif kind == "bullets":
            for item in payload:
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "numbers":
            start, items = num_parts(payload)
            for i, item in enumerate(items, start):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(f"{i}.  {item}"); r.font.size = Pt(10.5)
        elif kind == "table":
            add_table(doc, payload["headers"], payload["rows"])
        elif kind == "story":
            s = payload
            doc.add_heading(f"{s['id']}: {s['title']}", level=3)
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"As {s['persona']}, I need to be able to {s['need']}, because {s['reason']}.")
            r.font.italic = True; r.font.size = Pt(10.5)
            lbl = doc.add_paragraph(); lbl.paragraph_format.space_after = Pt(0)
            lr = lbl.add_run("Acceptance criteria"); lr.font.bold = True; lr.font.size = Pt(10)
            for item in s["ac"]:
                doc.add_paragraph(item, style="List Bullet")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "pagebreak":
            doc.add_page_break()

    doc.save(DOCX_OUT)


def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in rows:
        out.append("| " + " | ".join(c.replace("\n", " ") for c in row) + " |")
    return "\n".join(out)

def render_md():
    L = []
    for kind, payload in BLOCKS:
        if kind == "title": L.append(f"# {payload}\n")
        elif kind == "subtitle": L.append(f"> {payload}\n")
        elif kind == "meta":
            for k, v in payload: L.append(f"- **{k}:** {v}")
            L.append("")
        elif kind == "h1": L.append(f"\n## {payload}\n")
        elif kind == "h2": L.append(f"\n### {payload}\n")
        elif kind == "h3": L.append(f"\n#### {payload}\n")
        elif kind == "p": L.append(payload + "\n")
        elif kind == "lead": L.append(f"**{payload}**\n")
        elif kind == "callout": L.append(f"> {payload}\n")
        elif kind == "bullets":
            for item in payload: L.append(f"- {item}")
            L.append("")
        elif kind == "numbers":
            start, items = num_parts(payload)
            for i, item in enumerate(items, start): L.append(f"{i}. {item}")
            L.append("")
        elif kind == "table":
            L.append(md_table(payload["headers"], payload["rows"]) + "\n")
        elif kind == "story":
            s = payload
            L.append(f"\n#### {s['id']}: {s['title']}\n")
            L.append(f"*As {s['persona']}, I need to be able to {s['need']}, because {s['reason']}.*\n")
            L.append("Acceptance criteria:")
            for item in s["ac"]: L.append(f"- {item}")
            L.append("")
        elif kind == "pagebreak": L.append("\n---\n")
    with open(MD_OUT, "w") as f:
        f.write("\n".join(L).strip() + "\n")


if __name__ == "__main__":
    render_docx()
    render_md()
    print("Wrote:")
    print(" ", DOCX_OUT)
    print(" ", MD_OUT)
