"""Build the three FHIR Assistant agent .docx documents.

Outputs (1.1.0):
  docs/01_Requirements_User_Stories_FHIRAssistant_v1.1.0.docx
  docs/02_Technical_Build_Specification_FHIRAssistant_v1.1.0.docx
  docs/03_Lessons_Learned_FHIRAssistant_v1.1.0.docx

The FHIR Assistant agent (AgenticInterop.Agent.FHIRSpecialist) is the
second shipped agent — focused exclusively on the FHIR R4 platform:
FHIR Server (repository), Bulk FHIR Coordinator (BFC) exports, and
FHIR SQL Builder. It binds three MCP servers (FHIRServer, BulkFHIR,
Catalog) and five skills (FHIRServer, FHIRR4, SDA, BulkFHIR,
FHIRSQLBuilder). The chatbot surface is /csp/fhir-management (the
shipped FHIR Server Management page, instrumented with a launcher
button via AgenticInterop.Install.FHIRManagementPatch).

For the Health Interop generalist agent, see build_all_docx.py.

Requires: pip install python-docx Pillow
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

VERSION = "v1.1.0"
AGENT_SUFFIX = "FHIRAssistant"

DOCS = os.path.dirname(__file__)
IMG = os.path.join(DOCS, "img")


def img(name):
    return os.path.join(IMG, name)


def style_doc(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def add_title_page(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    doc.add_paragraph()
    doc.add_paragraph()


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph()


def add_screenshot(doc, name, caption, width_inches=6.0):
    # Allow callers to omit ".png" — append if missing.
    if not name.lower().endswith(".png"):
        name = name + ".png"
    path = img(name)
    if not os.path.exists(path):
        print(f"  WARN: screenshot missing — {name} (expected at {path})")
        return
    if True:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_inches))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)


# =============================================================================
# DOC 1 — Requirements and User Stories (FHIR Assistant)
# =============================================================================

def build_doc1():
    doc = Document()
    style_doc(doc)
    add_title_page(doc,
        "Agentic Health Interoperability",
        f"Requirements and User Stories — FHIR Assistant Agent — {VERSION}")
    doc.add_paragraph(
        "This document covers the FHIR Assistant agent (AgenticInterop.Agent.FHIRSpecialist) — "
        "a senior FHIR R4 specialist that helps platform Admins and End Users build, run, and "
        "operate an InterSystems IRIS for Health FHIR platform. Scope is strictly the FHIR side: "
        "FHIR R4 server (repository), Bulk FHIR Coordinator (BFC), and FHIR SQL Builder. The "
        "Health Interop generalist agent (productions, DTL/BPL/HL7 v2) is documented separately "
        "in 01_Requirements_User_Stories_HealthInterop_v1.1.0.docx."
    )

    doc.add_heading("Table of Contents", level=1)
    for line in [
        "1. Introduction",
        "    1.1 Product Vision",
        "    1.2 Why a Separate FHIR Assistant",
        "    1.3 Plan-Authorize-Act Flow",
        "2. Personas",
        "    2.1 AI Hub Admin (FHIR platform)",
        "    2.2 End User (FHIR consumer / data steward)",
        "3. Functional Requirements",
        "    3.1 FHIR Server Lifecycle",
        "    3.2 FHIR Resource Operations",
        "    3.3 Bulk FHIR Coordinator",
        "    3.4 FHIR SQL Builder",
        "    3.5 SDA / FHIR Transformation",
        "4. User Stories — AI Hub Admin",
        "5. User Stories — End User",
        "6. Non-functional Requirements",
        "7. Out of Scope (handled by Health Interop)",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    # ---- 1. Introduction ----
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Product Vision", level=2)
    doc.add_paragraph(
        "InterSystems IRIS for Health ships a FHIR R4 server, a Bulk FHIR Coordinator (BFC), a "
        "SDA-FHIR transformation pipeline, and a FHIR SQL Builder. These capabilities are powerful "
        "but require multi-step configuration that experienced IRIS engineers do quickly and "
        "newcomers do slowly: locating the correct foundation namespace, picking a storage strategy, "
        "loading profile packages, wiring SMART OAuth, and configuring BFC fetch/storage adapters. "
        "The FHIR Assistant gives every persona — Admin or end user — a conversational interface "
        "that handles those steps through real IRIS APIs."
    )
    doc.add_paragraph(
        "The FHIR Assistant operates under a strict plan-then-authorize-then-act discipline: it "
        "uses read-only tools to understand the situation, presents a plain-language plan, and only "
        "acts after the user explicitly approves. It refuses to expose internal class names, tool "
        "names, or JSON output; the user sees endpoint paths, namespace names, counts, and outcomes."
    )
    add_screenshot(doc, "21_fhir_chatbot", "Figure 1.1 — FHIR Assistant chatbot at /agentic/chat/index.html?chatbot=fhir-management. The agent operates under the FHIRSpecialist system prompt with plan-authorize-act discipline.")

    doc.add_heading("1.2 Why a Separate FHIR Assistant", level=2)
    doc.add_paragraph(
        "FHIR platform work has a different mental model than Ensemble production work. The user is "
        "interacting with a REST server, not a message-routing engine. Topics like business hosts, "
        "DTLs, HL7 v2 schemas, and routing rules are not relevant; conversely, topics like FHIR "
        "profiles, metadata packages, CapabilityStatement, $export sessions, SMART backend services, "
        "and SQL projections are central. A dedicated agent with its own MCP and Skill bindings keeps "
        "each conversation focused and avoids leaking irrelevant suggestions across domains."
    )
    doc.add_paragraph(
        "Personas reflect this split. The FHIR Assistant serves two personas — AI Hub Admin (sets up "
        "and operates the FHIR platform) and End User (queries data, validates resources, monitors "
        "exports). The Interface Engineer and Operator personas from the Health Interop world are out "
        "of scope here."
    )

    doc.add_heading("1.3 Plan-Authorize-Act Flow", level=2)
    doc.add_paragraph(
        "For any operation that creates, configures, loads, resets, deletes, enables, disables, or "
        "provisions state, the FHIR Assistant follows a three-step discipline:"
    )
    doc.add_paragraph("1. Plan quietly — uses read-only tools to discover namespaces, endpoints, packages, "
                      "counts, and existing config. Never asks permission just to look.",
                      style="List Number")
    doc.add_paragraph("2. Summarize and ask — replies with a short plain plan of exactly what will happen "
                      "and the choices it rests on (target namespace, endpoint path, FHIR version, storage "
                      "strategy). Ends with a clear question. Stops. Does not act.",
                      style="List Number")
    doc.add_paragraph("3. Act only after authorization — performs the action and reports the outcome in one "
                      "or two short sentences. No restating, no JSON, no architecture explanation.",
                      style="List Number")
    doc.add_paragraph(
        "Read-only questions (\"what endpoints exist?\", \"how big is the server?\") are answered directly "
        "without a plan or authorization step."
    )

    # ---- 2. Personas ----
    doc.add_page_break()
    doc.add_heading("2. Personas", level=1)
    doc.add_paragraph(
        "Two personas use the FHIR Assistant. The Developer and Interface Engineer / Operator personas "
        "from the Health Interop documentation are out of scope here."
    )

    doc.add_heading("2.1 AI Hub Admin (FHIR platform)", level=2)
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Primary interface", "FHIR Server Management page (/csp/fhir-management) with the injected FHIR Assistant launcher button; also reachable from /agentic/chat/index.html bound to the fhir-management chatbot key"],
            ["Security scope", "IRIS roles required for FHIR endpoint and BFC administration: %HS_Administrator (foundation namespace, endpoint create/delete), %HS_BFC_Administrator (BFC create/start), %Admin_Secure (SSL/OAuth provisioning)"],
            ["Primary outcomes", "FHIR R4 endpoints configured with the right profile packages; SMART backend OAuth servers + clients provisioned; BFC configurations created and activated; data loaded; storage/perf audited; access controls in place"],
            ["Not in scope", "Productions, DTLs, BPLs, routing rules, HL7 v2 — see the Health Interop agent"],
        ])

    doc.add_heading("2.2 End User (FHIR consumer / data steward)", level=2)
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Primary interface", "FHIR Server Management page or the standalone chatbot at /agentic/chat/index.html?chatbot=fhir-management"],
            ["Security scope", "Read access to one or more FHIR foundation namespaces; ability to call resource search / read / $validate; for staging data, write access to the FHIR upload staging folder"],
            ["Primary outcomes", "Locate the right FHIR endpoint, search and read resources, validate inbound FHIR payloads against the active profile package, kick off and monitor a bulk export, project FHIR data into SQL via the FHIR SQL Builder"],
            ["Not in scope", "Endpoint provisioning, OAuth setup, BFC config creation — those are Admin tasks"],
        ])

    # ---- 3. Functional Requirements ----
    doc.add_page_break()
    doc.add_heading("3. Functional Requirements", level=1)
    doc.add_paragraph(
        "Each functional area below maps to a set of tools the FHIR Assistant exposes via its three "
        "MCP servers. Section 3 of the Technical Build Specification has the full tool-by-tool catalog; "
        "this section describes the user-visible capability."
    )

    doc.add_heading("3.1 FHIR Server Lifecycle", level=2)
    doc.add_paragraph(
        "The agent can: discover which IRIS namespaces are FHIR foundation namespaces (have FHIR "
        "metadata defined), list and inspect existing endpoints, read each endpoint's CapabilityStatement, "
        "create new endpoints in the right foundation namespace with a chosen FHIR version and storage "
        "strategy, manage which profile packages (hl7.fhir.r4.core, us-core, etc.) are loaded, update "
        "endpoint configuration, enable or disable endpoints, and delete endpoints together with their data."
    )
    doc.add_paragraph(
        "Storage-strategy support: Advanced JSON, JsonAdvSQL, and any other strategies the running IRIS "
        "version exposes. The agent maps the user's plain-language choice (\"use Advanced JSON\") to the "
        "internal strategy name."
    )

    doc.add_heading("3.2 FHIR Resource Operations", level=2)
    doc.add_paragraph(
        "Searches, reads, creates, updates, deletes, and validates FHIR resources against any configured "
        "endpoint. Counts resources by type (Patient, Encounter, Observation, etc.). Loads a server-side "
        "directory of FHIR JSON files with progress reporting, ingestion metrics, and a durable run-history "
        "ledger. Resets data when explicitly authorized. Audits storage size and ingestion throughput, and "
        "runs a query-performance probe against representative resource types."
    )

    doc.add_heading("3.3 Bulk FHIR Coordinator (BFC)", level=2)
    doc.add_paragraph(
        "Lists existing BFC configurations and their fetch/storage adapters. Creates new configurations end "
        "to end: storage directory, SSL/TLS configuration, interop credential, and (for SMART-backed servers) "
        "an OAuth 2.0 server + client. Activates configurations, starts $export sessions against a source "
        "FHIR endpoint, and monitors session progress. Supports two storage modes: ndjson files on disk, "
        "or ingestion into a target FHIR server."
    )

    doc.add_heading("3.4 FHIR SQL Builder", level=2)
    doc.add_paragraph(
        "The FHIR SQL Builder has no automation interface. The agent walks the user through it manually "
        "using the FHIRSQLBuilder skill: Analysis (pick resources and elements), Specification (define "
        "columns, subtables, filters), Projection (generate and run). The agent answers SQL/JDBC/ODBC "
        "questions and explains how to read the projected tables."
    )

    doc.add_heading("3.5 SDA / FHIR Transformation", level=2)
    doc.add_paragraph(
        "The FHIR Assistant uses the SDA and FHIRR4 skills to answer questions about the SDA hub pattern "
        "and the SDA↔FHIR DTL library. It does not author HL7-to-SDA pipelines (Health Interop scope); it "
        "helps the user understand which built-in SDA-FHIR DTLs and lookup tables drive the FHIR-side of a "
        "transformation pipeline already in place."
    )

    # ---- 4. User Stories — Admin ----
    doc.add_page_break()
    doc.add_heading("4. User Stories — AI Hub Admin", level=1)

    admin_stories = [
        ("UC1 — Stand up a new FHIR R4 endpoint",
         "As an AI Hub Admin, I want to create a FHIR R4 endpoint at a path I choose, in the correct "
         "foundation namespace, with the storage strategy and profile packages my project requires, so "
         "that integration partners can start reading and writing FHIR resources.",
         [
             "The agent confirms which IRIS namespaces are FHIR foundation namespaces and proposes the right one",
             "The agent verifies the requested endpoint path is free",
             "The agent presents the plan in plain language (namespace, path, version, strategy, packages) and requests authorization",
             "On authorization, the endpoint is created and the agent reports the live URL and namespace",
             "Failure modes (path conflict, missing foundation namespace, insufficient role) are reported verbatim with a remediation hint"
         ]),
        ("UC2 — Add a profile package to an existing endpoint",
         "As an AI Hub Admin, I want to add US Core (or another profile package) to an existing FHIR endpoint "
         "so that validation matches my organization's interoperability requirements.",
         [
             "The agent lists the packages currently loaded on the endpoint",
             "The agent confirms the requested package is available in the IRIS instance's package registry",
             "On authorization, the package is added; the agent reports the new effective profile set",
         ]),
        ("UC3 — Provision SMART Backend OAuth for BFC",
         "As an AI Hub Admin, I want the FHIR Assistant to stand up the OAuth 2.0 server and a SMART backend "
         "client so that my Bulk FHIR Coordinator config can authenticate against a source endpoint.",
         [
             "The agent surfaces the four prerequisites (storage directory, SSL/TLS config, interop credential, OAuth server + client)",
             "The agent reports which are already present and which need creating",
             "On authorization, each missing prerequisite is provisioned idempotently; the agent reports the OAuth server URL and the client identifier",
             "If the logged-in user lacks the %Admin_Secure role, the agent reports the verbatim security error and stops",
         ]),
        ("UC4 — Create and activate a Bulk FHIR export configuration",
         "As an AI Hub Admin, I want to configure a BFC that pulls Patient and Observation resources from a "
         "source endpoint and writes them as ndjson files (or ingests them into a target endpoint) so that I "
         "can replicate data on a schedule.",
         [
             "The agent presents the fetch_config (source endpoint, OAuth client, resource types) and storage_config (ndjson directory OR target endpoint) in plain language",
             "On authorization, the configuration is created and the agent reports the config name and activation status",
             "The agent can start an $export session and monitor it; progress and session ID are reported plainly",
         ]),
        ("UC5 — Audit storage and ingestion performance",
         "As an AI Hub Admin, I want to know how much storage each FHIR endpoint is using and how fast the "
         "last ingestion run was, so I can decide whether to add capacity, archive, or tune.",
         [
             "The agent reports per-endpoint resource counts by type, storage size (database, journal, indices), and the last load run's resources/sec and bottleneck",
             "Run history is durable (FHIRLoadRun) so the user can compare runs over time",
         ]),
        ("UC6 — Reset endpoint data (destructive)",
         "As an AI Hub Admin, I want to wipe all resources from a non-production endpoint to start a clean "
         "load, with no ambiguity about what gets deleted.",
         [
             "The agent states the consequence plainly: \"This will delete every resource in <endpoint>; this cannot be undone.\"",
             "The agent proceeds only on the user's clear go-ahead for that specific action",
             "On completion, the agent reports the post-reset resource count (zero) and the endpoint is still configured",
         ]),
    ]
    for title, narrative, criteria in admin_stories:
        doc.add_heading(title, level=2)
        doc.add_paragraph(narrative)
        doc.add_paragraph("Acceptance criteria:")
        for c in criteria:
            doc.add_paragraph(c, style="List Bullet")

    # ---- 5. User Stories — End User ----
    doc.add_page_break()
    doc.add_heading("5. User Stories — End User", level=1)

    user_stories = [
        ("UC7 — Find the right FHIR endpoint",
         "As an End User, I want to ask \"what FHIR servers do we have and what versions do they support?\" "
         "and get an answer without knowing IRIS class names or namespace conventions.",
         [
             "The agent lists every endpoint across discovered FHIR foundation namespaces with URL, FHIR version, and enabled status",
             "For each endpoint, the agent surfaces the CapabilityStatement summary (supported resources, interactions, search params) on request",
         ]),
        ("UC8 — Search and read a resource",
         "As an End User, I want to ask \"how many Patients are in the dev endpoint?\" or \"show me Encounter 12345\" "
         "and get the result back in plain language.",
         [
             "Resource counts by type are reported with a single number per type",
             "Resource reads return the resource in human-readable form (key fields), not raw JSON",
             "If the user is not authorized to read the resource, the agent reports the underlying FHIR error verbatim",
         ]),
        ("UC9 — Validate a FHIR payload",
         "As an End User, I want to paste a FHIR resource I received from a partner and ask whether it passes "
         "validation against my endpoint's active profile set.",
         [
             "The agent runs $validate against the endpoint and returns the pass/fail status plus any validation issues",
             "If the user specified a profile URL, validation runs against that profile; otherwise the endpoint default",
         ]),
        ("UC10 — Load a folder of FHIR files",
         "As an End User, I want to drop a folder of FHIR JSON files onto the IRIS server's upload staging area "
         "and load them into an endpoint, with progress and final ingestion metrics.",
         [
             "The agent confirms the folder path exists and is on the IRIS server (not a user's desktop)",
             "The agent presents the file count + expected resource type breakdown and requests authorization",
             "Loading runs in the background; the agent reports the job ID, then status, then final metrics (count, duration, bytes, errors)",
             "A durable run history row is written so the load shows up in the audit/runs list later",
         ]),
        ("UC11 — Run a bulk export",
         "As an End User, I want to run an existing BFC configuration to export Patient and Observation data "
         "for the last 30 days, and monitor it until it finishes.",
         [
             "The agent lists existing BFC configurations and asks which to run if more than one matches",
             "On authorization, the $export starts; the agent reports the session ID immediately",
             "The user can ask \"how's the export going?\" any time; the agent reports the session status (in-progress / complete / failed) and resource counts so far",
         ]),
        ("UC12 — Project FHIR data into SQL",
         "As an End User, I want to be walked through the FHIR SQL Builder so I can create a relational projection "
         "of Patient and Observation data for a BI tool.",
         [
             "The agent explains the three-stage flow (Analysis, Specification, Projection) in plain language using the FHIRSQLBuilder skill",
             "The agent guides the user through column / subtable / filter choices but does NOT automate the Management Portal clicks",
             "The agent answers \"how do I query the resulting tables over JDBC?\" with concrete syntax",
         ]),
    ]
    for title, narrative, criteria in user_stories:
        doc.add_heading(title, level=2)
        doc.add_paragraph(narrative)
        doc.add_paragraph("Acceptance criteria:")
        for c in criteria:
            doc.add_paragraph(c, style="List Bullet")

    # ---- 6. Non-functional Requirements ----
    doc.add_page_break()
    doc.add_heading("6. Non-functional Requirements", level=1)
    add_table(doc,
        ["Requirement", "Target", "Rationale"],
        [
            ["Plan-Authorize-Act discipline", "Never act before explicit user OK on a mutating call", "Per the FHIR Assistant system prompt — having all details is not authorization"],
            ["No internals exposure", "Never show tool names, raw tool output, JSON, or class names", "Translate to user language: \"Advanced JSON storage strategy\", not the class name"],
            ["Strict scope", "Never introduce productions, business hosts, HL7 v2, or message routing topics the user did not raise", "Cross-domain leakage is a known failure mode; the FHIR Assistant stays in its lane"],
            ["Honest reporting", "Never claim that something was created, written, loaded, or run unless a tool actually confirmed it", "Same anti-fabrication rule as the Health Interop agent"],
            ["Style", "Plain prose, no markdown bold / italics / emojis, expand SDA/BFC on first use, common healthcare acronyms (HL7, FHIR, ADT) bare", "Chat surface renders markdown literally"],
            ["Namespace safety", "Every FHIR tool accepts an optional `namespace` argument and resolves to the request's X-IRIS-Namespace if absent", "Same pattern as Tool.FHIRServer; no cross-namespace leakage"],
            ["Security delegation", "Provisioning tools surface the underlying IRIS security error verbatim when the user lacks the required role", "%Admin_Secure / %HS_BFC_Administrator / %HS_Administrator are enforced by IRIS, not by the agent"],
        ])

    # ---- 7. Out of Scope ----
    doc.add_page_break()
    doc.add_heading("7. Out of Scope (handled by Health Interop)", level=1)
    doc.add_paragraph(
        "The FHIR Assistant deliberately does NOT cover these topics — they are the Health Interop agent's "
        "scope and are documented in 01_Requirements_User_Stories_HealthInterop_v1.1.0.docx:"
    )
    for line in [
        "Production class CRUD, business host lifecycle, routing rules, post-build validation",
        "DTL / BPL authoring and HL7 v2 → SDA pipeline construction (BuildHL7ToFHIRBPL et al.)",
        "HL7 v2 messaging — MLLP framing, ACK modes, batch handling, segment search tables",
        "Vector catalog rebuild (Ens.* / HS.*) — the Catalog skill is shared but the rebuild workflow is Admin-of-Health-Interop work",
        "Production deployment and System Default Settings dev→test→live promotion",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(
        "If a user asks the FHIR Assistant about one of those topics, the agent will direct them to the "
        "Health Interop chatbot."
    )

    path = os.path.join(DOCS, f"01_Requirements_User_Stories_{AGENT_SUFFIX}_{VERSION}.docx")
    doc.save(path)
    print(f"  saved {path}")


# =============================================================================
# DOC 2 — Technical Build Specification (FHIR Assistant)
# =============================================================================

def build_doc2():
    doc = Document()
    style_doc(doc)
    add_title_page(doc,
        "Agentic Health Interoperability",
        f"Technical Build Specification — FHIR Assistant Agent — {VERSION}")
    doc.add_paragraph(
        "This document specifies the FHIR Assistant agent's implementation: the agent class, the three "
        "MCP servers it binds, the 46 FHIR-related tools it exposes, the five skills it consults, the "
        "chatbot binding, and the security model. For the Health Interop generalist agent, see "
        "02_Technical_Build_Specification_HealthInterop_v1.1.0.docx."
    )

    doc.add_heading("Table of Contents", level=1)
    for line in [
        "1. Executive Summary",
        "2. Agent Class — FHIRSpecialist",
        "3. MCP Servers",
        "    3.1 MCP.FHIRServer",
        "    3.2 MCP.BulkFHIR",
        "    3.3 MCP.Catalog",
        "4. Tools (46 total)",
        "    4.1 Tool.FHIRServer (26 methods)",
        "    4.2 Tool.BulkFHIR (13 methods)",
        "    4.3 Tool.Catalog (7 methods)",
        "5. Skills",
        "    5.1 Skill.FHIRServer",
        "    5.2 Skill.FHIRR4",
        "    5.3 Skill.SDA",
        "    5.4 Skill.BulkFHIR",
        "    5.5 Skill.FHIRSQLBuilder",
        "6. Chatbot Surface",
        "7. Security Model",
        "8. Data Model",
        "9. Build Status Summary",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    # ---- 1. Executive Summary ----
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The FHIR Assistant is the second of two shipped agents in agentic_interop 1.1.0. It binds three "
        "MCP servers (FHIRServer, BulkFHIR, Catalog), exposes 46 FHIR-specific tools, and consults five "
        "skills (FHIRServer, FHIRR4, SDA, BulkFHIR, FHIRSQLBuilder). It is invoked from the standalone "
        "chatbot UI under the chatbot key fhir-management and also from a launcher button injected into "
        "the shipped IRIS FHIR Server Management page (/csp/fhir-management) by "
        "AgenticInterop.Install.FHIRManagementPatch."
    )
    add_table(doc,
        ["Component", "Class / Identifier", "Status"],
        [
            ["Agent", "AgenticInterop.Agent.FHIRSpecialist", "Shipped"],
            ["MCP Servers", "FHIRServer, BulkFHIR, Catalog", "Shipped"],
            ["Tools", "46 (26 + 13 + 7)", "Shipped"],
            ["Skills", "FHIRServer, FHIRR4, SDA, BulkFHIR, FHIRSQLBuilder", "Shipped"],
            ["Chatbot key", "fhir-management", "Seeded by Data.Chatbot.EnsureSeed"],
            ["Host page", "/csp/fhir-management (launcher button injected)", "Patched by Install.FHIRManagementPatch"],
            ["Plan-Authorize-Act flow", "Enforced by system prompt", "Documented in XData INSTRUCTIONS"],
        ])

    # ---- 2. Agent Class ----
    doc.add_page_break()
    doc.add_heading("2. Agent Class — FHIRSpecialist", level=1)
    doc.add_paragraph(
        "AgenticInterop.Agent.FHIRSpecialist extends %AI.Agent. The class parameters declare the three "
        "MCP bindings and the five skill bindings; the XData INSTRUCTIONS block contains the system prompt."
    )
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["MAXITERATIONS", "25"],
            ["MCPS", "AgenticInterop.MCP.FHIRServer, AgenticInterop.MCP.BulkFHIR, AgenticInterop.MCP.Catalog"],
            ["SKILLS", "AgenticInterop.Skill.FHIRServer, AgenticInterop.Skill.FHIRR4, AgenticInterop.Skill.SDA, AgenticInterop.Skill.BulkFHIR, AgenticInterop.Skill.FHIRSQLBuilder"],
        ])
    doc.add_paragraph(
        "The system prompt encodes the plan-authorize-act discipline, the absolute rules (no internals "
        "exposure, no scope drift, no claims without tool confirmation), and the FHIR-platform style "
        "guidance (plain prose, SDA/BFC acronym expansion, yes/no answer in first sentence). The full "
        "prompt lives in src/cls/AgenticInterop/Agent/FHIRSpecialist.cls."
    )
    add_screenshot(doc, "23_admin_agent_fhir", "Figure 2.1 — Admin UI Agents tab with FHIRSpecialist selected. The system prompt, MCP bindings, skill bindings, and provider can be edited without redeploying source.")

    # ---- 3. MCP Servers ----
    doc.add_page_break()
    doc.add_heading("3. MCP Servers", level=1)

    doc.add_heading("3.1 MCP.FHIRServer", level=2)
    doc.add_paragraph(
        "AgenticInterop.MCP.FHIRServer extends AgenticInterop.MCP.Base ← %AI.MCP.Service. Groups the 26 "
        "Tool.FHIRServer methods for the admin UI's MCP-binding configuration. Tool classes are "
        "registered directly via UseToolSet at agent build time — the MCP layer is admin-UI metadata."
    )

    doc.add_heading("3.2 MCP.BulkFHIR", level=2)
    doc.add_paragraph(
        "AgenticInterop.MCP.BulkFHIR groups the 13 Tool.BulkFHIR methods that wrap "
        "HS.BulkFHIR.Installer / Configuration / ExportManager / OAuth2Installer plus the Ens.Config "
        "Credentials and Security.SSLConfigs provisioning paths."
    )

    doc.add_heading("3.3 MCP.Catalog", level=2)
    doc.add_paragraph(
        "AgenticInterop.MCP.Catalog groups the 7 Tool.Catalog methods — class describe, namespace "
        "introspection, glossary, error-code lookup, plus the two vector-catalog search endpoints "
        "(search_ens, search_hs) shared with the Health Interop agent."
    )
    add_screenshot(doc, "24_admin_mcp_fhirserver", "Figure 3.1 — Admin UI MCPs tab showing MCP.FHIRServer. Each MCP groups one or more ToolSets for admin-UI configuration.")

    # ---- 4. Tools ----
    doc.add_page_break()
    doc.add_heading("4. Tools (46 total)", level=1)

    doc.add_heading("4.1 Tool.FHIRServer (26 methods)", level=2)
    add_table(doc,
        ["Method", "Purpose"],
        [
            ["DiscoverFHIRNamespaces", "List IRIS namespaces that are FHIR foundation namespaces (have FHIR metadata defined)"],
            ["ListFHIREndpoints", "Per-namespace inventory of FHIR endpoints with URL, FHIR version, enabled status"],
            ["GetFHIREndpoint", "Full configuration of one endpoint: storage strategy, profile packages, settings"],
            ["GetCapabilityStatement", "Read the endpoint's CapabilityStatement; optional full body"],
            ["ListAvailableFHIRPackages", "List FHIR profile packages installed on the IRIS instance"],
            ["AddFHIRPackages", "Add one or more packages to an endpoint's loaded set"],
            ["UpdateFHIREndpointConfig", "Apply config setting changes to an endpoint"],
            ["SetFHIREndpointEnabled", "Enable or disable an endpoint without deleting it"],
            ["CreateFHIREndpoint", "Create a new endpoint in a foundation namespace with chosen strategy and packages"],
            ["DeleteFHIREndpoint", "Remove an endpoint; optional deleteData=1 also drops the underlying data"],
            ["SearchFHIRResources", "FHIR search against a resource type with optional query string"],
            ["ReadFHIRResource", "Read a single resource by type + id"],
            ["CreateFHIRResource", "POST a new resource (server assigns id)"],
            ["UpdateFHIRResource", "PUT an existing resource by type + id"],
            ["DeleteFHIRResource", "DELETE a resource by type + id"],
            ["ValidateFHIRResource", "Run $validate against an endpoint with optional profile URL"],
            ["CountFHIRResources", "Per-type resource counts (uses _summary=count)"],
            ["ExecuteFHIRRequest", "Escape hatch — run an arbitrary FHIR HTTP request"],
            ["LoadFHIRData", "Synchronous load of one file or a small directory of FHIR JSON"],
            ["LoadFHIRDirectory", "Ordered async load of a server-side directory with infra-first option"],
            ["GetFHIRLoadStatus", "Status + progress of a running async load job"],
            ["GetFHIRLoadMetrics", "Final ingestion metrics for a completed (or last) load"],
            ["ListFHIRLoadRuns", "Durable history of all load runs (FHIRLoadRun table)"],
            ["GetFHIRServerStats", "Per-endpoint resource counts + storage size + journal/index stats"],
            ["GetFHIRQueryPerformance", "Run a query-performance probe against representative resource types"],
            ["ResetFHIRServerData", "Wipe all resources from a non-production endpoint (destructive)"],
        ])
    add_screenshot(doc, "25_admin_tool_fhirserver", "Figure 4.1 — Admin UI Tools tab with Tool.FHIRServer expanded. Each public ClassMethod is a tool the LLM can call; descriptions are LLM-facing contracts.")

    doc.add_heading("4.2 Tool.BulkFHIR (13 methods)", level=2)
    add_table(doc,
        ["Method", "Purpose"],
        [
            ["ListBulkFHIRConfigs", "Inventory of BFC configurations in the namespace"],
            ["GetBulkFHIRConfig", "Full configuration of one BFC config (fetch_config + storage_config)"],
            ["GetBulkFHIRSchema", "Schema for the fetch_config / storage_config classes (used for validation)"],
            ["CreateBulkFHIRConfig", "Create a new BFC config from a name + schema-conforming body"],
            ["ConfigureBulkFHIRConfig", "Apply settings updates to an existing config"],
            ["ActivateBulkFHIRConfig", "Activate a configured but inactive config"],
            ["DeleteBulkFHIRConfig", "Remove a BFC config"],
            ["StartBulkFHIRExport", "Kick off a $export session for the named config"],
            ["MonitorBulkFHIRSession", "Status / progress / counts for a running or completed session"],
            ["ProvisionStorageDir", "Create the storage directory the storage_config points at"],
            ["ProvisionSSLConfig", "Create or update a Security.SSLConfigs entry for outbound TLS"],
            ["ProvisionCredential", "Create or update an Ens.Config.Credentials entry"],
            ["ProvisionSMARTOAuth", "End-to-end: HS.HC.OAuth2.Server.Installer + OAuth2Installer client provisioning"],
        ])

    doc.add_heading("4.3 Tool.Catalog (7 methods)", level=2)
    add_table(doc,
        ["Method", "Purpose"],
        [
            ["GetUserNamespace", "Return the request's effective namespace (X-IRIS-Namespace or dispatch default)"],
            ["ListUserAccessibleNamespaces", "Namespaces the authenticated user can access"],
            ["DescribeClass", "%Dictionary introspection of a class — superclass, properties, methods"],
            ["ExplainStatus", "Decode a %Status into a human description"],
            ["LookupErrorCode", "Lookup a known IRIS error code from the seed table"],
            ["LookupGlossaryTerm", "Lookup a healthcare/IRIS glossary term"],
            ["SearchApiIndex", "Search the curated API topic index for a feature"],
        ])

    # ---- 5. Skills ----
    doc.add_page_break()
    doc.add_heading("5. Skills", level=1)
    doc.add_paragraph(
        "Skills are %AI.Agent.Skill subclasses with markdown INSTRUCTIONS distilled from the InterSystems "
        "IRIS for Health documentation. They load on demand — the agent invokes a skill rather than "
        "carrying all knowledge in the system prompt, which saves thousands of tokens per request. All "
        "five FHIR Assistant skills extend AgenticInterop.Skill.Base (the %OnNew JSON-marshal workaround)."
    )

    skills = [
        ("5.1 Skill.FHIRServer",
         "FHIR R4 server build and administration. Covers foundation namespaces, endpoint creation with "
         "storage strategies (Advanced JSON, JsonAdvSQL), profile package management (hl7.fhir.r4.core, "
         "us-core), CapabilityStatement reading, resource CRUD/search/$validate semantics, async directory "
         "load workflow, and guarded provisioning. Source: IRIS Health Connect FHIR Server documentation."),
        ("5.2 Skill.FHIRR4",
         "FHIR R4 resource modeling, references, search parameters, bundle handling (transaction / batch / "
         "collection), OAuth 2.0 + SMART on FHIR. Source: HL7 FHIR R4 specification + IRIS implementation notes. "
         "Includes the 'Tools you have for the FHIR pipeline' section pointing at ConfigureSDAToFHIRProcess / "
         "ConfigureFHIRToSDAProcess (Health Interop side, useful for FHIR Assistant when discussing pipeline architecture)."),
        ("5.3 Skill.SDA",
         "SDA3 as the canonical pivot between HL7 v2 / CDA / X12 and FHIR. Covers HS.SDA3.Container, the "
         "HS.FHIR.DTL.SDA3.vR4.* DTL library (171 standard mappings), and the lookup-table layer "
         "(^HS.XF.LookupTable). The FHIR Assistant uses this skill to answer 'how does my HL7 message end up "
         "in the FHIR server' questions without taking over the upstream pipeline."),
        ("5.4 Skill.BulkFHIR",
         "Bulk FHIR Coordinator (BFC) configuration model: fetch_config (source endpoint, $export or "
         "$everything, OAuth client, resource types) and storage_config (Storage.File / Storage.Ingestion). "
         "SMART backend authentication flow. Async REST flow for $export. Provisioning chain (storage dir, "
         "SSL/TLS, credential, OAuth server + client)."),
        ("5.5 Skill.FHIRSQLBuilder",
         "Manual walkthrough — no automation. Covers the three FHIR SQL Builder stages (Analysis, "
         "Specification, Projection), how to define columns / subtables / filters, and how to query the "
         "projected tables via SQL / JDBC / ODBC."),
    ]
    for title, body in skills:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    # ---- 6. Chatbot Surface ----
    doc.add_page_break()
    doc.add_heading("6. Chatbot Surface", level=1)
    doc.add_paragraph(
        "The FHIR Assistant is surfaced via two paths:"
    )
    doc.add_paragraph(
        "Path A — Floating launcher button in the FHIR Server Management page. "
        "AgenticInterop.Install.FHIRManagementPatch.Apply() (run on IPM Activate) injects a small "
        "<script> tag into /usr/irissys/csp/hslib/ui/isc-fhir/index.html so the page loads "
        "/agentic/inject.js?chatbot=fhir-management&mode=floating. Clicking the launcher opens an iframe "
        "to /agentic/chat/index.html?chatbot=fhir-management. Idempotent; reverted on Uninstall.",
        style="List Bullet")
    doc.add_paragraph(
        "Path B — Direct chatbot URL. "
        "/agentic/chat/index.html?chatbot=fhir-management resolves the chatbot via Data.Chatbot, looks up "
        "the bound AgentClass (AgenticInterop.Agent.FHIRSpecialist), and constructs an %AI.Agent.Session "
        "with the FHIR-specialist system prompt, MCPs, and skills.",
        style="List Bullet")
    doc.add_paragraph(
        "Both paths share the same Data.Chatbot row (key = fhir-management). Editing the row in the admin "
        "Chatbots tab — for example, swapping AgentClass to a custom subclass — takes effect on the next "
        "chat session without redeploying the patch."
    )
    add_screenshot(doc, "20_fhir_management", "Figure 6.1 — Shipped IRIS FHIR Server Management page at /csp/fhir-management. The FHIR Assistant launcher button is injected by AgenticInterop.Install.FHIRManagementPatch on IPM Activate.")
    add_screenshot(doc, "22_admin_chatbots", "Figure 6.2 — Admin UI Chatbots tab showing both shipped chatbot rows: interop (Health Interop) and fhir-management (FHIR Specialist). The agent class binding is editable without redeploying source.")

    # ---- 7. Security Model ----
    doc.add_page_break()
    doc.add_heading("7. Security Model", level=1)
    doc.add_paragraph(
        "The FHIR Assistant inherits the four-layer enforcement model shared with the Health Interop "
        "agent. FHIR-specific aspects:"
    )
    add_table(doc,
        ["Layer", "FHIR-side enforcement"],
        [
            ["1. Authentication", "JWT or Basic auth at the /api/agentic REST surface (CSP gateway). Unauthenticated requests are rejected before any tool runs."],
            ["2. Namespace validation", "Every FHIR tool resolves the target namespace via the X-IRIS-Namespace header and verifies it is a FHIR foundation namespace before any endpoint operation."],
            ["3. Permission delegation", "FHIR endpoint admin requires %HS_Administrator; BFC create/start requires %HS_BFC_Administrator; SSL/OAuth provisioning requires %Admin_Secure. The IRIS security model enforces these — the tools surface the verbatim error if the role is missing."],
            ["4. User confirmation", "All mutating tools (Create, Update, Delete, Activate, Reset, Provision, Load) are gated by the agent's plan-authorize-act flow. The user must explicitly authorize the specific action."],
        ])
    doc.add_paragraph(
        "FHIR endpoint OAuth — SMART backend services authentication for BFC fetch — is provisioned by "
        "ProvisionSMARTOAuth and runs entirely through HS.HC.OAuth2.Server.Installer. The agent never "
        "stores OAuth credentials itself; the IRIS OAuth server holds them under its own ACL."
    )

    # ---- 8. Data Model ----
    doc.add_page_break()
    doc.add_heading("8. Data Model", level=1)
    add_table(doc,
        ["Class", "Purpose", "Owned by"],
        [
            ["AgenticInterop.Data.FHIRLoadRun", "Durable history of every FHIR directory-load run (job ID, namespace, endpoint, directory, status, counts, duration, errors)", "Tool.FHIRServer.LoadFHIRDirectory writes one row per run; ListFHIRLoadRuns reads."],
            ["AgenticInterop.Data.Chatbot", "Key → AgentClass + host page + title mapping; fhir-management row binds to FHIRSpecialist", "Seeded by Data.Chatbot.EnsureSeed on IPM Activate; editable in the admin Chatbots tab."],
            ["AgenticInterop.Data.AuditLog", "Per-request audit trail (method, path, status, duration, user, namespace) — shared with all agents", "REST.Dispatch writes; admin Audit tab reads."],
            ["AgenticInterop.Data.AgentOverride / MCPOverride / ToolSetOverride", "Per-namespace overrides for agent / MCP / ToolSet configuration without touching source", "Admin UI writes; agent build resolves overrides via Editor.AgentService."],
            ["Ens.Config.DefaultSettings", "Out-of-band runtime overrides — used by FHIR Assistant for environment-specific BFC fetch endpoints, OAuth client IDs, storage directories", "Provisioned via the Health Interop tools (SetSystemDefaultSetting); the FHIR Assistant only consults via the FHIR tools' settings resolution."],
        ])

    # ---- 9. Build Status Summary ----
    doc.add_page_break()
    doc.add_heading("9. Build Status Summary", level=1)
    add_table(doc,
        ["Component", "Detail", "Status"],
        [
            ["Agent class", "AgenticInterop.Agent.FHIRSpecialist (extends %AI.Agent)", "Built"],
            ["MCP classes", "FHIRServer, BulkFHIR, Catalog (extend MCP.Base ← %AI.MCP.Service)", "Built"],
            ["Tool classes", "Tool.FHIRServer (26), Tool.BulkFHIR (13), Tool.Catalog (7) — 46 methods total", "Built"],
            ["Skill classes", "5 skills (FHIRServer, FHIRR4, SDA, BulkFHIR, FHIRSQLBuilder) — extend Skill.Base", "Built"],
            ["Chatbot binding", "Data.Chatbot row key=fhir-management → AgenticInterop.Agent.FHIRSpecialist", "Seeded on Activate"],
            ["FHIR Management UI patch", "AgenticInterop.Install.FHIRManagementPatch injects launcher button", "Applied on Activate; reverted on UnConfigure"],
            ["FHIR upload staging endpoint", "POST/GET/DELETE /api/agentic/fhir/upload — writes to mgr/Temp/agentic-fhir-upload/ then LoadFHIRDirectory reads", "Built"],
            ["FHIR Audit panel", "GET /api/agentic/fhir/audit + left-nav menu in the FHIR Management page", "Built"],
            ["IPM module version", "1.1.0", "Shipped"],
        ])

    path = os.path.join(DOCS, f"02_Technical_Build_Specification_{AGENT_SUFFIX}_{VERSION}.docx")
    doc.save(path)
    print(f"  saved {path}")


# =============================================================================
# DOC 3 — Lessons Learned (FHIR Assistant)
# =============================================================================

def build_doc3():
    doc = Document()
    style_doc(doc)
    add_title_page(doc,
        "Agentic Health Interoperability",
        f"Lessons Learned — FHIR Assistant Agent — {VERSION}\nFindings from building the FHIRSpecialist on IRIS for Health")
    doc.add_paragraph(
        "This document collects lessons learned while building the FHIR Assistant agent. For lessons "
        "specific to the Health Interop generalist (productions, DTL/BPL/HL7 v2 / SDA pipelines), see "
        "03_Lessons_Learned_HealthInterop_v1.1.0.docx."
    )

    doc.add_heading("Table of Contents", level=1)
    for line in [
        "1. Introduction",
        "2. Architecture decisions",
        "3. Plan-Authorize-Act discipline",
        "4. FHIR namespace and foundation rules",
        "5. Async load and durable run history",
        "6. Bulk FHIR and SMART backend",
        "7. FHIR SQL Builder — automation refusal as a feature",
        "8. UI integration with the shipped FHIR Management page",
        "9. Open issues and follow-ups",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    # ---- 1. Introduction ----
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The FHIR Assistant agent shipped after the Health Interop generalist as a deliberate second "
        "agent rather than an extension. The two agents share infrastructure (REST dispatch, audit, "
        "vector catalogs, chatbot config layer) but have separate MCP / Tool / Skill bindings and "
        "separate system prompts. Building the FHIR Assistant taught lessons that informed the chatbot "
        "config layer (which now supports two distinct chatbots — interop and fhir-management) and that "
        "would inform any third agent added later."
    )

    # ---- 2. Architecture decisions ----
    doc.add_heading("2. Architecture decisions", level=1)

    doc.add_heading("2.1 Two agents, not one extended agent", level=2)
    doc.add_paragraph(
        "An early option was to extend Agent.HealthInterop with FHIR-server tools and skills and keep "
        "one agent that switches mental modes based on the question. That option was rejected because "
        "the Health Interop persona — production-builder, plan-then-build, references business hosts and "
        "DTLs — leaks into FHIR conversations and confuses end users who do not work in productions at all. "
        "A separate agent with its own system prompt enforces FHIR-only style and scope."
    )

    doc.add_heading("2.2 Chatbot config layer as the binding mechanism", level=2)
    doc.add_paragraph(
        "The chatbot key resolves at request time from a Data.Chatbot row, not from a hard-coded class "
        "name in the dispatch. This makes adding a third agent later a configuration change rather than "
        "a code change. The FHIR Assistant's chatbot key (fhir-management) was added to the seed in the "
        "same commit that shipped the agent class, ensuring fresh installs get both chatbots."
    )

    doc.add_heading("2.3 MCP layer kept as metadata, not transport", level=2)
    doc.add_paragraph(
        "The MCP classes for the FHIR Assistant (FHIRServer, BulkFHIR, Catalog) follow the same pattern "
        "as the Health Interop MCPs: admin-UI metadata, not HTTP/SSE transport. Tools register directly "
        "via UseToolSet at agent build time. This kept the FHIR Assistant build symmetric with the "
        "existing Health Interop build and avoided introducing an HTTP loop for in-process calls."
    )

    # ---- 3. Plan-Authorize-Act ----
    doc.add_heading("3. Plan-Authorize-Act discipline", level=1)

    doc.add_heading("3.1 Why the FHIR Assistant is stricter than Health Interop", level=2)
    doc.add_paragraph(
        "Health Interop's user is typically an Interface Engineer who wants to build something, and the "
        "plan-then-approve flow there is mostly about avoiding wasted work. The FHIR Assistant's user is "
        "typically an Admin operating a live FHIR endpoint that other systems already depend on. A "
        "premature endpoint delete, package change, or data reset has immediate downstream impact. The "
        "FHIR Assistant system prompt makes the discipline absolute: \"Never create, configure, load, "
        "reset, delete, enable, disable, or provision ANYTHING until the user has explicitly authorized "
        "that exact action in a reply.\""
    )

    doc.add_heading("3.2 \"Having all the details is not authorization\"", level=2)
    doc.add_paragraph(
        "The single phrase that fixed the most fabrication-style failures during testing. Early prompts "
        "treated a detailed user request (\"create a FHIR R4 server with Advanced JSON at /a/b\") as "
        "implicit approval to act. That produced \"Done — your server is ready\" replies in the same turn "
        "as the request, before any tool ran. The new rule treats every mutating message as a request to "
        "PLAN; the agent presents the plan and waits for the user's go-ahead."
    )

    doc.add_heading("3.3 Read-only is exempt", level=2)
    doc.add_paragraph(
        "\"What endpoints exist?\", \"how big is the server?\", \"what's the CapabilityStatement?\" — these "
        "answers are direct, no plan, no confirmation. The discipline applies only to state-changing calls."
    )

    # ---- 4. FHIR namespace rules ----
    doc.add_heading("4. FHIR namespace and foundation rules", level=1)
    doc.add_paragraph(
        "An IRIS namespace must be a FHIR foundation namespace before a FHIR endpoint can be created in "
        "it. The agent uses DiscoverFHIRNamespaces to surface candidates and refuses to create an "
        "endpoint in a plain application namespace. This caught several would-be misconfigurations during "
        "testing where the user assumed any namespace would work."
    )
    doc.add_paragraph(
        "FHIR lookup tables (^HS.XF.LookupTable) have a similar gate: the HS.FHIR.DTL.Util.API.LookupTable "
        "methods refuse with \"<NS> is not a valid namespace for Lookup Table\" unless "
        "HS.Util.Installer.ConfigItem has a row for the namespace. The FHIRR4 / SDA skills now spell this "
        "out so the agent does not promise lookup-table edits in an application namespace."
    )

    # ---- 5. Async load + durable run history ----
    doc.add_heading("5. Async load and durable run history", level=1)
    doc.add_paragraph(
        "LoadFHIRDirectory runs as a background job. Early versions reported only the synchronous "
        "result of \"job started\" and lost track once the chat session ended. The fix was two-part:"
    )
    doc.add_paragraph(
        "Part 1 — GetFHIRLoadStatus and GetFHIRLoadMetrics tools that take a jobId and return current "
        "progress or final metrics. The user can ask \"how's that load doing?\" any time, in any session.",
        style="List Bullet")
    doc.add_paragraph(
        "Part 2 — AgenticInterop.Data.FHIRLoadRun, a persistent row written when the load starts and "
        "updated on completion. ListFHIRLoadRuns reads this; the FHIR Audit panel uses it for the "
        "ingestion-history view. The class is namespace-aware so each foundation namespace has its own "
        "run history.",
        style="List Bullet")

    # ---- 6. BFC + SMART ----
    doc.add_heading("6. Bulk FHIR and SMART backend", level=1)

    doc.add_heading("6.1 Four prerequisites, one tool", level=2)
    doc.add_paragraph(
        "A working BFC export against a SMART-protected source endpoint needs four pieces of "
        "infrastructure: a storage directory on the IRIS server, an SSL/TLS configuration for outbound "
        "TLS, an Ens.Config.Credentials entry, and an OAuth 2.0 server + client. Asking the user to set "
        "these up by hand against the management portal was a frequent failure point. The Provision* "
        "tools (ProvisionStorageDir / ProvisionSSLConfig / ProvisionCredential / ProvisionSMARTOAuth) "
        "automate this end-to-end, are idempotent, and surface IRIS security errors verbatim when the "
        "user lacks %Admin_Secure or %HS_BFC_Administrator."
    )

    doc.add_heading("6.2 Storage modes — file vs ingestion", level=2)
    doc.add_paragraph(
        "BFC supports two storage modes: Storage.File writes ndjson to a directory; Storage.Ingestion "
        "writes to a target FHIR endpoint. The agent disambiguates by asking what the user wants the "
        "exported data for — local files for a downstream batch process, or a replica server. Mixing the "
        "two by mistake (e.g. ingestion config pointing at a directory) used to produce confusing "
        "errors; the GetBulkFHIRSchema tool surfaces the schema constraints so the validation happens "
        "before creation."
    )

    # ---- 7. FHIR SQL Builder ----
    doc.add_heading("7. FHIR SQL Builder — automation refusal as a feature", level=1)
    doc.add_paragraph(
        "FHIR SQL Builder is a Management Portal workflow. There is no programmatic API to script the "
        "three-stage Analysis → Specification → Projection process; trying to drive the UI from code is "
        "fragile and brittle across IRIS versions. The FHIR Assistant explicitly refuses to automate it. "
        "Instead, the FHIRSQLBuilder skill walks the user through it conversationally and answers SQL / "
        "JDBC / ODBC questions about the resulting projected tables. This was a deliberate design choice: "
        "agents that pretend to drive UIs they cannot drive lose user trust quickly."
    )

    # ---- 8. UI integration ----
    doc.add_heading("8. UI integration with the shipped FHIR Management page", level=1)
    doc.add_paragraph(
        "The shipped IRIS FHIR Server Management page at /csp/fhir-management already has the right "
        "audience: anyone administering FHIR endpoints. Putting the FHIR Assistant launcher on that page "
        "(via AgenticInterop.Install.FHIRManagementPatch) means users discover the agent without "
        "navigating elsewhere. The patch is idempotent on apply and reverted on uninstall, and the script "
        "tag injection is small enough that an IRIS upgrade overwriting the host page is easy to repatch."
    )
    doc.add_paragraph(
        "The launcher mode (floating button) and chatbot key (fhir-management) are parameterized in "
        "inject.js — the same /agentic/inject.js asset serves both the Interop Editor (mode=header, "
        "chatbot=interop) and the FHIR Management page (mode=floating, chatbot=fhir-management). No "
        "code duplication."
    )
    add_screenshot(doc, "20_fhir_management", "Figure 8.1 — FHIR Server Management page with the FHIR Assistant launcher injected. Users discover the agent without navigating away.")
    add_screenshot(doc, "21_fhir_chatbot", "Figure 8.2 — Same FHIR Assistant available standalone at /agentic/chat/index.html?chatbot=fhir-management. Both launcher and standalone routes resolve through the same Data.Chatbot row (key=fhir-management).")

    # ---- 9. Open issues ----
    doc.add_heading("9. Open issues and follow-ups", level=1)
    doc.add_paragraph(
        "Open items for the next round of FHIR Assistant work:"
    )
    doc.add_paragraph(
        "FHIR endpoint role-based access tooling — granting users specific FHIR endpoint permissions "
        "(read vs write per resource type) is still a manual Management Portal task. A "
        "ConfigureFHIREndpointAccess tool would mirror the Configure* helpers on the Health Interop side.",
        style="List Bullet")
    doc.add_paragraph(
        "FHIR Lookup table editing from the FHIR Assistant — currently lives on Tool.Transform (Health "
        "Interop side). A pass-through wrapper on Tool.FHIRServer would let the FHIR Assistant edit "
        "lookup tables without crossing into Health Interop scope.",
        style="List Bullet")
    doc.add_paragraph(
        "Streaming progress for long-running loads — GetFHIRLoadStatus polls for progress; an SSE "
        "stream that pushes progress events while a load is running would feel more responsive to the "
        "user. Backend already emits ToolStart/ToolDone via Agent.Manager, but the FHIR-load loop is "
        "still poll-only.",
        style="List Bullet")
    doc.add_paragraph(
        "$everything as a first-class BFC fetch mode — currently $export is the documented path. "
        "$everything is supported but underplayed in the schema; the FHIRR4 skill should be expanded.",
        style="List Bullet")

    path = os.path.join(DOCS, f"03_Lessons_Learned_{AGENT_SUFFIX}_{VERSION}.docx")
    doc.save(path)
    print(f"  saved {path}")


if __name__ == "__main__":
    print("Building FHIR Assistant .docx documents...")
    build_doc1()
    build_doc2()
    build_doc3()
    print("\nAll three FHIR Assistant documents built.")
