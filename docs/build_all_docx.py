"""Build the three project .docx documents with embedded screenshots.

Outputs:
  docs/01_Requirements_User_Stories.docx
  docs/02_Technical_Build_Specification.docx
  docs/03_Lessons_Learned.docx

Requires: pip install python-docx Pillow
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOCS = os.path.dirname(__file__)
IMG = os.path.join(DOCS, "img")


def img(name):
    """Return full path to a screenshot."""
    return os.path.join(IMG, name)


def style_doc(doc):
    """Apply consistent styling to the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 5):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Calibri"
        hstyle.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.styles["Heading 1"].font.size = Pt(22)
    doc.styles["Heading 2"].font.size = Pt(16)
    doc.styles["Heading 3"].font.size = Pt(13)


def add_title_page(doc, title, subtitle, version="Version 3.0 | May 2026 | InterSystems AI Hub"):
    """Add a title page."""
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(version)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.add_page_break()


def add_figure(doc, image_path, caption, width=Inches(6.2)):
    """Add an image with caption below it."""
    if not os.path.exists(image_path):
        p = doc.add_paragraph(f"[Screenshot: {caption}]")
        p.italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure: {caption}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    doc.add_paragraph("")


def add_table(doc, headers, rows):
    """Add a styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "1A1A2E"
        })
        shading.append(bg)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")


# =============================================================================
# Document 1: Requirements & User Stories
# =============================================================================
def build_doc1():
    doc = Document()
    style_doc(doc)
    add_title_page(doc,
        "Agentic Health Interoperability",
        "Requirements and User Stories")

    # Table of contents placeholder
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("    1.1 Product Vision")
    doc.add_paragraph("    1.2 The %AI Framework Foundation")
    doc.add_paragraph("    1.3 Framework Extensions")
    doc.add_paragraph("2. Personas")
    doc.add_paragraph("    2.1 Developer")
    doc.add_paragraph("    2.2 AI Hub Admin")
    doc.add_paragraph("    2.3 End User: Interface Engineer")
    doc.add_paragraph("    2.4 End User: Operator")
    doc.add_paragraph("    2.5 Interface Engineer vs Operator: Tool Access")
    doc.add_paragraph("3. Building the Foundation: Vector Catalogs")
    doc.add_paragraph("4. Building the Foundation: LLM Connections")
    doc.add_paragraph("5. Core Use Cases")
    doc.add_paragraph("    5.1 Build Productions (Interface Engineer)")
    doc.add_paragraph("    5.2 Review and Improve Existing Productions (Operator, with Interface Engineer escalation)")
    doc.add_paragraph("    5.3 Create and Optimize Transformations (Interface Engineer)")
    doc.add_paragraph("6. Developer Experience")
    doc.add_paragraph("7. AI Hub Admin Experience")
    doc.add_paragraph("8. End User Experience")
    doc.add_paragraph("    8.1 The Chatbot")
    doc.add_paragraph("    8.2 Performance Guardrails")
    doc.add_paragraph("    8.3 Interface Engineer User Stories")
    doc.add_paragraph("    8.4 Operator User Stories")
    doc.add_paragraph("9. Audit and Security Requirements")
    doc.add_paragraph("10. Source Control and Change Control Integration")
    doc.add_paragraph("11. End-to-End Scenario")
    doc.add_paragraph("12. Non-Functional Requirements")
    doc.add_paragraph("Appendix A: Admin UI Tab Summary")
    doc.add_paragraph("Appendix B: %AI Framework Primitives Used")
    doc.add_paragraph("Appendix C: Security Enforcement Layers")
    doc.add_page_break()

    # =========================================================================
    # 1. Introduction
    # =========================================================================
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Product Vision", level=2)
    doc.add_paragraph(
        "A configuration-driven AI Copilot for InterSystems IRIS for Health that helps "
        "integration engineers build Productions, create Transformations, test HL7/FHIR "
        "messages, and explore the IRIS class catalog -- entirely through natural-language "
        "conversation or a structured admin UI.")
    doc.add_paragraph(
        "The copilot bridges the gap between healthcare data expertise and InterSystems "
        "platform knowledge. Instead of navigating Management Portal screens and writing "
        "ObjectScript by hand, engineers describe what they need in plain English and the "
        "copilot builds it using real IRIS APIs.")
    add_figure(doc, img("15_chatbot.png"),
        "The chatbot interface -- the End User's primary interface for interacting with the copilot")

    doc.add_heading("1.2 The %AI Framework Foundation", level=2)
    doc.add_paragraph(
        "The entire solution is built on the InterSystems %AI Framework -- the native AI "
        "infrastructure shipped with IRIS for Health 2026.2. The framework provides six "
        "core primitives that the copilot uses directly:")
    add_table(doc,
        ["%AI Primitive", "Role in the Copilot"],
        [
            ["%AI.Agent", "The central orchestrator. One agent instance (HealthInterop) receives user messages, decides which tools or skills to invoke, and orchestrates multi-step workflows"],
            ["%AI.MCP.Service", "Model Context Protocol servers that group related capabilities into named service domains (Production, Transform, Testing, Catalog)"],
            ["%AI.ToolSet", "Tool grouping classes that organize tools by domain. Each ToolSet maps to one MCP server"],
            ["%AI.Tool", "Individual tool implementations. Each public ClassMethod in a Tool class is a callable tool with JSON Schema input/output"],
            ["%AI.Agent.Skill", "Declarative sub-agents that package domain knowledge as markdown INSTRUCTIONS. The main agent delegates domain questions to the appropriate skill"],
            ["%AI.RAG.KnowledgeBase", "Vector storage with HNSW index for semantic search over the IRIS class library (Business Hosts, adapters, transformation classes)"],
        ])
    doc.add_paragraph(
        "The copilot does not replace or modify the %AI Framework. It builds on top of it "
        "-- every Agent, MCP, ToolSet, Tool, and Skill is a standard %AI subclass that the "
        "framework manages natively.")

    doc.add_heading("1.3 Framework Extensions", level=2)
    doc.add_paragraph(
        "During development, we encountered three issues in the %AI Framework that required "
        "application-level workarounds. These are extensions, not modifications -- the "
        "framework classes are untouched.")

    doc.add_paragraph(
        "Extension 1: AgenticInterop.Skill.Base (workaround for %AI.Agent.Skill %OnNew bug)")
    doc.add_paragraph(
        "The %AI.Agent.Skill %OnNew method passes a %DynamicObject to $ZF (the Foreign "
        "Function Interface call to the Rust LLM bridge), but the bridge expects a JSON "
        "string. This throws a <FUNCTION> error that prevents any skill from being "
        "instantiated. We created AgenticInterop.Skill.Base that overrides %OnNew to "
        "serialize the object before the $ZF call. All 15 domain skills extend this base class "
        "instead of %AI.Agent.Skill directly.",
        style="List Bullet")

    doc.add_paragraph(
        "Extension 2: Anthropic direct provider (workaround for Bedrock tool-result hang)")
    doc.add_paragraph(
        "When the agent calls a tool and receives the result, the Rust bridge hangs "
        "indefinitely when sending the tool result back to the AWS Bedrock Converse API. "
        "The hang occurs below the ObjectScript API surface. We switched to the Anthropic "
        "direct provider, which works correctly. The same agent, tools, and skills work "
        "without any code changes -- only the LLM connection configuration changes.",
        style="List Bullet")

    doc.add_paragraph(
        "Extension 3: AgenticInterop.Policy.ToolFilter (framework tool cleanup)")
    doc.add_paragraph(
        "The %AI Framework exposes default tools (FileSystem, SQL, ShellTools) that are "
        "irrelevant for healthcare interoperability and waste LLM tokens. The ToolFilter "
        "policy strips these before each LLM call, leaving only the 118 healthcare-specific "
        "tools defined by the project.",
        style="List Bullet")

    # =========================================================================
    # 2. Personas
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("2. Personas", level=1)
    doc.add_paragraph(
        "The system serves four distinct personas. The End User role is split into Interface Engineer "
        "and Operator because the distinction affects which tools are available, which "
        "permissions are required, and whether the agent operates as a dev-time or run-time tool.")

    doc.add_heading("2.1 Developer", level=2)
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Role", "Builds agent infrastructure: writes Tool classes in ObjectScript, authors Skill documents with INSTRUCTIONS, builds vector catalog embeddings, packages and deploys via IPM"],
            ["Primary interface", "VS Code with InterSystems ObjectScript extension, terminal"],
            ["Security scope", "Full %DB access, %Dictionary write, source control, IPM packaging"],
            ["Deliverable", "Compiled classes inside an IPM package that the AI Hub Admin configures"],
        ])
    doc.add_paragraph(
        "The Developer defines what the copilot can do. They write the code that implements "
        "tools, skills, and MCP servers. Their work happens in VS Code and ships as compiled "
        "classes.")

    doc.add_heading("2.2 AI Hub Admin", level=2)
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Role", "Configures all AI settings: creates agents with custom system prompts, assembles MCP Servers from available ToolSets, links Skills to Agents, manages LLM connections, builds vector catalogs, reviews audit logs. Assigns Interface Engineer and Operator roles to end users"],
            ["Primary interface", "IRIS Management Portal -- AI Hub admin UI at /agentic/admin/"],
            ["Security scope", "%ISCMgtPortal group membership, /api/agentic/ endpoints, Secured Wallet write for API keys"],
            ["Deliverable", "A fully configured agent ready for end users to interact with, with appropriate tool access per role"],
        ])
    doc.add_paragraph(
        "The AI Hub Admin decides how the copilot behaves. They configure the agent's "
        "personality, which tools are available, which skills are loaded, and which LLM "
        "provider powers the responses. They also control which tools are available to "
        "Interface Engineers versus Operators through ToolSet configuration and role-based tool "
        "filtering. No code editing required -- everything is configuration through the "
        "admin UI.")
    add_figure(doc, img("01_agents_list.png"),
        "Admin UI -- the AI Hub Admin's primary interface for configuring the copilot")

    doc.add_heading("2.3 End User: Interface Engineer", level=2)
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Role", "Uses the chatbot to create new integration artifacts: productions, DTLs, BPLs, routing rules, lookup tables. This is a dev-time role -- the Interface Engineer authors new content that will be deployed"],
            ["Primary interface", "Chatbot at /agentic/chat/index.html (standalone or embedded in the Interop Editor)"],
            ["Security scope", "Chat access plus create/update/delete permissions on interoperability classes. All mutating operations require explicit approval via the confirmation gate. Changes flow through source control hooks"],
            ["Deliverable", "New or modified productions, transformations, and routing rules -- exported to source control"],
            ["Primary use cases", "UC-1 (Build Productions), UC-3 (Create Transformations)"],
        ])
    doc.add_paragraph(
        "The Interface Engineer creates new integration artifacts through conversation. Because Interface Engineers "
        "create and modify class definitions, their work triggers source control hooks (see "
        "Section 10) and feeds into the CI/CD pipeline. Interface Engineer tool access includes "
        "create_production, add_business_host, create_dtl, compile_dtl, create_routing_rule, "
        "and other mutating tools.")

    doc.add_heading("2.4 End User: Operator", level=2)
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Role", "Uses the chatbot to monitor, triage, and review existing integrations at run-time. The Operator does not create new productions or DTLs -- they observe, diagnose, and recommend"],
            ["Primary interface", "Chatbot at /agentic/chat/index.html (standalone or embedded in the Interop Editor)"],
            ["Security scope", "Chat access plus read-only access to production configurations and monitoring data. May adjust operational settings (pool size, throttle, retry intervals) but not structural changes"],
            ["Deliverable", "Diagnosis reports, remediation recommendations, settings adjustments, modernization advice"],
            ["Primary use cases", "UC-2 (Review and Improve Existing Productions)"],
        ])
    doc.add_paragraph(
        "The Operator focuses on run-time concerns: error triage, throughput monitoring, "
        "queue depth analysis, and production health assessment. They can recommend changes "
        "(refactor this DTL, add a dead-letter queue, increase pool size) but structural "
        "changes require an Interface Engineer. Operator tool access includes get_production, "
        "query_event_log, top_errors, message_summary, queue_status, and other read/monitoring "
        "tools.")

    doc.add_heading("2.5 Interface Engineer vs Operator: Tool Access", level=2)
    doc.add_paragraph(
        "The distinction between Interface Engineer and Operator is enforced through tool availability. "
        "The AI Hub Admin configures which ToolSets are available to each role:")
    add_table(doc,
        ["Tool Category", "Int. Eng.", "Operator", "Examples"],
        [
            ["Create/Update/Delete", "Yes", "No", "create_production, create_dtl, add_business_host"],
            ["Start/Stop", "Yes", "No", "start_production, stop_production"],
            ["Compile", "Yes", "No", "compile_dtl"],
            ["Read/Inspect", "Yes", "Yes", "get_production, list_dtls, describe_class"],
            ["Monitoring", "Yes", "Yes", "query_event_log, top_errors, message_summary"],
            ["Search", "Yes", "Yes", "search_ens, search_hs"],
            ["Testing", "Yes", "Limited", "send_hl7, validate_hl7_structure"],
            ["Settings adjustment", "Yes", "Yes (operational only)", "update_business_host_settings"],
        ])
    doc.add_paragraph(
        "This separation matters for security: the Operator's agent instance cannot create or "
        "modify class definitions, which means it cannot accidentally (or through prompt "
        "injection) alter production behavior. The AI Hub Admin enforces this by binding "
        "different ToolSets to Interface Engineer-mode and Operator-mode agent configurations.")

    # =========================================================================
    # 3. Building the Foundation: Vector Catalogs
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("3. Building the Foundation: Vector Catalogs", level=1)
    doc.add_paragraph(
        "Before the copilot can help end users build productions or create transformations, "
        "it needs to know what Business Hosts, adapters, and transformation classes are "
        "available in the IRIS class library. This knowledge comes from two semantic search "
        "catalogs built with IRIS Vector Search.")

    doc.add_heading("3.1 Why Vector Catalogs", level=2)
    doc.add_paragraph(
        "Integration engineers face a discovery problem: IRIS for Health ships hundreds of "
        "Business Host classes (services, processes, operations) and transformation classes "
        "(DTL, SDA helpers, FHIR mappers). Finding the right class for a given requirement "
        "means searching documentation, browsing %Dictionary, or asking a colleague.")
    doc.add_paragraph(
        "The vector catalogs solve this by embedding curated descriptions of every relevant "
        "class into a searchable vector index. When an end user asks 'build a production "
        "that receives HL7 messages over TCP', the agent searches the Ens.* catalog for "
        "'HL7 TCP inbound service' and gets back EnsLib.HL7.Service.TCPService as the top "
        "result -- without the user needing to know the exact class name.")

    doc.add_heading("3.2 The Ens.* Catalog (164 Business Hosts and Adapters)", level=2)
    doc.add_paragraph(
        "The search_ens catalog indexes every class extending Ens.Host, Ens.BusinessService, "
        "Ens.BusinessProcess, Ens.BusinessOperation, Ens.OutboundAdapter, and "
        "Ens.InboundAdapter. For each class, the builder extracts class name and description, "
        "superclass hierarchy, key configurable parameters, and message types accepted "
        "and produced.")
    doc.add_paragraph(
        "The extraction uses %Dictionary.ClassDefinition as the source of truth. Curated "
        "prose descriptions feed the embeddings -- auto-generated accessor methods and "
        "structural boilerplate are stripped to keep the semantic signal strong in the "
        "384-dimensional embedding space.")

    doc.add_heading("3.3 The HS.* Catalog (58 Transformation Classes)", level=2)
    doc.add_paragraph(
        "The search_hs catalog indexes HealthShare-specific transformation classes: DTL "
        "classes, FHIR mappers, SDA helpers, and HL7 gateways under the HS.* hierarchy. "
        "This catalog powers the transformation use case: when the agent needs to find "
        "existing transformations for a format pair, it searches search_hs rather than "
        "listing every DTL class in the namespace.")

    doc.add_heading("3.4 Technical Implementation", level=2)
    add_table(doc,
        ["Attribute", "Detail"],
        [
            ["Embedding model", "FastEmbed (384-dimensional vectors, bundled with IRIS)"],
            ["Vector storage", "%AI.RAG.KnowledgeBase with HNSW index"],
            ["Query path", "%AI.ToolMgr.ExecuteTool(kbName, args) -- the only working path. SQL EMBEDDING() does not work with bundled FastEmbed"],
            ["Document format", "Curated prose descriptions (class name + description + superclass + key parameters), not raw class dumps"],
            ["Rebuild trigger", "Admin UI Catalogs tab or API call"],
        ])

    doc.add_heading("3.5 Building and Rebuilding Catalogs", level=2)
    doc.add_paragraph(
        "The AI Hub Admin builds the catalogs from the admin UI Catalogs tab. Each catalog "
        "shows its row count, last rebuild timestamp, kind breakdown (by superclass), and "
        "provides a test search panel for validating search quality.")
    add_figure(doc, img("11_catalogs.png"),
        "Catalogs tab showing search_ens (164 rows indexed) and search_hs (58 rows indexed) with kind breakdowns, rebuild buttons, test search panel, and browse panel")
    doc.add_paragraph(
        "Catalogs should be rebuilt when IRIS is upgraded (new classes may be available) or "
        "when the catalog builder logic changes. For all other cases, the persistent HNSW "
        "index serves queries without rebuilding.")

    # =========================================================================
    # 4. Building the Foundation: LLM Connections
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("4. Building the Foundation: LLM Connections", level=1)
    doc.add_paragraph(
        "The copilot requires an LLM provider to generate responses. The AI Hub Admin "
        "configures LLM connections through the admin UI -- no environment variables, "
        "no config files, no code changes.")

    doc.add_heading("4.1 Supported Providers", level=2)
    add_table(doc,
        ["Provider", "Config Fields", "Secret Field"],
        [
            ["AWS Bedrock", "region, model", "AWS_BEARER_TOKEN_BEDROCK"],
            ["Anthropic", "model, base URL", "ANTHROPIC_API_KEY"],
            ["OpenAI", "model, base URL", "OPENAI_API_KEY"],
            ["Azure OpenAI", "endpoint, deployment, API version", "AZURE_OPENAI_API_KEY"],
            ["Google Gemini", "model, region", "GEMINI_API_KEY"],
            ["NVIDIA NIM", "model, base URL", "NIM_API_KEY"],
        ])

    doc.add_heading("4.2 Connection Lifecycle", level=2)
    doc.add_paragraph(
        "1. Create: The AI Hub Admin enters connection details in the admin UI Connections "
        "tab -- provider type, model name, region, base URL")
    doc.add_paragraph(
        "2. Store secret: The API key is entered in a masked input field. On Save, the key "
        "is written to the IRIS Secured Wallet under collection AgenticInteropConnections. "
        "The key is never stored in SQL tables, globals, or source code")
    doc.add_paragraph(
        "3. Test: The 'Test Connection' button sends a minimal completion request (1 token) "
        "to the configured provider. On success, it displays the model name and response "
        "latency. On failure, it shows the error text verbatim")
    doc.add_paragraph(
        "4. Status: Green dot (last test OK), red dot (last test failed), gray dot (never tested)")
    doc.add_paragraph(
        "5. Bind: The agent configuration references a connection by name. At request time, "
        "the Manager loads the secret from the Wallet and configures the LLM client")

    add_figure(doc, img("09_connections_list.png"),
        "Connections tab showing the configured connection with DEFAULT and CORE badges, green status dot")
    add_figure(doc, img("10_connection_detail.png"),
        "Connection detail editor showing provider, model, region, API key (masked), and Test Connection button with green TESTED OK status")

    doc.add_heading("4.3 Security: The IRIS Secured Wallet", level=2)
    doc.add_paragraph(
        "All API keys are stored exclusively in the IRIS Secured Wallet (%Wallet.KeyValue, "
        "collection AgenticInteropConnections). Security invariants:")
    doc.add_paragraph("API keys are NEVER stored in SQL tables, globals, or source code", style="List Bullet")
    doc.add_paragraph("API keys are NEVER returned by any REST endpoint (not even masked)", style="List Bullet")
    doc.add_paragraph("API keys are NEVER logged in the audit trail", style="List Bullet")
    doc.add_paragraph("The Wallet is the single source of truth for secrets", style="List Bullet")
    doc.add_paragraph("Connection test results (including provider error messages) ARE logged for debugging", style="List Bullet")

    # =========================================================================
    # 5. Core Use Cases
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("5. Core Use Cases", level=1)
    doc.add_paragraph(
        "The copilot addresses three primary use cases that cover the full lifecycle of "
        "healthcare integration work inside IRIS for Health. Each use case maps to a primary "
        "End User persona (Interface Engineer or Operator).")
    add_figure(doc, img("15_chatbot.png"),
        "Chat interface showing starter prompts organized by use case category: Build, Transform, Operate, and Review")

    # UC-1: Build Productions (Interface Engineer)
    doc.add_heading("5.1 Use Case 1: Build Productions (Interface Engineer)", level=2)
    doc.add_paragraph(
        "The most common task for an integration engineer is building new Productions -- "
        "the runtime message-processing pipelines in IRIS for Health. A Production consists "
        "of Business Services (inbound), Business Processes (routing/orchestration), and "
        "Business Operations (outbound), wired together with settings, routing rules, and "
        "message transformations.")
    doc.add_paragraph("The agent assists the Interface Engineer through the entire production lifecycle:")
    doc.add_paragraph(
        "Discovery: The End User describes their integration goal in plain English (e.g., "
        "'build a production that receives ADT messages over MLLP, transforms them to FHIR "
        "R4, and sends them to a REST endpoint'). The agent searches the Ens.* vector "
        "catalog to find the right Business Host classes (EnsLib.HL7.Service.TCPService, "
        "EnsLib.FHIR.Operation.REST, etc.).",
        style="List Bullet")
    doc.add_paragraph(
        "Proposal: The agent presents a production layout -- which hosts to add, what "
        "settings to configure, which adapters to use -- and asks the Interface Engineer to approve "
        "before making changes.",
        style="List Bullet")
    doc.add_paragraph(
        "Build: Upon approval, the agent creates the production class, adds each Business "
        "Host with appropriate settings, creates routing rules, and compiles everything. "
        "Each mutating step goes through the confirmation gate.",
        style="List Bullet")
    doc.add_paragraph(
        "Validation: The agent runs PostBuildValidation to check for configuration errors, "
        "sends a test HL7 message through the pipeline, and verifies that messages flow "
        "end-to-end without errors.",
        style="List Bullet")
    doc.add_paragraph(
        "Source Control: After successful build, the newly created classes (production "
        "definition, routing rules) are captured by Health Connect Cloud source control "
        "hooks and exported to the Git repository for CI/CD review (see Section 10).",
        style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Tools involved:")
    doc.add_paragraph("search_ens -- find the right Business Host classes from the vector catalog", style="List Bullet")
    doc.add_paragraph("describe_class -- inspect class details, parameters, and settings", style="List Bullet")
    doc.add_paragraph("create_production -- create the production class definition", style="List Bullet")
    doc.add_paragraph("add_business_host -- add Business Services, Processes, and Operations", style="List Bullet")
    doc.add_paragraph("update_business_host_settings -- configure adapter settings, file paths, connection parameters", style="List Bullet")
    doc.add_paragraph("create_routing_rule -- create routing rules with conditions and actions", style="List Bullet")
    doc.add_paragraph("start_production / stop_production -- lifecycle management", style="List Bullet")
    doc.add_paragraph("PostBuildValidation -- automated post-build health check", style="List Bullet")
    doc.add_paragraph("BuildAndSendHL7TestMessage -- generate and send test messages through the pipeline", style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Skills involved: Productions, Adapters, HL7v2, RoutingRules")
    doc.add_paragraph("")
    doc.add_paragraph("Example prompts:")
    doc.add_paragraph(
        '"Build a complete production that receives HL7 v2.5 ADT^A01 admission messages '
        'over an inbound folder, transforms each ADT into an ORU^R01 observation report, '
        'routes the transformed messages to an outbound folder, and sends failures to a '
        'dead-letter folder."',
        style="List Bullet")
    doc.add_paragraph(
        '"I need a production that receives X12 270 eligibility inquiries over SFTP, calls '
        'our internal eligibility REST API, constructs the X12 271 response, and writes it '
        "back to the payer's SFTP outbound folder.\"",
        style="List Bullet")

    # UC-2: Review and Improve Productions (Operator, with Interface Engineer escalation)
    doc.add_heading("5.2 Use Case 2: Review and Improve Existing Productions (Operator, with Interface Engineer escalation)", level=2)
    doc.add_paragraph(
        "Integration engineers inherit productions built by others, or maintain productions "
        "that were built months or years ago. They need to understand what a production does, "
        "identify problems, and find opportunities to modernize it using newer IRIS features "
        "and best practices.")
    doc.add_paragraph(
        "This use case spans both personas. The Operator handles the read-only investigative "
        "work (error triage, health assessment, throughput analysis). When the investigation "
        "reveals changes that need to be made (refactor a DTL, add a dead-letter queue, "
        "restructure routing rules), those changes escalate to the Interface Engineer.")

    doc.add_paragraph("")
    doc.add_paragraph("Operator activities (read-only, run-time):")
    doc.add_paragraph(
        "Error Triage: The agent queries the Event Log and Message Header tables to find "
        "recent errors, groups them by Business Host, identifies the most frequent error "
        "messages, and recommends remediation steps. It can spot suspended or errored "
        "messages that need manual intervention.",
        style="List Bullet")
    doc.add_paragraph(
        "Production Health Assessment: The agent inspects the production configuration, "
        "checks queue depths, reviews throughput statistics, and identifies bottlenecks. "
        "It can recommend settings changes (pool size, throttle, retry intervals) based "
        "on what it observes.",
        style="List Bullet")
    doc.add_paragraph(
        "Operational Settings Adjustment: The Operator can adjust operational settings "
        "(pool size, throttle, retry intervals, call interval) through the agent without "
        "Interface Engineer involvement. These are runtime tuning changes, not structural modifications.",
        style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Interface Engineer activities (mutating, dev-time):")
    doc.add_paragraph(
        "DTL Review and Refactoring: The agent reviews Data Transformation Language (DTL) "
        "definitions and identifies hardcoded values that should be lookup tables, missing "
        "null checks on source fields, incorrect handling of repeating fields, and segments "
        "being dropped. It suggests refactored versions with explanations. Implementing the "
        "refactored DTL requires Interface Engineer permissions.",
        style="List Bullet")
    doc.add_paragraph(
        "Modernization: The agent knows about newer IRIS features (via Skills) and can "
        "recommend upgrades -- for example, replacing a custom BPL with a built-in DTL, "
        "using record maps instead of custom parsers, or adopting the HL7-to-SDA-to-FHIR "
        "pipeline instead of point-to-point transformations. Implementing these changes "
        "requires Interface Engineer permissions.",
        style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Tools involved (Operator):")
    doc.add_paragraph("get_production -- inspect the full production configuration and all hosts", style="List Bullet")
    doc.add_paragraph("query_event_log -- search the Event Log for errors, warnings, and trace messages", style="List Bullet")
    doc.add_paragraph("top_errors -- group errors by frequency and identify systemic issues", style="List Bullet")
    doc.add_paragraph("query_message_status -- find messages in Error or Suspended state", style="List Bullet")
    doc.add_paragraph("message_summary -- throughput statistics across all hosts", style="List Bullet")
    doc.add_paragraph("queue_status -- check for queue buildup indicating backpressure", style="List Bullet")
    doc.add_paragraph("describe_class -- look up what a Business Host class does and its available settings", style="List Bullet")
    doc.add_paragraph("list_dtls / get_dtl -- review existing transformation logic", style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Tools involved (Interface Engineer escalation):")
    doc.add_paragraph("update_business_host_settings -- reconfigure adapter settings", style="List Bullet")
    doc.add_paragraph("update_dtl / compile_dtl -- modify and compile transformation logic", style="List Bullet")
    doc.add_paragraph("create_routing_rule -- add new routing rules", style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Skills involved: Productions, DTL, BPL, Adapters, ESBPattern")
    doc.add_paragraph("")
    doc.add_paragraph("Example prompts (Operator):")
    doc.add_paragraph(
        '"Review the last 2 hours of errors across all productions. Group them by Business '
        'Host, show the top 5 most frequent error messages with counts, identify messages '
        'stuck in Suspended or Error state, and recommend remediation steps."',
        style="List Bullet")
    doc.add_paragraph(
        '"What is the throughput of the ADT_Router process over the last 24 hours? Are '
        'there any queue depth spikes?"',
        style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph("Example prompts (Interface Engineer):")
    doc.add_paragraph(
        '"Review our current ADT_A08_to_SDA3 DTL for: hardcoded values that should be '
        'lookup tables, missing null checks on source fields, incorrect handling of '
        'repeating PID-3 identifiers, and segments we are dropping that we should not."',
        style="List Bullet")
    doc.add_paragraph(
        '"Refactor the ADT routing to use the built-in HL7-to-SDA-to-FHIR pipeline '
        'instead of our custom point-to-point DTLs."',
        style="List Bullet")

    # UC-3: Create and Optimize Transformations (Interface Engineer)
    doc.add_heading("5.3 Use Case 3: Create and Optimize Transformations (Interface Engineer)", level=2)
    doc.add_paragraph(
        "Data Transformations are the heart of healthcare interoperability. Integration "
        "engineers spend most of their time writing, debugging, and optimizing DTL (Data "
        "Transformation Language) and BPL (Business Process Language) definitions that "
        "convert messages between formats -- HL7 v2 to SDA3, SDA3 to FHIR R4, CDA to "
        "SDA3, and more.")
    doc.add_paragraph("The agent assists the Interface Engineer with transformation work at every stage:")
    doc.add_paragraph(
        "Pipeline Discovery: The agent traces the full transformation pipeline for any "
        "format pair (e.g., HL7 v2 to FHIR R4) showing which IRIS classes handle each "
        "step, what intermediate formats are used, and where the data flows. The "
        "Transformation and Mapping Catalog (Transforms tab) provides this information "
        "visually at the field level.",
        style="List Bullet")
    doc.add_paragraph(
        "DTL Creation: The agent creates new DTL definitions by first searching the HS.* "
        "catalog for existing transformations that handle the same or similar format pair, "
        "then scaffolding a new DTL with the correct source/target classes and document "
        "types.",
        style="List Bullet")
    doc.add_paragraph(
        "Schema Introspection: The agent can introspect HL7 v2 message schemas (segments, "
        "fields, components) and FHIR R4 resource structures so the Interface Engineer understands "
        "what data is available at each point in the pipeline.",
        style="List Bullet")
    doc.add_paragraph(
        "Dry-Run Testing: The agent can execute a DTL against a sample message (DryRunDTL) "
        "to verify the transformation produces the expected output without deploying to a "
        "production.",
        style="List Bullet")
    doc.add_paragraph(
        "Cross-Format Mapping Insights: Through the Transformation and Mapping Catalog, "
        "the agent can see exactly which HL7 fields map through SDA3 to FHIR, which "
        "fields are inbound-only, and which are outbound-only. This enables gap analysis "
        "before writing any code.",
        style="List Bullet")
    doc.add_paragraph(
        "Source Control: After creating or modifying DTLs and BPLs, the source control "
        "hooks capture the new class definitions and export them to Git for CI/CD review "
        "(see Section 10).",
        style="List Bullet")

    add_figure(doc, img("13_transforms_hl7_fhir.png"),
        "Transformation and Mapping Catalog: HL7 v2 to FHIR R4 via SDA3.Address, showing sub-field level mappings, coverage filter chips, and IRIS class names inline")

    doc.add_paragraph("")
    doc.add_paragraph("Tools involved:")
    doc.add_paragraph("list_dtls -- discover existing transformations in the namespace", style="List Bullet")
    doc.add_paragraph("create_dtl -- scaffold a new DTL with source/target classes", style="List Bullet")
    doc.add_paragraph("update_dtl / compile_dtl -- modify and compile transformation logic", style="List Bullet")
    doc.add_paragraph("dry_run_dtl -- test a transformation against sample data", style="List Bullet")
    doc.add_paragraph("list_sda_fhir_dtls -- find built-in SDA3-to-FHIR transformations", style="List Bullet")
    doc.add_paragraph("describe_transformation_pipeline -- trace the full format conversion path", style="List Bullet")
    doc.add_paragraph("get_hl7_schema_map / get_hl7_segment_fields -- introspect HL7 v2 message structures", style="List Bullet")
    doc.add_paragraph("search_hs -- semantic search for transformation classes in the HS.* catalog", style="List Bullet")
    doc.add_paragraph("compare_messages -- field-level diff between two messages", style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Skills involved: DTL, HL7v2, FHIRR4, SDA, CDA, X12")
    doc.add_paragraph("")
    doc.add_paragraph("Example prompts:")
    doc.add_paragraph(
        '"Create an interface that accepts any HL7 v2 message (ADT, ORU, ORM, MDM, SIU) '
        'on a single inbound MLLP service, transforms it to the appropriate FHIR R4 '
        'resources using the built-in HL7-to-SDA-to-FHIR pipeline, and POSTs the '
        'resulting Bundle to our FHIR Server."',
        style="List Bullet")
    doc.add_paragraph(
        '"Build a production that ingests C-CDA documents via a REST endpoint, validates '
        'them against the C-CDA R2.1 schema, transforms them to FHIR R4 Composition + '
        'DocumentReference + Patient/Encounter/Condition resources using SDA3 as the '
        'intermediate model, and persists the Bundle to our FHIR repository."',
        style="List Bullet")

    # =========================================================================
    # 6. Developer Experience
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("6. Developer Experience", level=1)
    doc.add_paragraph(
        "Developers work exclusively through VS Code with the InterSystems ObjectScript "
        "extension. They write classes (Agent, MCP, ToolSet, Tool, Skill), compile them, "
        "and deploy via IPM. The admin UI is not their primary interface -- their deliverable "
        "is code that the AI Hub Admin configures.")

    doc.add_heading("6.1 User Stories", level=2)
    stories_dev = [
        ("US-D01: Create a New Tool",
         "As a Developer, I want to write an ObjectScript class that extends %AI.Tool with "
         "defined parameters (NAME, DESCRIPTION, INPUT, OUTPUT), so that the tool appears in "
         "the agent's catalog and can be invoked during chat conversations.",
         ["Tool class compiles without errors",
          "Tool appears in the admin UI Tools tab after compilation",
          "Tool description follows the contract format: imperative verb, scope, side effects",
          "Tool input/output schemas are valid JSON Schema",
          "Tool includes at least one happy-path unit test",
          "Tool specifies whether it is available to Interface Engineer, Operator, or both roles"]),
        ("US-D02: Create a New Skill",
         "As a Developer, I want to write a Skill class with INSTRUCTIONS content (markdown "
         "text up to 32K characters), so that the agent can delegate domain-specific questions "
         "to a specialist sub-agent.",
         ["Skill class extends AgenticInterop.Skill.Base (not %AI.Agent.Skill directly -- see Section 1.3)",
          "INSTRUCTIONS parameter contains domain knowledge in plain prose",
          "Skill registers automatically via SkillLoader at agent build time",
          "AI Hub Admin can override INSTRUCTIONS content in the admin UI without code changes"]),
        ("US-D03: Create a New MCP Server",
         "As a Developer, I want to write an MCP Server class that groups related ToolSets "
         "under a named service, so that AI Hub Admins can enable/disable entire capability "
         "domains from the admin UI.",
         ["MCP class extends AgenticInterop.MCP.Base (which extends %AI.MCP.Service)",
          "Parameters: NAME, DESCRIPTION, TOOLSETS (comma-separated list)",
          "MCP appears in the admin UI MCPs tab after compilation"]),
        ("US-D04: Deploy via IPM",
         "As a Developer, I want to package the entire project as an IPM module (agentic-interop), "
         "so that an AI Hub Admin can install it on any IRIS for Health 2026.2+ instance.",
         ["module.xml defines all sources, CSP applications, seed data, and install hooks",
          "zpm load on a clean namespace produces a working system",
          "Install hooks: CSP timeout patch applied, Interop Editor patched with AI buttons",
          "Uninstall hooks: Interop Editor reverted to original state"]),
        ("US-D05: Write Custom Tool Implementations",
         "As a Developer, I want to implement tools using SQL statements, ObjectScript class "
         "methods, or Embedded Python, so that I can leverage the best language for each task.",
         ["SQL tools execute parameterized queries (no string concatenation)",
          "ObjectScript tools use $namespace for namespace-agnostic operation",
          "Python tools use ##class(%SYS.Python).Import() for LLM/MCP glue",
          "All tools handle errors gracefully and return structured error objects"]),
        ("US-D06: Build and Maintain Vector Catalogs",
         "As a Developer, I want to rebuild the Ens.* and HS.* vector catalogs from %Dictionary, "
         "so that the agent can semantically search for Business Hosts and transformation classes.",
         ["AgenticInterop.Catalog.Builder walks %Dictionary.ClassDefinition for relevant superclasses",
          "Embeddings use FastEmbed (384-dim HNSW vectors) via %AI.RAG.KnowledgeBase",
          "Curated prose descriptions (not auto-generated accessor signatures) feed the embeddings",
          "Catalog rebuild can be triggered from admin UI or scheduled"]),
        ("US-D07: Implement Permission-Aware Tools",
         "As a Developer, I want to write tools that validate the authenticated user's permissions "
         "before executing, so that the LLM agent cannot bypass the user's security constraints "
         "even if the tool is technically available.",
         ["Every mutating tool receives the authenticated username from the REST layer",
          "The tool checks the user's roles and database privileges before executing (e.g., $System.Security.Check())",
          "If the user lacks permission, the tool returns a structured error explaining which permission is missing",
          "The agent receives this error and explains to the user what happened (not a silent failure)",
          "Foundation namespace (HSLIB, HSSYS, ENSLIB) writes are always rejected regardless of user role"]),
    ]
    for title, story, criteria in stories_dev:
        doc.add_heading(title, level=3)
        doc.add_paragraph(story)
        doc.add_paragraph("Acceptance criteria:")
        for c in criteria:
            doc.add_paragraph(c, style="List Bullet")

    # =========================================================================
    # 7. AI Hub Admin Experience
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("7. AI Hub Admin Experience", level=1)
    doc.add_paragraph(
        "AI Hub Admins work through the IRIS Management Portal. They configure agents, "
        "manage LLM connections, review transformation mappings, tune skills, and prepare "
        "the copilot for end users. No code editing required.")
    doc.add_paragraph(
        "The admin UI is a vanilla JavaScript SPA served at /agentic/admin/. It communicates "
        "with /api/agentic/ REST endpoints using JWT or Basic authentication. Every action "
        "is logged to the audit trail.")

    # 7.2 Agent configuration
    doc.add_heading("7.1 Configure the Agent", level=2)
    doc.add_paragraph(
        "The AI Hub Admin customizes the agent's system prompt, temperature, max iterations, "
        "bound MCPs, and skills so that the agent behaves according to the organization's "
        "integration requirements.")
    add_figure(doc, img("02_agent_detail.png"),
        "Agent editor showing the HealthInterop agent configuration: class name, temperature slider, max iterations, tool binding mode, system prompt with persona and formatting rules")
    doc.add_paragraph("Agent configuration fields:", style="List Bullet")
    doc.add_paragraph("Name, description, and instructions (system prompt textarea)", style="List Bullet 2")
    doc.add_paragraph("Temperature slider (default 0.3) and max iterations (default 25)", style="List Bullet 2")
    doc.add_paragraph("MCP binding: checkbox list of available MCP servers", style="List Bullet 2")
    doc.add_paragraph("Skill binding: checkbox list of available skills", style="List Bullet 2")
    doc.add_paragraph("Provider selection: dropdown of configured connections", style="List Bullet 2")
    doc.add_paragraph("Tool binding mode: MCP chain or bypass (Agent -> Tool directly)", style="List Bullet 2")
    doc.add_paragraph("Changes saved as override rows that survive IPM upgrades", style="List Bullet 2")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: The AI Hub Admin can configure separate agent profiles for Interface Engineer and "
        "Operator roles, each with different ToolSet bindings. This enforces the privilege "
        "separation described in Section 2.5.")

    # 7.3 MCP Servers
    doc.add_heading("7.2 Configure MCP Servers", level=2)
    doc.add_paragraph(
        "The AI Hub Admin enables/disables MCP servers and customizes their descriptions "
        "to control which capability domains the agent has access to.")
    add_figure(doc, img("03_mcps_list.png"),
        "MCPs tab showing the six MCP servers: catalog, production, testing, transform, fhirserver, and bulkfhir, each with their class name, description, and toolset count")
    add_table(doc,
        ["MCP Server", "ToolSets", "Purpose"],
        [
            ["mcp.production", "Production", "CRUD productions, business hosts, settings, start/stop"],
            ["mcp.transform", "Transform", "CRUD DTL/BPL, routing rules, lookup tables, HL7 schema"],
            ["mcp.testing", "Testing", "Send HL7/FHIR messages, validate, compare"],
            ["mcp.catalog", "Catalog, Monitoring", "Vector search, class introspection, event log, throughput"],
        ])

    # 7.4 ToolSets and Tools
    doc.add_heading("7.3 Configure ToolSets and Tools", level=2)
    doc.add_paragraph(
        "The AI Hub Admin views and customizes ToolSets and their individual tools -- "
        "tuning descriptions, toggling tools on/off, and dry-running tools to verify behavior.")
    add_figure(doc, img("04_toolsets_list.png"),
        "ToolSets tab showing the seven ToolSets: Catalog (7 tools), Monitoring (5 tools), Production (29 tools), Testing (8 tools), Transform (30 tools), FHIRServer (26 tools), BulkFHIR (13 tools)")
    add_figure(doc, img("05_tools_list.png"),
        "Tools tab listing all 118 tools across the seven ToolSet classes")
    doc.add_paragraph("Clicking a tool opens the detail editor with the LLM-facing description and schema:")
    add_figure(doc, img("06_tool_detail.png"),
        "Tool detail editor for DescribeClass showing the contract-style description, method signature, JSON input schema, implementation kind, timeout, and confirmation requirement")

    # 7.5 Skills
    doc.add_heading("7.4 Configure Skills", level=2)
    doc.add_paragraph(
        "The AI Hub Admin views and edits the INSTRUCTIONS content for each skill, refining "
        "the agent's domain knowledge without Developer involvement.")
    add_figure(doc, img("07_skills_list.png"),
        "Skills tab showing the 12 shipped skills: Adapters, BPL, CDA, DTL, ESBPattern, FHIRR4, HL7v2, Productions, and more")
    add_figure(doc, img("08_skill_detail.png"),
        "Skill editor showing class name, description, bound ToolSets, and expandable class source viewer")
    add_table(doc,
        ["Skill", "Domain"],
        [
            ["Productions", "Production anatomy, BS/BP/BO patterns, lifecycle management"],
            ["DTL", "DTL syntax, foreach, subtransforms, lookup tables, virtual documents"],
            ["BPL", "BPL activities, compensation handlers, async patterns"],
            ["RoutingRules", "Rule sets, constraints, when-conditions, dead-letter handling"],
            ["HL7v2", "Message types, segments, ACK semantics, schema navigation"],
            ["FHIRR4", "Resources, references, search parameters, R4 bundles"],
            ["SDA", "SDA3 model as transformation hub, HL7-to-SDA-to-FHIR pipeline"],
            ["RestInProductions", "REST services and operations inside productions"],
            ["ESBPattern", "Using a production as an Enterprise Service Bus"],
            ["X12", "HIPAA EDI transactions, envelope structures, schemas"],
            ["CDA", "CDA/C-CDA documents, XSLT pipelines, SDA conversion"],
            ["Adapters", "File/TCP/HTTP/REST/FTP/SQL/MQTT/SOAP adapter selection"],
        ])

    # 7.6 Transformation and Mapping Catalog
    doc.add_heading("7.5 Review Transformation Mappings", level=2)
    doc.add_paragraph(
        "The Transforms tab provides the Transformation and Mapping Catalog: a visual "
        "field-level mapping explorer showing how data flows between external formats "
        "through the SDA3 canonical model.")
    add_figure(doc, img("12_transforms_empty.png"),
        "Transforms tab initial state: format pair selection (Data From / Data To dropdowns) with 1,538 pre-computed rows")
    add_figure(doc, img("13_transforms_hl7_fhir.png"),
        "Transformation and Mapping Catalog with HL7 v2 to FHIR R4 selected, showing SDA3.Address mappings at sub-field level with IRIS class names and coverage filters")
    doc.add_paragraph("Features:")
    doc.add_paragraph("Format pair selection: HL7 v2, FHIR R4, FHIR STU3, CDA, X12, SDA3", style="List Bullet")
    doc.add_paragraph("SDA3 type sidebar: 110 data types, browsable and filterable", style="List Bullet")
    doc.add_paragraph("Sub-field level detail: PID.11.3 City, not just PID-11 PatientAddress", style="List Bullet")
    doc.add_paragraph("IRIS class names inline for each direction of the mapping", style="List Bullet")
    doc.add_paragraph("Coverage filter chips: End-to-end (green), Inbound only (blue), Outbound only (yellow)", style="List Bullet")
    doc.add_paragraph("1,538 pre-computed rows, rebuilt on demand in ~0.2 seconds", style="List Bullet")

    # 7.7 Audit
    doc.add_heading("7.6 View Audit Log", level=2)
    doc.add_paragraph(
        "The Audit tab shows all API requests with filters by kind, username, and date range.")
    add_figure(doc, img("14_audit.png"),
        "Audit tab showing the request log with status codes, HTTP methods, paths, kind classification, timestamps, usernames, namespace, duration, and byte counts")

    # 7.9 AI Hub Covered Stories
    doc.add_heading("7.7 User Stories Covered by the AI Hub Framework", level=2)
    doc.add_paragraph(
        "The following capabilities are required by this project but are expected to be "
        "provided by the AI Hub framework itself rather than implemented in the application "
        "layer. They are listed here to ensure coverage tracking:")
    add_table(doc,
        ["Story", "Requirement", "Expected Coverage"],
        [
            ["US-H01", "Agent lifecycle management (create, configure, delete agents)", "AI Hub admin UI"],
            ["US-H02", "MCP server registration and discovery", "AI Hub framework (%AI.MCP.Service)"],
            ["US-H03", "Tool schema validation and registration", "AI Hub framework (%AI.Tool)"],
            ["US-H04", "LLM provider abstraction (Bedrock, Anthropic, OpenAI, etc.)", "AI Hub framework (LLM bridge)"],
            ["US-H05", "SSE streaming from agent to client", "AI Hub framework (%AI.Agent.StreamChat)"],
            ["US-H06", "Conversation state management across turns", "AI Hub framework (conversation context)"],
            ["US-H07", "Role-based access control for AI Hub admin operations", "AI Hub framework (to be confirmed)"],
            ["US-H08", "Source control hook integration for class changes", "Health Connect Cloud (%SourceControl hooks)"],
        ])
    doc.add_paragraph(
        "If any of these are not provided by the framework, they must be built at the "
        "application layer. The current implementation includes application-level workarounds "
        "for US-H04 (Bedrock hang) and US-H07 (custom role checks).")

    # =========================================================================
    # 8. End User Experience
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("8. End User Experience", level=1)

    doc.add_heading("8.1 The Chatbot", level=2)
    doc.add_paragraph(
        "The End User (Interface Engineer or Operator) interacts with the copilot through a streaming "
        "chat interface. The chatbot is available at /agentic/chat/index.html (standalone) "
        "or embedded in the Interop Editor via an AI button (iframe mode).")
    add_figure(doc, img("15_chatbot.png"),
        "Chat interface showing starter prompts (Build, Transform, Operate, Review categories), conversation history sidebar, namespace and connection indicators in the top bar")
    doc.add_paragraph("Key capabilities:")
    doc.add_paragraph("SSE streaming: tokens appear in real time, no loading spinner", style="List Bullet")
    doc.add_paragraph("Tool calls render as inline cards with name, arguments, status, and collapsible result", style="List Bullet")
    doc.add_paragraph("Mutating tool calls pause with Approve/Reject prompt (ConfirmationGate policy)", style="List Bullet")
    doc.add_paragraph("Conversation history rail with search, resume, and rename", style="List Bullet")
    doc.add_paragraph("Starter prompts organized by use case: Build, Transform, Operate, Review", style="List Bullet")
    doc.add_paragraph("Top bar shows agent name, connection status, and New Chat button", style="List Bullet")

    doc.add_heading("8.2 Performance Guardrails", level=2)
    doc.add_paragraph("Monitor enforces 60-second deadline + 50,000 token budget per turn", style="List Bullet")
    doc.add_paragraph("Complex tasks broken into multiple short turns with visible progress at each phase", style="List Bullet")
    doc.add_paragraph("If a turn exceeds limits, the monitor triggers a graceful stop and the agent summarizes partial results", style="List Bullet")

    # Interface Engineer User Stories
    doc.add_heading("8.3 Interface Engineer User Stories", level=2)

    doc.add_heading("US-E01: Build a Production Through Conversation", level=3)
    doc.add_paragraph(
        "As an Interface Engineer, I want to describe an integration requirement in plain English and "
        "have the agent build the production, so that I can create working integrations "
        "without manually navigating the Management Portal.")
    doc.add_paragraph("Acceptance criteria:")
    doc.add_paragraph("Agent searches catalogs for appropriate Business Hosts", style="List Bullet")
    doc.add_paragraph("Agent proposes a production layout and waits for approval", style="List Bullet")
    doc.add_paragraph("Each mutating step goes through the confirmation gate", style="List Bullet")
    doc.add_paragraph("PostBuildValidation runs after build to verify configuration", style="List Bullet")
    doc.add_paragraph("Test message sent through the pipeline to verify end-to-end flow", style="List Bullet")
    doc.add_paragraph("Newly created classes are captured by source control hooks (if configured)", style="List Bullet")

    doc.add_heading("US-E02: Create a Transformation Through Conversation", level=3)
    doc.add_paragraph(
        "As an Interface Engineer, I want to describe a data transformation requirement and have the "
        "agent scaffold the DTL, so that I can create transformations with the correct "
        "source/target classes and field mappings.")
    doc.add_paragraph("Acceptance criteria:")
    doc.add_paragraph("Agent searches HS.* catalog for existing transformations", style="List Bullet")
    doc.add_paragraph("Agent creates DTL with correct source/target document types", style="List Bullet")
    doc.add_paragraph("Dry-run executes against sample data before deployment", style="List Bullet")
    doc.add_paragraph("Compiled DTL is captured by source control hooks (if configured)", style="List Bullet")

    # Operator User Stories
    doc.add_heading("8.4 Operator User Stories", level=2)

    doc.add_heading("US-E03: Triage Production Errors", level=3)
    doc.add_paragraph(
        "As an Operator, I want to ask the agent to review recent errors and group them "
        "by cause, so that I can quickly identify systemic issues and prioritize remediation.")
    doc.add_paragraph("Acceptance criteria:")
    doc.add_paragraph("Agent queries Event Log for recent errors (no mutating operations)", style="List Bullet")
    doc.add_paragraph("Results grouped by Business Host and error message", style="List Bullet")
    doc.add_paragraph("Suspended and errored messages identified", style="List Bullet")
    doc.add_paragraph("Remediation steps recommended (but not automatically applied)", style="List Bullet")
    doc.add_paragraph("Agent does NOT have access to create/delete tools in Operator mode", style="List Bullet")

    doc.add_heading("US-E04: Assess Production Health", level=3)
    doc.add_paragraph(
        "As an Operator, I want to ask the agent for a health assessment of a running "
        "production, so that I can identify bottlenecks and tuning opportunities without "
        "reading raw metrics.")
    doc.add_paragraph("Acceptance criteria:")
    doc.add_paragraph("Agent inspects production configuration, queue depths, and throughput", style="List Bullet")
    doc.add_paragraph("Agent recommends operational settings changes (pool size, throttle, retry)", style="List Bullet")
    doc.add_paragraph("Agent can apply operational settings adjustments if the Operator approves", style="List Bullet")
    doc.add_paragraph("Agent cannot make structural changes (add/remove hosts, create classes)", style="List Bullet")

    # =========================================================================
    # 9. Audit and Security Requirements
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("9. Audit and Security Requirements", level=1)

    doc.add_heading("9.1 Authentication", level=2)
    add_table(doc,
        ["Method", "Use Case"],
        [
            ["Basic auth", "Direct admin UI and chatbot access (standalone mode)"],
            ["JWT Bearer", "Embedded access from Interop Editor (token passed via postMessage bridge)"],
        ])
    doc.add_paragraph(
        "All REST endpoints require authentication. UnauthenticatedEnabled=0 on the "
        "/api/agentic/ web application. No UI element (banner, button, link, modal, or "
        "text) is visible before login.")

    doc.add_heading("9.2 Authorization and Role-Based Access", level=2)
    doc.add_paragraph("AI Hub Admin operations require %ISCMgtPortal group membership", style="List Bullet")
    doc.add_paragraph("End User chat access requires authenticated IRIS user", style="List Bullet")
    doc.add_paragraph(
        "Interface Engineer-mode tools (create, update, delete, compile, start, stop) require "
        "specific roles assigned by the AI Hub Admin",
        style="List Bullet")
    doc.add_paragraph(
        "Operator-mode tools (read, monitor, query) require basic authenticated access",
        style="List Bullet")
    doc.add_paragraph(
        "Mutating operations require explicit Approve from the End User via the "
        "ConfirmationGate policy -- the agent cannot modify the system without user consent",
        style="List Bullet")
    doc.add_paragraph(
        "Cross-namespace access validated via database-level read permissions and "
        "X-IRIS-Namespace header",
        style="List Bullet")

    doc.add_heading("9.3 LLM Service Identity and Permission Delegation", level=2)
    doc.add_paragraph(
        "The LLM agent executes tool calls within the IRIS process. This creates a critical "
        "security requirement: the agent must NEVER bypass the authenticated user's "
        "permissions, even if the agent's service identity has broader access.")
    doc.add_paragraph(
        "The principle: The agent acts on behalf of the authenticated user, not on its own "
        "behalf. Every tool call is executed under the authenticated user's security context. "
        "If the user cannot perform an action through the Management Portal, the agent must "
        "not perform it either.")
    doc.add_paragraph("")
    doc.add_paragraph("Implementation requirements:")
    doc.add_paragraph(
        "Every tool receives the authenticated username from the REST layer (extracted "
        "from Basic auth or JWT)",
        style="List Bullet")
    doc.add_paragraph(
        "Before executing a mutating operation, the tool checks the user's roles and "
        "database privileges via $System.Security.Check() or equivalent",
        style="List Bullet")
    doc.add_paragraph(
        "If the user lacks the required permission, the tool returns a structured error: "
        '{ "error": { "code": "PERMISSION_DENIED", "message": "...", "required_permission": "..." } }',
        style="List Bullet")
    doc.add_paragraph(
        "The agent surfaces this error to the user in plain language",
        style="List Bullet")
    doc.add_paragraph(
        "The audit log records both the attempted action and the permission denial",
        style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph("Example scenarios:")
    add_table(doc,
        ["User Request", "User Has Permission?", "Agent Behavior"],
        [
            ["Create a new production", "Yes (%DB_WRITE)", "Agent proceeds with confirmation gate"],
            ["Create a new production", "No (read-only access)", "Agent explains the user lacks write permission"],
            ["Create an OAuth 2.0 client", "No (requires %Admin_Secure)", "Agent explains the user needs security admin privileges"],
            ["Query the event log", "Yes (any authenticated user)", "Agent executes the read-only query"],
            ["Modify classes in HSLIB", "No (Foundation namespace)", "Agent explains Foundation namespaces are read-only"],
        ])

    doc.add_heading("9.4 Namespace Constraints", level=2)
    doc.add_paragraph(
        "Not all namespaces are equal. IRIS for Health distinguishes between Foundation "
        "namespaces (shipped with the product) and non-Foundation namespaces (created by "
        "users for their integration work).")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Foundation namespaces (HSLIB, HSSYS, ENSLIB, and others marked as Foundation): "
        "always read-only for the agent, regardless of user permissions. The agent can search "
        "and inspect classes but MUST NOT create, modify, or delete any class definition. "
        "This is a hard constraint -- the ConfirmationGate cannot override it.")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Non-Foundation namespaces (HSCUSTOM, user-created namespaces): the agent respects "
        "the authenticated user's database-level permissions. Users with %DB_WRITE can create "
        "and modify classes (through the agent with confirmation). Users with read-only "
        "access can inspect and search but not modify.")

    doc.add_heading("9.5 Interface Engineer vs Operator Privilege Separation", level=2)
    doc.add_paragraph(
        "The security distinction between Interface Engineer and Operator is enforced at two levels:")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Level 1: Tool availability (configured by AI Hub Admin). The AI Hub Admin binds "
        "different ToolSets to Interface Engineer-mode and Operator-mode agent configurations. Operator "
        "agents do not have access to mutating tools. Even if an Operator user has "
        "database-level write access, the agent cannot call tools that are not bound to the "
        "Operator configuration. This prevents prompt injection attacks from escalating "
        "Operator sessions to Interface Engineer-level access.")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Level 2: Permission validation (enforced by tools). Even in Interface Engineer mode, every "
        "mutating tool validates the user's permissions before executing. The ConfirmationGate "
        "provides a third layer: the user must explicitly approve each mutating action.")

    doc.add_heading("9.6 Audit Logging", level=2)
    doc.add_paragraph("Every REST request is captured in AgenticInterop.Data.AuditLog:")
    add_table(doc,
        ["Field", "Description"],
        [
            ["Created", "When the request was received"],
            ["Username", "Authenticated IRIS user"],
            ["Namespace", "Active namespace at request time"],
            ["SessionId", "Browser session identifier"],
            ["Job", "IRIS job number"],
            ["Method", "GET, POST, PUT, DELETE"],
            ["Path", "/api/agentic/chat/stream, /api/agentic/registry/agents, etc."],
            ["StatusCode", "HTTP response status"],
            ["RequestSize", "Bytes received"],
            ["ResponseSize", "Bytes sent"],
            ["DurationMs", "End-to-end request time in milliseconds"],
            ["ErrorText", "Error message (if status >= 400), including permission denials"],
            ["Kind", "Classification: registry, editor.agent, chat, namespace, health"],
        ])
    add_figure(doc, img("14_audit.png"),
        "Audit log showing all API requests with status codes, methods, paths, timestamps, users, and durations")

    doc.add_heading("9.7 Secret Management", level=2)
    doc.add_paragraph(
        "All API keys are stored in the IRIS Secured Wallet (%Wallet.KeyValue, collection "
        "AgenticInteropConnections):")
    doc.add_paragraph("API keys are NEVER stored in SQL tables, globals, or source code", style="List Bullet")
    doc.add_paragraph("API keys are NEVER returned by any REST endpoint", style="List Bullet")
    doc.add_paragraph("API keys are NEVER logged in the audit trail", style="List Bullet")
    doc.add_paragraph("The Wallet is the single source of truth for secrets", style="List Bullet")

    doc.add_heading("9.8 Security Policies", level=2)
    doc.add_paragraph(
        "ConfirmationGate policy: Mutating tools (create, update, delete, start, stop) "
        "pause execution and surface an Approve/Reject prompt in the chat UI. The agent "
        "cannot modify productions, transformations, or routing rules without the End User "
        "clicking Approve. This is the last line of defense after permission checks and "
        "namespace validation.")
    doc.add_paragraph(
        "ToolFilter policy: Strips framework-default tools (FileSystem, SQL, ShellTools) "
        "from the LLM's tool catalog, leaving only the 118 healthcare-specific "
        "tools, saving ~5K tokens per request. This policy also prevents the LLM from using "
        "tools that could bypass the permission model (e.g., raw SQL execution).")

    # =========================================================================
    # 10. Source Control and Change Control Integration
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("10. Source Control and Change Control Integration", level=1)

    doc.add_heading("10.1 The Problem", level=2)
    doc.add_paragraph(
        "When the agent creates a new production, DTL, BPL, or routing rule, it is creating "
        "or modifying IRIS class definitions. These class changes must flow through the same "
        "change control process as any manual edit -- version control tracking, peer review, "
        "and CI/CD validation. Without source control integration, agent-created artifacts "
        "become invisible to the deployment pipeline.")

    doc.add_heading("10.2 Health Connect Cloud Source Control Model", level=2)
    doc.add_paragraph(
        "Health Connect Cloud provides version control hooks via the %SourceControl "
        "framework. These hooks intercept class save and compile events, export class "
        "definitions to a local Git repository, track which user made which change, and "
        "feed into a GitLab-based CI/CD pipeline that validates, tests, and deploys changes "
        "across environments (dev -> staging -> production).")
    doc.add_paragraph(
        "The agent's work must integrate with this model. Agent-created artifacts are NOT "
        "a special case -- they follow the same path as manually created artifacts.")

    doc.add_heading("10.3 Integration Requirements", level=2)
    doc.add_paragraph(
        "Artifact capture: When the agent creates or modifies a class, the %SourceControl "
        "hooks fire automatically because the agent uses standard IRIS APIs "
        "($System.OBJ.Compile, %Dictionary) that trigger the hooks.",
        style="List Bullet")
    doc.add_paragraph(
        "User attribution: Changes are attributed to the authenticated user (the Interface Engineer), "
        "not to the agent's service account. The source control hooks capture the $username "
        "from the IRIS process context.",
        style="List Bullet")
    doc.add_paragraph(
        "No bypass: The agent does not use any mechanism that would bypass source control "
        "hooks. All class modifications go through the documented %Dictionary and "
        "$System.OBJ APIs.",
        style="List Bullet")

    doc.add_heading("10.4 CI/CD Pipeline Flow", level=2)
    add_figure(doc, img("16_cicd_pipeline.png"),
        "CI/CD Pipeline Flow: how agent-created artifacts flow from conversation to production deployment through source control hooks, Git, and the CI/CD pipeline")

    doc.add_heading("10.5 User Stories", level=2)
    doc.add_heading("US-SC01: Agent-Created Artifacts Appear in Source Control", level=3)
    doc.add_paragraph(
        "As an Interface Engineer, I want the productions, DTLs, and BPLs that the agent creates to be "
        "captured by source control hooks, so that they flow through the same CI/CD pipeline "
        "as manually created artifacts.")
    doc.add_paragraph("Acceptance criteria:")
    doc.add_paragraph("Agent uses standard IRIS APIs for all class operations", style="List Bullet")
    doc.add_paragraph("%SourceControl hooks fire on every class save and compile", style="List Bullet")
    doc.add_paragraph("Exported class definitions appear in the local Git working directory", style="List Bullet")
    doc.add_paragraph("Changes are attributed to the authenticated user, not the agent", style="List Bullet")

    doc.add_heading("US-SC02: Agent Reports Source Control Status", level=3)
    doc.add_paragraph(
        "As an Interface Engineer, I want the agent to report whether source control captured the new "
        "classes after a build, so that I know whether my changes are tracked.")
    doc.add_paragraph("Acceptance criteria:")
    doc.add_paragraph("After creating classes, the agent checks whether %SourceControl hooks are active", style="List Bullet")
    doc.add_paragraph("If active: agent reports 'Classes exported to source control'", style="List Bullet")
    doc.add_paragraph("If not active: agent warns 'Source control is not configured in this namespace'", style="List Bullet")

    doc.add_heading("US-SC03: Agent Respects Source Control Locks", level=3)
    doc.add_paragraph(
        "As an Interface Engineer, I want the agent to respect source control locks on classes, so that "
        "it does not modify a class that another user has locked for editing.")
    doc.add_paragraph("Acceptance criteria:")
    doc.add_paragraph("Before modifying a class, the agent checks whether it is locked", style="List Bullet")
    doc.add_paragraph("If locked: agent reports 'This class is locked by [user]. Cannot modify.'", style="List Bullet")
    doc.add_paragraph("If unlocked: agent proceeds with the normal confirmation gate flow", style="List Bullet")

    # =========================================================================
    # 11. End-to-End Scenario
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("11. End-to-End Scenario", level=1)
    doc.add_paragraph("This scenario demonstrates all four personas working together:")
    steps = [
        ("Developer", "writes a new Tool class that creates HL7 routing rules, compiles it, and deploys via zpm load"),
        ("AI Hub Admin", "opens the admin UI, sees the new tool in the Tools tab, reviews its description and tests it with the dry-run panel"),
        ("AI Hub Admin", "configures Interface Engineer and Operator agent profiles with appropriate ToolSet bindings"),
        ("AI Hub Admin", "goes to the Connections tab, verifies the LLM connection shows a green status dot"),
        ("AI Hub Admin", "opens the Catalogs tab, verifies both catalogs (search_ens: 164 classes, search_hs: 58 classes) are indexed"),
        ("AI Hub Admin", "opens the Transforms tab, selects HL7 v2 -> FHIR R4, reviews Address field mappings to verify the Transformation and Mapping Catalog is populated"),
        ("Int. Eng.", "opens the chatbot and asks: 'Build me a production that receives ADT^A04 messages via MLLP, transforms patient demographics to FHIR R4, and sends them to a REST endpoint'"),
        ("Agent", "validates that the Interface Engineer has %DB_WRITE permission on the target namespace"),
        ("Agent", "searches the Ens.* catalog for appropriate Business Hosts (EnsLib.HL7.Service.TCPService, EnsLib.FHIR.Operation.REST)"),
        ("Agent", "proposes the production layout and asks the Interface Engineer to approve"),
        ("Int. Eng.", "clicks Approve"),
        ("Agent", "creates the production, adds the hosts, configures settings. Source control hooks capture each class change"),
        ("Agent", "builds and sends a test HL7 ADT^A04 message"),
        ("Agent", "validates the result and reports success, confirms source control captured the changes"),
        ("Operator", "opens the chatbot the next day and asks: 'How is the ADT production performing? Any errors in the last 24 hours?'"),
        ("Agent", "queries the Event Log (read-only), reports throughput statistics and any issues"),
        ("Operator", "asks: 'Increase the pool size on the TCPService to 3' -- the agent adjusts the operational setting after approval"),
        ("AI Hub Admin", "reviews the audit log to see the complete trace of all actions by both Interface Engineer and Operator"),
    ]
    for i, (persona, action) in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. ")
        run = p.add_run(persona)
        run.bold = True
        p.add_run(f" {action}")

    # =========================================================================
    # 12. Non-Functional Requirements
    # =========================================================================
    doc.add_heading("12. Non-Functional Requirements", level=1)
    add_table(doc,
        ["Requirement", "Target"],
        [
            ["Response latency", "First token in < 2 seconds; full response in < 90 seconds per turn"],
            ["Concurrent users", "5 simultaneous chat sessions (single IRIS instance)"],
            ["Catalog rebuild", "< 30 seconds for full Ens.* + HS.* re-index"],
            ["Field mapping rebuild", "< 1 second for full HL7/SDA3/FHIR trace (1,538 rows)"],
            ["Availability", "System operational whenever IRIS is running; no external dependencies except LLM provider"],
            ["Data retention", "Audit logs retained indefinitely; no automatic purge"],
            ["Browser support", "Chrome 120+, Edge 120+, Firefox 120+ (ES2020 baseline)"],
            ["Permission check latency", "< 10ms per tool call (cached role lookups)"],
        ])

    # =========================================================================
    # Appendices
    # =========================================================================
    doc.add_page_break()
    doc.add_heading("Appendix A: Admin UI Tab Summary", level=1)
    add_table(doc,
        ["Tab", "Purpose", "Persona", "Entity Count"],
        [
            ["Agents", "Agent configuration (system prompt, MCPs, skills, provider)", "AI Hub Admin", "2 (HealthInterop, FHIRSpecialist)"],
            ["MCPs", "MCP server enable/disable and description", "AI Hub Admin", "6 servers"],
            ["ToolSets", "ToolSet grouping and description", "AI Hub Admin", "7 ToolSets"],
            ["Tools", "Individual tool schemas and dry-run", "AI Hub Admin", "118 tools"],
            ["Skills", "Skill INSTRUCTIONS editor", "AI Hub Admin", "15 domain skills"],
            ["Connections", "LLM provider credentials and health check", "AI Hub Admin", "N (user-configured)"],
            ["Catalogs", "Vector catalog status, rebuild, search", "AI Hub Admin", "2 catalogs"],
            ["Transforms", "Field-level mapping explorer (Transformation and Mapping Catalog)", "AI Hub Admin / End User", "1,538 rows"],
            ["Chatbots", "Bind each chatbot surface to an agent", "AI Hub Admin", "2 (Interop, FHIR Management)"],
            ["Audit", "Request audit trail", "AI Hub Admin", "All API calls"],
        ])

    doc.add_heading("Appendix B: %AI Framework Primitives Used", level=1)
    add_table(doc,
        ["Framework Class", "Application Subclass", "Purpose"],
        [
            ["%AI.Agent", "AgenticInterop.Agent.HealthInterop + FHIRSpecialist", "Main + FHIR-specialist agent instances"],
            ["%AI.MCP.Service", "AgenticInterop.MCP.Base + 6 servers", "MCP server grouping"],
            ["%AI.ToolSet", "7 ToolSet classes", "Tool grouping by domain"],
            ["%AI.Tool", "7 Tool classes (118 methods)", "Individual tool implementations"],
            ["%AI.Agent.Skill", "AgenticInterop.Skill.Base + 15 domain skills", "Domain knowledge sub-agents"],
            ["%AI.RAG.KnowledgeBase", "search_ens, search_hs", "Vector search catalogs"],
            ["%AI.ToolMgr", "Used at query time", "RAG query execution"],
            ["%AI.Agent.Policy", "ConfirmationGate, ToolFilter", "Security and token policies"],
        ])

    doc.add_heading("Appendix C: Security Enforcement Layers", level=1)
    doc.add_paragraph(
        "The agent's actions pass through four enforcement layers before any mutation occurs:")
    add_table(doc,
        ["Layer", "Enforced By", "What It Checks"],
        [
            ["1. Tool availability", "AI Hub Admin (ToolSet binding)", "Is this tool bound to the user's agent profile (Interface Engineer vs Operator)?"],
            ["2. Namespace validation", "Tool implementation", "Is the target namespace non-Foundation? Does the user have database access?"],
            ["3. Permission delegation", "Tool implementation", "Does the authenticated user have the required IRIS role for this operation?"],
            ["4. User confirmation", "ConfirmationGate policy", "Did the user click Approve for this specific mutating action?"],
        ])
    doc.add_paragraph(
        "A mutating action must pass all four layers. A failure at any layer prevents "
        "execution and returns a clear error to the agent, which explains the denial to "
        "the user.")

    path = os.path.join(DOCS, "01_Requirements_User_Stories.docx")
    doc.save(path)
    print(f"  saved {path}")


# =============================================================================
# Document 2: Technical Build Specification
# =============================================================================
def build_doc2():
    doc = Document()
    style_doc(doc)
    add_title_page(doc,
        "Agentic Health Interoperability",
        "Technical Build Specification")

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document specifies the components needed to deliver a production-grade "
        "Agentic Health Interoperability solution. The system is a configuration-driven AI "
        "Copilot embedded in IRIS for Health that enables integration engineers to build "
        "Productions, create Transformations, and test healthcare messages through "
        "natural-language conversation.")
    doc.add_paragraph(
        "The solution is built on the IRIS %AI Framework (Agent, MCP, ToolSet, Tool, Skill "
        "primitives) and extends it with application-specific infrastructure: a chat UX, an "
        "admin UI, vector catalogs, the Transformation and Mapping Catalog, connection "
        "management, and audit/security controls.")
    doc.add_paragraph(
        "Four personas interact with the system: Developers who build capabilities in code, "
        "AI Hub Admins who configure the copilot through the admin UI, Interface Engineers (dev-time "
        "End Users) who create new integration artifacts, and Operators (run-time End Users) "
        "who monitor, triage, and review existing integrations.")

    # 2. Chatbot UX
    doc.add_page_break()
    doc.add_heading("2. Chatbot UX", level=1)
    doc.add_paragraph(
        "A streaming chat interface embedded in the IRIS Management Portal that connects to "
        "a %AI.Agent via Server-Sent Events (SSE). The chatbot is the End User's primary "
        "interface for interacting with the agent.")
    add_figure(doc, img("15_chatbot.png"),
        "Chat interface with streaming responses, starter prompts, conversation sidebar, and connection status")

    doc.add_heading("2.1 Streaming Responses", level=2)
    doc.add_paragraph(
        "SSE endpoint (POST /chat/stream) emits tokens as they arrive from the LLM. "
        "Each token is a data: event with JSON payload. Tool lifecycle events (tool_start, "
        "tool_args, tool_result, tool_error) render as inline cards.")

    doc.add_heading("2.2 Confirmation Gate", level=2)
    doc.add_paragraph(
        "Mutating tools (create, update, delete) pause execution and surface an "
        "Approve/Reject prompt. The AgenticInterop.Policy.ConfirmationGate policy "
        "intercepts tool calls before execution. Rejection feeds back to the agent as a "
        "tool error.")

    doc.add_heading("2.3 Performance Guardrails", level=2)
    doc.add_paragraph(
        "AgenticInterop.Agent.Monitor enforces per-turn limits: 60-second deadline and "
        "50,000 token budget. If a turn exceeds limits, the monitor triggers a graceful "
        "stop and the agent summarizes partial results.")

    # 3. Agent
    doc.add_page_break()
    doc.add_heading("3. Agent", level=1)
    doc.add_paragraph(
        "A single %AI.Agent instance ('HealthInterop') serves as the router agent. It "
        "receives user messages, decides which tools or skills to invoke, and orchestrates "
        "multi-step workflows.")
    add_figure(doc, img("02_agent_detail.png"),
        "Agent configuration editor showing system prompt, temperature, max iterations, tool binding mode, and MCP/skill bindings")

    doc.add_heading("3.1 Agent Architecture", level=2)
    doc.add_paragraph("AgenticInterop.Agent.HealthInterop (extends %AI.Agent)")
    doc.add_paragraph("  Manager: builds configured agent at request time (loads provider, binds MCPs, loads skills, attaches policies)", style="List Bullet")
    doc.add_paragraph("  Monitor: iteration callback enforcing 60s deadline + 50K token budget", style="List Bullet")
    doc.add_paragraph("  SkillLoader: discovers Skill.Base subclasses and registers as tools", style="List Bullet")

    doc.add_heading("3.2 Runtime Configuration", level=2)
    add_table(doc,
        ["Parameter", "Default", "Configurable Via"],
        [
            ["System prompt", "Shipped class INSTRUCTIONS", "Admin UI Agent editor (AI Hub Admin)"],
            ["Temperature", "0.3", "Admin UI slider"],
            ["Max iterations", "25", "Admin UI input"],
            ["Bound MCPs", "All 4", "Admin UI checkbox list"],
            ["Bound skills", "All 12", "Admin UI checkbox list"],
            ["LLM provider", "bedrock-default", "Admin UI dropdown"],
            ["Tool binding", "MCP chain", "Admin UI radio (mcp/bypass)"],
        ])

    doc.add_heading("3.3 Overlay Pattern", level=2)
    doc.add_paragraph(
        "Shipped class defaults and user customizations both survive across IPM upgrades. "
        "Override tables (AgentOverride, MCPOverride, ToolSetOverride) store admin UI "
        "changes made by the AI Hub Admin. At build time, the Overlay class merges overrides "
        "on top of compiled defaults. 'Reset to defaults' deletes the override row.")

    # 4. Skills
    doc.add_page_break()
    doc.add_heading("4. Skills", level=1)
    doc.add_paragraph(
        "Declarative sub-agents that package domain knowledge as markdown INSTRUCTIONS. "
        "The main agent delegates domain-specific questions to the appropriate skill rather "
        "than answering from general training data.")
    add_figure(doc, img("07_skills_list.png"),
        "Skills tab listing the 12 shipped skills covering IRIS interoperability domains")
    add_figure(doc, img("08_skill_detail.png"),
        "Skill editor showing class name, description, bound ToolSets, and expandable class source")

    doc.add_heading("4.1 Skill Catalog (v1)", level=2)
    add_table(doc,
        ["Skill", "Domain", "Content Source"],
        [
            ["Productions", "Production anatomy, BS/BP/BO patterns, settings", "IRIS documentation"],
            ["DTL", "Data Transformation Language syntax, foreach, subtransforms", "IRIS documentation"],
            ["BPL", "Business Process Language activities, compensation, async", "IRIS documentation"],
            ["RoutingRules", "Rule sets, constraints, when conditions, HL7 routing", "IRIS documentation"],
            ["HL7v2", "Message types, segments, ACK semantics, composite types", "HL7 v2 specification"],
            ["FHIRR4", "Resources, references, search params, server operations", "FHIR R4 specification"],
            ["SDA", "SDA3 model, common pitfalls, mapping to FHIR", "IRIS documentation"],
            ["RestInProductions", "RESTful services inside productions, HTTP adapters", "IRIS documentation"],
            ["ESBPattern", "Enterprise Service Bus patterns, routing, transformation", "Architecture guides"],
            ["X12", "X12 transaction sets, EDI healthcare claims", "X12 specification"],
            ["CDA", "Clinical Document Architecture, CCD, C-CDA", "CDA specification"],
            ["Adapters", "Inbound/outbound adapters, TCP, HTTP, SOAP, file, FTP", "IRIS documentation"],
        ])

    doc.add_heading("4.2 Skill Registration", level=2)
    doc.add_paragraph(
        "Skills extend AgenticInterop.Skill.Base (not %AI.Agent.Skill directly -- see "
        "Lessons Learned for the bug that required this extension). At agent build time, "
        "SkillLoader discovers all Skill.Base subclasses, instantiates them, and registers "
        "them as tools in the agent's tool catalog.")

    # 5. MCP Servers
    doc.add_page_break()
    doc.add_heading("5. MCP Servers", level=1)
    doc.add_paragraph(
        "Four internal MCP (Model Context Protocol) servers group related capabilities "
        "into named service domains. Each MCP maps to one or more ToolSets.")
    add_figure(doc, img("03_mcps_list.png"),
        "MCP servers: catalog, production, testing, and transform")

    add_table(doc,
        ["MCP Server", "ToolSets", "Purpose"],
        [
            ["mcp.production", "Production", "CRUD productions, business hosts, settings, start/stop"],
            ["mcp.transform", "Transform", "CRUD DTL/BPL, routing rules, lookup tables, HL7 schema"],
            ["mcp.testing", "Testing", "Send HL7/FHIR messages, validate, compare"],
            ["mcp.catalog", "Catalog, Monitoring", "Vector search, class introspection, event log, throughput"],
        ])

    # 6. Tools
    doc.add_page_break()
    doc.add_heading("6. Tools", level=1)
    doc.add_paragraph(
        "118 tools across 7 tool classes implement the agent's capabilities. Each tool is "
        "a method with [Tool] annotation, JSON Schema input/output, and a natural-language "
        "description.")
    add_figure(doc, img("05_tools_list.png"),
        "All 118 tools listed across the seven ToolSet classes")
    add_figure(doc, img("06_tool_detail.png"),
        "Tool detail editor showing the LLM contract description, method signature, input schema, and configuration")

    doc.add_heading("6.1 Tool Catalog by ToolSet", level=2)
    add_table(doc,
        ["ToolSet", "Tool Count", "Tools"],
        [
            ["Production", "13", "list_productions, get_production, create_production, delete_production, start_production, stop_production, add_business_host, remove_business_host, update_settings, PostBuildValidation, ..."],
            ["Transform", "13", "list_dtls, get_dtl, create_dtl, update_dtl, compile_dtl, delete_dtl, list_bpls, create_bpl, list_routing_rules, create_routing_rule, introspect_hl7_schema, ..."],
            ["Testing", "8", "send_hl7, send_fhir, validate_hl7_structure, validate_fhir_resource, compare_messages, BuildAndSendHL7TestMessage, ..."],
            ["Catalog", "8", "search_ens, search_hs, describe_class, get_namespace, list_classes, lookup_reference, search_glossary, ..."],
            ["Monitoring", "6", "query_event_log, group_errors, message_status, throughput_stats, queue_depth, ..."],
        ])

    doc.add_heading("6.2 Tool Policies", level=2)
    doc.add_paragraph("ConfirmationGate: Mutating tools require explicit End User approval before execution.", style="List Bullet")
    doc.add_paragraph("ToolFilter: Strips framework waste tools (FileSystem, SQL, ShellTools) from the LLM catalog. Saves ~5K tokens per request.", style="List Bullet")

    # 7. Catalogs
    doc.add_page_break()
    doc.add_heading("7. Vector Catalogs", level=1)
    doc.add_paragraph(
        "Two semantic search catalogs index the IRIS class library so the agent can find "
        "relevant Business Hosts, adapters, and transformation classes by natural-language "
        "query.")
    add_figure(doc, img("11_catalogs.png"),
        "Catalogs tab showing search_ens (164 rows) and search_hs (58 rows) with kind breakdowns, rebuild controls, and test search panel")

    add_table(doc,
        ["Catalog", "Name", "Row Count", "Source"],
        [
            ["Ens.*", "search_ens", "164", "%Dictionary for Ens.Host, Ens.BusinessService, Ens.BusinessProcess, Ens.BusinessOperation, adapters"],
            ["HS.*", "search_hs", "58", "%Dictionary for HS.* transformation classes, FHIR mappers, SDA helpers"],
        ])

    doc.add_heading("7.1 Technical Implementation", level=2)
    doc.add_paragraph("Embedding model: FastEmbed (384-dimensional vectors, bundled with IRIS)", style="List Bullet")
    doc.add_paragraph("Vector storage: %AI.RAG.KnowledgeBase with HNSW index", style="List Bullet")
    doc.add_paragraph("Query path: %AI.ToolMgr.ExecuteTool(kbName, args) -- the only working path (SQL EMBEDDING() does not work with FastEmbed)", style="List Bullet")
    doc.add_paragraph("Document format: Curated prose descriptions, not raw method signatures", style="List Bullet")

    # 8. Transformation and Mapping Catalog
    doc.add_page_break()
    doc.add_heading("8. Transformation and Mapping Catalog", level=1)
    doc.add_paragraph(
        "A visual field-level mapping explorer showing how data flows between external "
        "formats (HL7 v2, FHIR R4, CDA, X12) through the SDA3 canonical model.")
    add_figure(doc, img("12_transforms_empty.png"),
        "Transforms tab initial state: format pair selection with 1,538 pre-computed rows")
    add_figure(doc, img("13_transforms_hl7_fhir.png"),
        "Transformation and Mapping Catalog with HL7 v2 to FHIR R4 selected, showing SDA3.Address mappings at sub-field level with IRIS class names and coverage filters")

    doc.add_heading("8.1 Data Flow Model", level=2)
    doc.add_paragraph("SDA3 is the universal pivot: all external formats map through it.")
    doc.add_paragraph("HL7 v2 (inbound) --> SDA3 (canonical) --> FHIR R4 (outbound)")
    doc.add_paragraph("Example: PID.11.3 City --> City --> city")

    doc.add_heading("8.2 Data Sources", level=2)
    doc.add_paragraph("HL7 to SDA3: Programmatic extraction from HS.Gateway.HL7.HL7ToSDA3 ObjectScript methods (not DTL)", style="List Bullet")
    doc.add_paragraph("SDA3 to FHIR: DTL class analysis via HS.FHIR.DTL.SDA3.vR4.* with backward-walk algorithm", style="List Bullet")
    doc.add_paragraph("Sub-field enrichment: Static lookup mapping composite HL7 types (XAD, XPN, CX, XTN) to component fields", style="List Bullet")

    # 9. Connection Management
    doc.add_page_break()
    doc.add_heading("9. Connection Management", level=1)
    doc.add_paragraph(
        "A multi-provider LLM connection manager that stores credentials securely and "
        "provides health-check (connection test) functionality. Managed by the AI Hub Admin.")
    add_figure(doc, img("10_connection_detail.png"),
        "Connection editor showing Bedrock configuration with provider, model, region, masked API key, and Test Connection button")

    add_table(doc,
        ["Provider", "Config Fields", "Secret Field"],
        [
            ["AWS Bedrock", "region, model", "AWS_BEARER_TOKEN_BEDROCK (Wallet)"],
            ["Anthropic", "model, base URL", "ANTHROPIC_API_KEY (Wallet)"],
            ["OpenAI", "model, base URL", "OPENAI_API_KEY (Wallet)"],
            ["Azure OpenAI", "endpoint, deployment, API version", "AZURE_OPENAI_API_KEY (Wallet)"],
            ["Google Gemini", "model, region", "GEMINI_API_KEY (Wallet)"],
            ["NVIDIA NIM", "model, base URL", "NIM_API_KEY (Wallet)"],
        ])

    doc.add_heading("9.1 Security Invariants", level=2)
    doc.add_paragraph("API keys are NEVER stored in SQL tables, globals, or source code", style="List Bullet")
    doc.add_paragraph("API keys are NEVER returned by any REST endpoint", style="List Bullet")
    doc.add_paragraph("API keys are NEVER logged in audit trail", style="List Bullet")
    doc.add_paragraph("The only storage location is the IRIS Secured Wallet", style="List Bullet")

    # 10. Audit & Security
    doc.add_page_break()
    doc.add_heading("10. Audit and Security", level=1)
    add_figure(doc, img("14_audit.png"),
        "Audit log showing all API requests with status codes, methods, paths, timestamps, users, and durations")

    doc.add_heading("10.1 Authentication", level=2)
    add_table(doc,
        ["Method", "Use Case"],
        [
            ["Basic auth", "Direct admin UI access (standalone mode)"],
            ["JWT Bearer", "Embedded access from Interop Editor (token passed via postMessage bridge)"],
        ])
    doc.add_paragraph(
        "All REST endpoints require authentication. No UI element is visible before login.")

    doc.add_heading("10.2 LLM Permission Delegation", level=2)
    doc.add_paragraph(
        "The agent must NEVER bypass the authenticated user's permissions. Every tool call "
        "is executed under the authenticated user's security context. If the user cannot "
        "perform an action through the Management Portal, the agent must not perform it "
        "either. Tools validate permissions via $System.Security.Check() before executing "
        "mutating operations.")

    doc.add_heading("10.3 Namespace Constraints", level=2)
    doc.add_paragraph(
        "Foundation namespaces (HSLIB, HSSYS, ENSLIB) are always read-only for the agent, "
        "regardless of user permissions. Non-Foundation namespaces respect the authenticated "
        "user's database-level permissions.")

    doc.add_heading("10.4 Interface Engineer vs Operator Privilege Separation", level=2)
    doc.add_paragraph(
        "Enforced at two levels: (1) Tool availability -- the AI Hub Admin binds different "
        "ToolSets to Interface Engineer-mode and Operator-mode agent configurations; (2) Permission "
        "validation -- every mutating tool checks the user's IRIS roles before executing.")

    doc.add_heading("10.5 Audit Log Fields", level=2)
    add_table(doc,
        ["Field", "Description"],
        [
            ["Timestamp", "When the request was received"],
            ["Username", "Authenticated IRIS user"],
            ["Namespace", "Active namespace at request time"],
            ["SessionId", "Browser session identifier"],
            ["Job", "IRIS job number"],
            ["Method", "GET, POST, PUT, DELETE"],
            ["Path", "/api/agentic/chat/stream"],
            ["StatusCode", "HTTP response status"],
            ["Duration", "End-to-end request time"],
            ["Kind", "registry, editor.agent, chat, namespace, etc."],
        ])

    doc.add_heading("10.6 Security Policies", level=2)
    doc.add_paragraph(
        "ConfirmationGate: Mutating operations require explicit End User approval. This is "
        "the last line of defense after permission checks and namespace validation.",
        style="List Bullet")
    doc.add_paragraph(
        "ToolFilter: Strips framework waste tools from the LLM catalog, preventing access "
        "to generic system tools that could bypass the permission model.",
        style="List Bullet")

    doc.add_heading("10.7 Source Control Integration", level=2)
    doc.add_paragraph(
        "Agent-created artifacts (productions, DTLs, BPLs, routing rules) integrate with "
        "Health Connect Cloud's %SourceControl hooks. The agent uses standard IRIS APIs "
        "($System.OBJ.Compile, %Dictionary) that trigger the hooks automatically. Changes "
        "are attributed to the authenticated user, not the agent's service account. The "
        "CI/CD pipeline (GitLab) validates, tests, and deploys changes across environments.")
    add_figure(doc, img("16_cicd_pipeline.png"),
        "CI/CD Pipeline Flow: agent-created artifacts flow through source control hooks to Git and CI/CD")

    # 11. Performance
    doc.add_page_break()
    doc.add_heading("11. Performance", level=1)
    add_table(doc,
        ["Metric", "Target", "How Achieved"],
        [
            ["First token latency", "< 2 seconds", "SSE streaming, no buffering"],
            ["Turn completion", "< 90 seconds", "Monitor: 60s deadline + 50K token budget"],
            ["Catalog search", "< 500ms", "Pre-built HNSW vector index, in-process query"],
            ["Field mapping load", "< 100ms", "Pre-computed SQL table (1,538 rows)"],
            ["Catalog rebuild", "< 30 seconds", "Batch %Dictionary walk + FastEmbed"],
            ["Concurrent users", "5 sessions", "Per-process agent instances, no shared state"],
            ["Token efficiency", "< 50K tokens/task", "ToolFilter, curated skills, no markdown overhead"],
        ])

    # 12. Implementation Summary
    doc.add_heading("12. Implementation Summary", level=1)
    add_table(doc,
        ["Component", "Classes", "Status"],
        [
            ["Agent", "4 classes (HealthInterop, Manager, Monitor, SkillLoader)", "Built"],
            ["MCP Servers", "5 classes (Base + 4 servers)", "Built"],
            ["ToolSets", "5 classes", "Built"],
            ["Tools", "7 classes, 118 tools", "Built"],
            ["Skills", "16 classes (Base + 15 domain skills)", "Built"],
            ["Data model", "6 persistent classes", "Built"],
            ["REST API", "1 dispatcher + 13 service classes", "Built"],
            ["Admin UI", "3 HTML + 2 JS + 2 CSS files", "Built"],
            ["Chat UI", "1 HTML + 1 JS + 1 CSS file", "Built"],
            ["Vector catalogs", "2 classes (Builder, Attach)", "Built"],
            ["Transformation and Mapping Catalog", "2 classes (TransformService, FieldMapping)", "Built"],
            ["Policies", "2 classes (ConfirmationGate, ToolFilter)", "Built"],
            ["Install hooks", "2 classes (CSPTimeoutPatch, InteropEditorPatch)", "Built"],
            ["IPM package", "module.xml", "Built"],
            ["Total", "61 ObjectScript classes, 8 web files, 100+ commits", ""],
        ])

    path = os.path.join(DOCS, "02_Technical_Build_Specification.docx")
    doc.save(path)
    print(f"  saved {path}")


# =============================================================================
# Document 3: Lessons Learned
# =============================================================================
def build_doc3():
    doc = Document()
    style_doc(doc)
    add_title_page(doc,
        "Agentic Health Interoperability",
        "Lessons Learned\nFindings from the experimental build of an Agentic AI Copilot on IRIS for Health")

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document captures the key lessons learned during the experimental build of "
        "the Agentic Health Interoperability project on InterSystems IRIS for Health 2026.2 "
        "with the %AI Framework (build 162.0). These findings cover framework bugs that "
        "required application-level extensions, performance optimizations using Vector "
        "Search, and strategies for reducing token consumption and improving response speed.")
    doc.add_paragraph(
        "The solution is built entirely on the %AI Framework primitives (Agent, MCP, "
        "ToolSet, Tool, Skill, KnowledgeBase). Where we extended the framework, it was to "
        "work around specific bugs -- not to replace framework functionality. These "
        "extensions are documented here as recommendations for InterSystems to address in "
        "future framework releases.")

    # 2. Bugs in the %AI Framework
    doc.add_page_break()
    doc.add_heading("2. %AI Framework Bugs and Extensions", level=1)
    doc.add_paragraph(
        "During development, we encountered three bugs in the %AI Framework that required "
        "application-level workarounds. In each case, the framework classes are untouched "
        "-- we built extension classes that work around the issues.")

    doc.add_heading("2.1 BUG: %AI.Agent.Skill %OnNew $ZF Marshaling Error", level=2)
    doc.add_paragraph("Severity: Blocker (prevents skill instantiation)")
    doc.add_paragraph(
        "When instantiating a skill via %New(), the parent %OnNew method passes a "
        "%DynamicObject to $ZF (the Foreign Function Interface call to the Rust LLM "
        "bridge). The $ZF call expects a JSON string, not an object reference, and throws "
        "a <FUNCTION> error.")
    doc.add_paragraph(
        "Extension: Created AgenticInterop.Skill.Base that overrides %OnNew to serialize "
        "the DynamicObject to a JSON string before the $ZF call using literal macro values "
        "(IrisLLMLibrary=1042, LLMBUILDSKILLFROMJSON=66). All 15 domain skills extend this base "
        "class instead of %AI.Agent.Skill directly.")
    doc.add_paragraph(
        "Recommendation: Fix %AI.Agent.Skill.%OnNew to serialize the DynamicObject to a "
        "JSON string before the $ZF call.")

    doc.add_heading("2.2 BUG: Bedrock Tool-Result Round-Trip Hang", level=2)
    doc.add_paragraph("Severity: Blocker (prevents multi-turn tool use with Bedrock)")
    doc.add_paragraph(
        "When the agent calls a tool and receives the result, the Rust bridge hangs "
        "indefinitely when sending the tool result back to the Bedrock Converse API. The "
        "hang occurs below the ObjectScript API surface -- the StreamChat() or Run() call "
        "never returns.")
    doc.add_paragraph(
        "Tested with Claude Sonnet 4 and Claude Haiku 3.5 via Bedrock (us-east-1), both "
        "cross-region and direct endpoints. All configurations reproduce the hang.")
    doc.add_paragraph(
        "Extension: Switched the runtime LLM provider from Bedrock to Anthropic direct. "
        "The same agent, tools, and skills work correctly -- only the connection "
        "configuration changes. No code modifications required.")
    doc.add_paragraph(
        "Recommendation: Investigate the Rust LLM bridge's Bedrock Converse API "
        "integration. WRC ticket recommended.")

    doc.add_heading("2.3 Extension: ToolFilter Policy (Framework Tool Cleanup)", level=2)
    doc.add_paragraph("Severity: Performance (wastes ~5K tokens per request)")
    doc.add_paragraph(
        "The %AI Framework exposes default tools (FileSystem, SQL, ShellTools) that are "
        "irrelevant for healthcare interoperability tasks. These tools consume LLM tokens "
        "and can cause the model to make irrelevant tool calls.")
    doc.add_paragraph(
        "Extension: AgenticInterop.Policy.ToolFilter strips framework-default tools "
        "before each LLM call, leaving only the 118 healthcare-specific tools.")
    doc.add_paragraph(
        "Recommendation: Add a ToolFilter-like policy to the %AI Framework defaults, or "
        "provide an opt-in mechanism for framework tools.")

    doc.add_heading("2.4 ISSUE: %FromJSON Instance Method Returns Empty String", level=2)
    doc.add_paragraph("Severity: Medium (causes silent data loss)")
    doc.add_paragraph(
        "Calling {}.%FromJSON(jsonString) returns '' (empty string) instead of a populated "
        "object. The instance form of %FromJSON appears non-functional on this IRIS build.")
    doc.add_paragraph(
        "Workaround: Use the class-method form: "
        "Set obj = ##class(%DynamicAbstractObject).%FromJSON(jsonString)")

    doc.add_heading("2.5 ISSUE: $get() on %DynamicObject Properties Throws <INVALID CLASS>", level=2)
    doc.add_paragraph("Severity: Medium (causes runtime errors in otherwise valid code)")
    doc.add_paragraph(
        "$get(dynamicObj.propertyName, defaultValue) throws <INVALID CLASS> when the "
        "property doesn't exist. ObjectScript $get doesn't handle %DynamicObject property "
        "access the way it handles local variables.")
    doc.add_paragraph(
        "Workaround: Use $select with %IsDefined: "
        "Set value = $select(obj.%IsDefined(\"propName\"): obj.propName, 1: defaultValue)")

    doc.add_heading("2.6 ISSUE: CSP UseSession Deadlock on REST Endpoints", level=2)
    doc.add_paragraph("Severity: High (causes 401 errors on second request in same browser session)")
    doc.add_paragraph(
        "When UseSession=1 (the default) on a %CSP.REST class, the CSP gateway validates "
        "the CSRF token on every request. When called from an iframe, the second request "
        "in the same browser session gets a 401 because the CSRF token validation races "
        "with the iframe's credential handling.")
    doc.add_paragraph("Workaround: Set UseSession=0 on every REST dispatch class.")

    doc.add_heading("2.7 ISSUE: %OpenId Returns Stale Data in Cross-Process Polling", level=2)
    doc.add_paragraph("Severity: High (causes infinite polling loops)")
    doc.add_paragraph(
        "%OpenId(id) uses a process-local OREF cache. In a polling loop where Process A "
        "writes a row and Process B polls with %OpenId, Process B's cache returns the stale "
        "pre-update state indefinitely. The poll never sees the update.")
    doc.add_paragraph(
        "Workaround: Use SQL queries for cross-process polling instead of %OpenId.")

    # 3. ObjectScript Language Gotchas
    doc.add_page_break()
    doc.add_heading("3. ObjectScript Language Gotchas", level=1)
    doc.add_paragraph(
        "These are not framework bugs but language behaviors that caused significant "
        "debugging time during development.")

    doc.add_heading("3.1 Numeric Comparison Operators on Strings", level=2)
    doc.add_paragraph(
        'ObjectScript >=, <=, >, < operators are NUMERIC, not string-based. The expression '
        'ch >= "A" && ch <= "Z" always evaluates to true because both sides coerce to 0.')
    doc.add_paragraph(
        "Solution: Use $ascii() for character range comparisons: "
        "Set code = $ascii(ch); If (code >= 65) && (code <= 90)")

    doc.add_heading("3.2 QUIT Inside try/catch Blocks", level=2)
    doc.add_paragraph(
        "quit value only works at the method top level. Inside a block (if/for/while/"
        "try/catch), quit exits the block without a return value. This causes silent "
        "method exits.")
    doc.add_paragraph(
        "Solution: Capture to a local variable and quit after the block, or use RETURN "
        "value which works from anywhere.")

    doc.add_heading("3.3 Comment Syntax Inside Method Bodies", level=2)
    doc.add_paragraph(
        "// comments work at the class level but cause #1002 compile errors inside "
        "ClassMethod/Method bodies on some IRIS versions.")
    doc.add_paragraph(
        "Solution: Use ; (semicolon) for comments inside method bodies.")

    # 4. Vector Search
    doc.add_page_break()
    doc.add_heading("4. Vector Search: Improving Velocity", level=1)

    add_figure(doc, img("11_catalogs.png"),
        "Vector catalogs admin panel showing indexed class counts and kind breakdowns")

    doc.add_heading("4.1 The Embedding Quality Problem", level=2)
    doc.add_paragraph(
        "The initial catalog builder indexed every class by dumping its full "
        "%Dictionary.ClassDefinition content -- including auto-generated accessor methods, "
        "storage definitions, and parameter boilerplate. In a 384-dimensional embedding "
        "space, this structural noise drowned out the semantic signal (class description, "
        "purpose, key behaviors).")
    doc.add_paragraph(
        'Result: Searches for "HL7 TCP service" returned generic base classes instead of '
        "EnsLib.HL7.Service.TCPService because the relevant description text was a tiny "
        "fraction of the indexed content.")

    doc.add_heading("4.2 The Fix: Curated Prose Over Raw Metadata", level=2)
    add_table(doc,
        ["Before (Noisy)", "After (Curated)"],
        [
            ["Full class dump including storage, indices, XData", "Class name + description + superclass + key parameters"],
            ["Auto-generated accessor methods (Get/Set/IsValid)", "Removed entirely"],
            ["Parameter definitions with internal flags", "Only parameters the user would configure"],
            ["Inherited method signatures from 5 levels of superclass", "Only overridden methods with their descriptions"],
        ])
    doc.add_paragraph(
        'Impact: Search relevance improved dramatically. "HL7 TCP inbound service" now '
        "returns EnsLib.HL7.Service.TCPService as the top result with a cosine similarity "
        "of 0.85+.")

    doc.add_heading("4.3 Query Path Discovery", level=2)
    doc.add_paragraph(
        "The documented EMBEDDING() SQL function does NOT work with the bundled FastEmbed "
        "embedding model. After extensive testing, the only working query path is "
        "%AI.ToolMgr.ExecuteTool(kbName, args), which routes through the "
        "%AI.RAG.KnowledgeBase internal search pipeline.")

    # 5. Reducing Token Consumption
    doc.add_page_break()
    doc.add_heading("5. Reducing Token Consumption and Improving Speed", level=1)

    doc.add_heading("5.1 The Problem", level=2)
    doc.add_paragraph(
        "A naive agent configuration with all framework tools exposed consumed 15K+ tokens "
        "per request just for the tool catalog. Complex multi-turn tasks could exceed 100K "
        "tokens before producing a useful result. Response times were 30-60 seconds per turn.")

    doc.add_heading("5.2 Strategy 1: ToolFilter Policy (saved ~5K tokens per request)", level=2)
    doc.add_paragraph(
        "The %AI Framework exposes default tools (FileSystem, SQL, ShellTools) that are "
        "irrelevant for healthcare interoperability tasks. The ToolFilter strips these "
        "before each LLM call.")
    doc.add_paragraph("Before: 57 tools in the LLM catalog (~15K tokens)")
    doc.add_paragraph("After: 42 tools in the LLM catalog (~10K tokens)")

    doc.add_heading("5.3 Strategy 2: Concise Tool Descriptions (saved ~2K tokens)", level=2)
    doc.add_paragraph(
        "Early tool descriptions were verbose explanations. We rewrote them as imperative "
        "contracts: verb, scope, side effects, expected inputs. The concise versions convey "
        "the same information in 80% fewer tokens.")

    doc.add_heading("5.4 Strategy 3: No Markdown Formatting (eliminated rendering bugs)", level=2)
    doc.add_paragraph(
        "The chat UI renders plain text. Markdown bold (**text**) and italic (*text*) "
        "appear as literal asterisks. We instructed the agent to use plain prose with line "
        "breaks and - bullets.")

    doc.add_heading("5.5 Strategy 4: Monitor Callback with Token Budget", level=2)
    doc.add_paragraph(
        "AgenticInterop.Agent.Monitor enforces a 50K token budget per turn. This prevents "
        "infinite tool-call loops, excessive catalog searches, and runaway conversation depth.")

    doc.add_heading("5.6 Strategy 5: Breaking Complex Tasks into Multiple Short Turns", level=2)
    doc.add_paragraph("Instead of one massive prompt, the agent works in phases:")
    doc.add_paragraph("1. Research phase: Search catalogs, introspect schemas (read-only, fast)")
    doc.add_paragraph("2. Proposal phase: Present a plan, ask for approval (no tool calls)")
    doc.add_paragraph("3. Build phase: Execute step by step (mutating tools, confirmation gates)")
    doc.add_paragraph("4. Validation phase: Run PostBuildValidation, send test messages")
    doc.add_paragraph("5. Report phase: Summarize what was done")
    doc.add_paragraph("Each phase is a short LLM turn (< 30 seconds, < 20K tokens).")

    doc.add_heading("5.7 Results", level=2)
    add_table(doc,
        ["Metric", "Before Optimization", "After Optimization"],
        [
            ["Tool catalog tokens", "~15K", "~10K"],
            ["System prompt tokens", "~5K", "~2K"],
            ["Simple query (list productions)", "~25K total tokens", "~15K total tokens"],
            ["Complex task (build production)", "~120K total tokens", "~50K total tokens"],
            ["Simple query latency", "8-12 seconds", "3-5 seconds"],
            ["Complex task latency", "90+ seconds (often timeout)", "45-60 seconds (across 3-4 turns)"],
        ])

    # 6. Additional Findings
    doc.add_page_break()
    doc.add_heading("6. Additional Findings", level=1)

    doc.add_heading("6.1 SDA3 as the Universal Pivot", level=2)
    doc.add_paragraph(
        "The most powerful architectural insight for the Transformation and Mapping Catalog "
        "was that SDA3 is the universal pivot format in IRIS for Health. Every external "
        "format (HL7 v2, FHIR R4, CDA, X12) maps through SDA3. This means field-level "
        "mappings can be pre-computed as a three-column join: Source -> SDA3 -> Target.")
    add_figure(doc, img("13_transforms_hl7_fhir.png"),
        "Transformation and Mapping Catalog showing the SDA3 pivot pattern: HL7 v2 fields map into SDA3, SDA3 maps out to FHIR R4")

    doc.add_heading("6.2 HL7 v2 Programmatic Mappings Are NOT DTL", level=2)
    doc.add_paragraph(
        "The HL7 -> SDA3 direction is implemented in ObjectScript methods "
        "(HS.Gateway.HL7.HL7ToSDA3), not in DTL classes. This required a completely "
        "different extraction approach -- parsing ObjectScript source for Set statements "
        "targeting SDA properties.")

    doc.add_heading("6.3 Vanilla JS vs. React for the Admin UI", level=2)
    doc.add_paragraph(
        "The original plan called for React 18 + TypeScript + Vite. During implementation, "
        "we switched to vanilla JavaScript because: the IRIS CSP gateway serves static "
        "files (no build step), the admin UI is a configuration tool (not complex "
        "interactive), and it eliminates framework bundle size (React + ReactDOM = 130KB "
        "gzipped).")

    doc.add_heading("6.4 The Overlay Pattern for Configuration Persistence", level=2)
    doc.add_paragraph(
        "The 'class-as-data' model created a tension between shipped defaults and user "
        "customizations. IPM zpm load recompiles shipped classes, overwriting manual edits. "
        "The overlay pattern resolves this: shipped class parameters define defaults, "
        "override tables store AI Hub Admin edits, at build time the overlay merges "
        "overrides on top, and 'Reset to defaults' deletes the row.")

    # 7. Recommendations
    doc.add_page_break()
    doc.add_heading("7. Summary of Recommendations for InterSystems", level=1)
    add_table(doc,
        ["#", "Category", "Recommendation", "Impact"],
        [
            ["1", "Framework bug", "Fix %AI.Agent.Skill.%OnNew $ZF marshaling", "Eliminates need for Skill.Base extension"],
            ["2", "Framework bug", "Fix Bedrock tool-result round-trip hang", "Enables Bedrock as a production runtime provider"],
            ["3", "API fix", "Fix %DynamicObject.%FromJSON() instance method", "Eliminates silent data loss"],
            ["4", "API fix", "Document that $get() doesn't work on %DynamicObject properties", "Prevents runtime errors"],
            ["5", "Documentation", "Document %AI.ToolMgr.ExecuteTool() as the RAG query path", "SQL EMBEDDING() doesn't work with FastEmbed"],
            ["6", "Performance", "Add ToolFilter-like policy to framework defaults", "Prevents token waste from generic tools"],
            ["7", "Performance", "Expose token usage metrics from Rust bridge", "Enables application-level token budgeting"],
            ["8", "CSP", "Add Cache-Control: no-cache option for development web apps", "Eliminates static file caching during development"],
        ])

    path = os.path.join(DOCS, "03_Lessons_Learned.docx")
    doc.save(path)
    print(f"  saved {path}")


if __name__ == "__main__":
    print("Building .docx documents with screenshots...")
    build_doc1()
    build_doc2()
    build_doc3()
    print("\nAll three documents built.")
